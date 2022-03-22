from django.http import HttpResponse,HttpResponseRedirect
import numpy as np
import xlrd
from docx import Document
import math
from report.models import *
from report.effectnum import *
from datetime import datetime
import re

def Matrix_effect_fileread(files,reportinfo,project,platform,manufacturers,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB, Number_of_compounds):

    # 第一步:后台数据抓取（回收率上下限）
    id1 = Special.objects.get(project=project).id  
    id2 = Matrixeffectspecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if Matrixeffectspecialmethod.objects.filter(matrixeffectspecial=id2): 
        maxBias=Matrixeffectspecialmethod.objects.get(matrixeffectspecial=id2).bias #最大允许偏差
   
    else:
        general_1 = General.objects.get(name="通用性项目")
        matrixeffect_general = Matrixeffectgeneral.objects.get(general=general_1)
        maxBias=Matrixeffectgeneralmethod.objects.get(matrixeffectgeneral=matrixeffect_general).bias

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)

    #  第二步:开始文件读取

    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典Matrix_effect_dict,数据格式如下：
    print(Matrix_effect_dict):
    {"化合物1":{"A":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3],
    "B":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3],...,
    "AL":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3],...}
     "化合物2":{"A":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3],
     "B":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3],...,
     "AL":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3],...}
    }
    '''

    # 创建第二步需要生成的结果容器
    Matrix_effect_dict={} # 定义需要生成的字典
    
    for file in files:
        if platform=="液质":
            if manufacturers =="Agilent":
                csv_file = file.seek(0)  # 此网址查找到的答案:https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)): 
                    if len(lines[i])!=0:
                        lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                    else:
                        lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i] #最后一行如为空行，则删除该元素

                # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                norm=[] #化合物列表
                for j in range(len(lines[0])):  #从第一行开始
                    if "-Q Results" in lines[0][j]:
                        if lines[0][j].split("-Q")[0][0]!='"':  # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除  
                            norm.append(lines[0][j].split("-Q")[0])
                        else:
                            norm.append(lines[0][j].split("-Q")[0][1:])
      
                # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                nameindex=0  #实验号索引
                Areaindex=[]
                ISAreaindex=[]
                AreaRatioindex=[]
                for j in range(len(lines[1])):  #从第二行开始       
                    if lines[1][j] == "Sample Name" :
                        nameindex=j
                    elif lines[1][j]  == "Area" :
                        Areaindex.append(j)
                    elif lines[1][j]  == "IS Area" :
                        ISAreaindex.append(j)
                    elif lines[1][j]  == "Area Ratio" :
                        AreaRatioindex.append(j)
        
                # 循环每个化合物
                for k in range(len(norm)):  

                    # 每个化合物的字典
                    normdict={}  

                    # 病人基质样本列表(PS-A-1) 
                    PSlist=[]
                    for i in range(len(lines)): 
                        if "PS" in lines[i][nameindex] and lines[i][nameindex][0:4] not in PSlist:
                            PSlist.append(lines[i][nameindex][0:4])

                    # 标准溶液列表
                    STDlist=[]
                    for i in range(len(lines)): 
                        if "STD" in lines[i][nameindex] and lines[i][nameindex][0:5] not in STDlist:
                            STDlist.append(lines[i][nameindex][0:5])

                    # 混合样本列表
                    MIXlist=[]
                    for i in range(len(lines)): 
                        if "MIX" in lines[i][nameindex] and lines[i][nameindex][0:6] not in MIXlist:
                            MIXlist.append(lines[i][nameindex][0:6])

                    # 循环病人基质样本列表
                    for j in PSlist:
                        middlelist=[] 
                        for i in range(len(lines)):
                            if j in lines[i][nameindex]:
                                middlelist.append(effectnum(lines[i][Areaindex[k]],digits))
                                middlelist.append(effectnum(lines[i][ISAreaindex[k]],digits))
                                middlelist.append(new_round(float(lines[i][AreaRatioindex[k]]),3))
                        normdict[j]=middlelist

                    # 循环标准溶液列表
                    for j in STDlist:
                        middlelist=[] 
                        for i in range(len(lines)):
                            if j in lines[i][nameindex]:
                                middlelist.append(effectnum(lines[i][Areaindex[k]],digits))
                                middlelist.append(effectnum(lines[i][ISAreaindex[k]],digits))
                                middlelist.append(new_round(float(lines[i][AreaRatioindex[k]]),3))
                        normdict[j]=middlelist

                    # 循环混合样本列表
                    for j in MIXlist:
                        middlelist=[] 
                        for i in range(len(lines)):
                            if j in lines[i][nameindex]:
                                middlelist.append(effectnum(lines[i][Areaindex[k]],digits))
                                middlelist.append(effectnum(lines[i][ISAreaindex[k]],digits))
                                middlelist.append(new_round(float(lines[i][AreaRatioindex[k]]),3))
                        normdict[j]=middlelist

                    Matrix_effect_dict[norm[k]]=normdict

            elif manufacturers =="Waters":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                norm=[] #化合物列表
                norm_row=[] #化合物所在行
                for j in range(nrows):
                    for i in PTnorm:
                        if i in str(file_data.row_values(j)[0]):
                            norm.append(i)
                            norm_row.append(j)

                nameindex=0  #实验号索引
                Areaindex=0
                ISAreaindex=0
                AreaRatioindex=0
                for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    if file_data.row_values(norm_row[0]+2)[i]=="Name":
                        nameindex=i
                    elif file_data.row_values(norm_row[0]+2)[i]=="Area":
                        Areaindex=i
                    elif file_data.row_values(norm_row[0]+2)[i]=="IS Area":
                        ISAreaindex=i
                    elif file_data.row_values(norm_row[0]+2)[i]=="Area Ratio":
                        AreaRatioindex=i

                for j in range(len(norm)):                       
                    A=[]
                    B=[]
                    C=[]
                    D=[]
                    E=[]
                    F=[]
                    L=[]
                    H=[]
                    AL=[]
                    AH=[]
                    BL=[]
                    BH=[]
                    CL=[]
                    CH=[]
                    DL=[]
                    DH=[]
                    EL=[]
                    EH=[]
                    FL=[]
                    FH=[]
                    group_Me={} #每个化合物的字典
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]): 
                            if file_data.row_values(i)[nameindex]=="A":
                                A.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                A.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                A.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="B":                 
                                B.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                B.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                B.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="C":                 
                                C.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                C.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                C.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="D":                 
                                D.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                D.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                D.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="E":                 
                                E.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                E.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                E.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="F":                 
                                F.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                F.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                F.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="L":                 
                                L.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                L.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                L.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="H":                 
                                H.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                H.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                H.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="AL":                 
                                AL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                AL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="AH":                 
                                AH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                AH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="BL":                 
                                BL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                BL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="BH":                 
                                BH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                BH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="CL":                 
                                CL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                CL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="CH":                 
                                CH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                CH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="DL":                 
                                DL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                DL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="DH":                 
                                DH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                DH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="EL":                 
                                EL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                EL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="EH":                 
                                EH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                EH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="FL":                 
                                FL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                FL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="FH":                 
                                FH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                FH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        
                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],nrows): 
                            if file_data.row_values(i)[nameindex]=="A":
                                A.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                A.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                A.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="B":                 
                                B.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                B.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                B.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="C":                 
                                C.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                C.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                C.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="D":                 
                                D.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                D.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                D.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="E":                 
                                E.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                E.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                E.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="F":                 
                                F.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                F.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                F.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="L":                 
                                L.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                L.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                L.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="H":                 
                                H.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                H.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                H.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="AL":                 
                                AL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                AL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="AH":                 
                                AH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                AH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="BL":                 
                                BL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                BL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="BH":                 
                                BH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                BH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="CL":                 
                                CL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                CL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="CH":                 
                                CH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                CH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="DL":                 
                                DL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                DL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="DH":                 
                                DH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                DH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="EL":                 
                                EL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                EL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="EH":                 
                                EH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                EH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="FL":                 
                                FL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                FL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                            elif file_data.row_values(i)[nameindex]=="FH":                 
                                FH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                FH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))

                    group_Me["A"]=A
                    group_Me["B"]=B
                    group_Me["C"]=C
                    group_Me["D"]=D
                    group_Me["E"]=E
                    group_Me["F"]=F
                    group_Me["L"]=L
                    group_Me["H"]=H
                    group_Me["AL"]=AL
                    group_Me["AH"]=AH
                    group_Me["BL"]=BL
                    group_Me["BH"]=BH
                    group_Me["CL"]=CL
                    group_Me["CH"]=CH
                    group_Me["DL"]=DL
                    group_Me["DH"]=DH
                    group_Me["EL"]=EL
                    group_Me["EH"]=EH
                    group_Me["FL"]=FL
                    group_Me["FH"]=FH

                    Matrix_effect_dict[norm[j]]=group_Me

            elif manufacturers =="Thermo":
                Thermo = Special.objects.get(project=project) 
                pt_special = PTspecial.objects.get(special=Thermo)
                pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
                PTnorm=[] # 待测物质列表
                for i in pt_accept:
                    PTnorm.append(i.norm)

                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm=[] #Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                sheetindex=[] #需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    if data.sheet_names()[index] in PTnorm:
                        norm.append(data.sheet_names()[index])
                        sheetindex.append(index)

                # 循环读取每个sheet工作表,即为每个化合物的表
                for index in range(len(sheetindex)):
                    print(index)
                    file_data = data.sheets()[sheetindex[index]]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    #第一行确定samplename和浓度所在列
                    nameindex=0
                    Areaindex=0
                    ISAreaindex=0
                    AreaRatioindex=0
                    for i in range(len(file_data.row_values(0))):  
                        if file_data.row_values(0)[i]=="Compound":
                            nameindex=i
                        elif file_data.row_values(0)[i]=="Area":
                            Areaindex=i
                        elif file_data.row_values(0)[i]=="ISTD Response":
                            ISAreaindex=i
                        elif file_data.row_values(0)[i]=="Response Ratio":
                            AreaRatioindex=i

                    A=[]
                    B=[]
                    C=[]
                    D=[]
                    E=[]
                    F=[]
                    L=[]
                    H=[]
                    AL=[]
                    AH=[]
                    BL=[]
                    BH=[]
                    CL=[]
                    CH=[]
                    DL=[]
                    DH=[]
                    EL=[]
                    EH=[]
                    FL=[]
                    FH=[]
                    group_Me={} #每个化合物的字典
                    for i in range(nrows): 
                        if file_data.row_values(i)[nameindex]=="A":
                            A.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            A.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            A.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="B":                 
                            B.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            B.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            B.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="C":                 
                            C.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            C.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            C.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="D":                 
                            D.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            D.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            D.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="E":                 
                            E.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            E.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            E.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="F":                 
                            F.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            F.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            F.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="L":                 
                            L.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            L.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            L.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="H":                 
                            H.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            H.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            H.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="AL":                 
                            AL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            AL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            AL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="AH":                 
                            AH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            AH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            AH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="BL":                 
                            BL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            BL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            BL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="BH":                 
                            BH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            BH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            BH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="CL":                 
                            CL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            CL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            CL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="CH":                 
                            CH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            CH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            CH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="DL":                 
                            DL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            DL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            DL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="DH":                 
                            DH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            DH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            DH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="EL":                 
                            EL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            EL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            EL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="EH":                 
                            EH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            EH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            EH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="FL":                 
                            FL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            FL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            FL.append(new_round(file_data.row_values(i)[AreaRatioindex],3))
                        elif file_data.row_values(i)[nameindex]=="FH":                 
                            FH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                            FH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                            FH.append(new_round(file_data.row_values(i)[AreaRatioindex],3))

                    group_Me["A"]=A
                    group_Me["B"]=B
                    group_Me["C"]=C
                    group_Me["D"]=D
                    group_Me["E"]=E
                    group_Me["F"]=F
                    group_Me["L"]=L
                    group_Me["H"]=H
                    group_Me["AL"]=AL
                    group_Me["AH"]=AH
                    group_Me["BL"]=BL
                    group_Me["BH"]=BH
                    group_Me["CL"]=CL
                    group_Me["CH"]=CH
                    group_Me["DL"]=DL
                    group_Me["DH"]=DH
                    group_Me["EL"]=EL
                    group_Me["EH"]=EH
                    group_Me["FL"]=FL
                    group_Me["FH"]=FH

                    Matrix_effect_dict[norm[index]]=group_Me

            elif manufacturers =="岛津":
                content= []
                for line in file:
                    content.append(line.decode("GB2312").replace("\r\n", "").split("\t"))

                nameindex=0  #实验号索引
                Areaindex=0
                ISAreaindex=0
                AreaRatioindex=0
                norm=[] #化合物列表
                norm_row=[] #化合物所在行
                
                for i in range(len(content[2])):  #第三行确定samplename和浓度所在列
                    if content[2][i]=="数据文件名":
                        nameindex=i 
                    elif content[2][i]=="Area":
                        Areaindex=i 
                    elif content[2][i]=="IS Area":
                        ISAreaindex=i 
                    elif content[2][i]=="Area Ratio":
                        AreaRatioindex=i 

                for i in range(len(content)): 
                    if content[i][0]=="Name": #如果某一行第一列为"Name"，则该行第二列为化合物名称
                        norm.append(content[i][1])
                        norm_row.append(i)

                for j in range(len(norm)):                       
                    A=[]
                    B=[]
                    C=[]
                    D=[]
                    E=[]
                    F=[]
                    L=[]
                    H=[]
                    AL=[]
                    AH=[]
                    BL=[]
                    BH=[]
                    CL=[]
                    CH=[]
                    DL=[]
                    DH=[]
                    EL=[]
                    EH=[]
                    FL=[]
                    FH=[]
                    group_Me={} #每个化合物的字典
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]): 
                            if content[i][nameindex]=="A":
                                A.append(effectnum(content[i][Areaindex],digits))
                                A.append(effectnum(content[i][ISAreaindex],digits))
                                A.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="B":                 
                                B.append(effectnum(content[i][Areaindex],digits))
                                B.append(effectnum(content[i][ISAreaindex],digits))
                                B.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="C":                 
                                C.append(effectnum(content[i][Areaindex],digits))
                                C.append(effectnum(content[i][ISAreaindex],digits))
                                C.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="D":                 
                                D.append(effectnum(content[i][Areaindex],digits))
                                D.append(effectnum(content[i][ISAreaindex],digits))
                                D.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="E":                 
                                E.append(effectnum(content[i][Areaindex],digits))
                                E.append(effectnum(content[i][ISAreaindex],digits))
                                E.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="F":                 
                                F.append(effectnum(content[i][Areaindex],digits))
                                F.append(effectnum(content[i][ISAreaindex],digits))
                                F.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="L":                 
                                L.append(effectnum(content[i][Areaindex],digits))
                                L.append(effectnum(content[i][ISAreaindex],digits))
                                L.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="H":                 
                                H.append(effectnum(content[i][Areaindex],digits))
                                H.append(effectnum(content[i][ISAreaindex],digits))
                                H.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="AL":                 
                                AL.append(effectnum(content[i][Areaindex],digits))
                                AL.append(effectnum(content[i][ISAreaindex],digits))
                                AL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="AH":                 
                                AH.append(effectnum(content[i][Areaindex],digits))
                                AH.append(effectnum(content[i][ISAreaindex],digits))
                                AH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="BL":                 
                                BL.append(effectnum(content[i][Areaindex],digits))
                                BL.append(effectnum(content[i][ISAreaindex],digits))
                                BL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="BH":                 
                                BH.append(effectnum(content[i][Areaindex],digits))
                                BH.append(effectnum(content[i][ISAreaindex],digits))
                                BH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="CL":                 
                                CL.append(effectnum(content[i][Areaindex],digits))
                                CL.append(effectnum(content[i][ISAreaindex],digits))
                                CL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="CH":                 
                                CH.append(effectnum(content[i][Areaindex],digits))
                                CH.append(effectnum(content[i][ISAreaindex],digits))
                                CH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="DL":                 
                                DL.append(effectnum(content[i][Areaindex],digits))
                                DL.append(effectnum(content[i][ISAreaindex],digits))
                                DL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="DH":                 
                                DH.append(effectnum(content[i][Areaindex],digits))
                                DH.append(effectnum(content[i][ISAreaindex],digits))
                                DH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="EL":                 
                                EL.append(effectnum(content[i][Areaindex],digits))
                                EL.append(effectnum(content[i][ISAreaindex],digits))
                                EL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="EH":                 
                                EH.append(effectnum(content[i][Areaindex],digits))
                                EH.append(effectnum(content[i][ISAreaindex],digits))
                                EH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="FL":                 
                                FL.append(effectnum(content[i][Areaindex],digits))
                                FL.append(effectnum(content[i][ISAreaindex],digits))
                                FL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="FH":                 
                                FH.append(effectnum(content[i][Areaindex],digits))
                                FH.append(effectnum(content[i][ISAreaindex],digits))
                                FH.append(new_round(float(content[i][AreaRatioindex]),3))
                        
                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],len(content)): 
                            if content[i][nameindex]=="A":
                                A.append(effectnum(content[i][Areaindex],digits))
                                A.append(effectnum(content[i][ISAreaindex],digits))
                                A.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="B":                 
                                B.append(effectnum(content[i][Areaindex],digits))
                                B.append(effectnum(content[i][ISAreaindex],digits))
                                B.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="C":                 
                                C.append(effectnum(content[i][Areaindex],digits))
                                C.append(effectnum(content[i][ISAreaindex],digits))
                                C.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="D":                 
                                D.append(effectnum(content[i][Areaindex],digits))
                                D.append(effectnum(content[i][ISAreaindex],digits))
                                D.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="E":                 
                                E.append(effectnum(content[i][Areaindex],digits))
                                E.append(effectnum(content[i][ISAreaindex],digits))
                                E.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="F":                 
                                F.append(effectnum(content[i][Areaindex],digits))
                                F.append(effectnum(content[i][ISAreaindex],digits))
                                F.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="L":                 
                                L.append(effectnum(content[i][Areaindex],digits))
                                L.append(effectnum(content[i][ISAreaindex],digits))
                                L.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="H":                 
                                H.append(effectnum(content[i][Areaindex],digits))
                                H.append(effectnum(content[i][ISAreaindex],digits))
                                H.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="AL":                 
                                AL.append(effectnum(content[i][Areaindex],digits))
                                AL.append(effectnum(content[i][ISAreaindex],digits))
                                AL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="AH":                 
                                AH.append(effectnum(content[i][Areaindex],digits))
                                AH.append(effectnum(content[i][ISAreaindex],digits))
                                AH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="BL":                 
                                BL.append(effectnum(content[i][Areaindex],digits))
                                BL.append(effectnum(content[i][ISAreaindex],digits))
                                BL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="BH":                 
                                BH.append(effectnum(content[i][Areaindex],digits))
                                BH.append(effectnum(content[i][ISAreaindex],digits))
                                BH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="CL":                 
                                CL.append(effectnum(content[i][Areaindex],digits))
                                CL.append(effectnum(content[i][ISAreaindex],digits))
                                CL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="CH":                 
                                CH.append(effectnum(content[i][Areaindex],digits))
                                CH.append(effectnum(content[i][ISAreaindex],digits))
                                CH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="DL":                 
                                DL.append(effectnum(content[i][Areaindex],digits))
                                DL.append(effectnum(content[i][ISAreaindex],digits))
                                DL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="DH":                 
                                DH.append(effectnum(content[i][Areaindex],digits))
                                DH.append(effectnum(content[i][ISAreaindex],digits))
                                DH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="EL":                 
                                EL.append(effectnum(content[i][Areaindex],digits))
                                EL.append(effectnum(content[i][ISAreaindex],digits))
                                EL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="EH":                 
                                EH.append(effectnum(content[i][Areaindex],digits))
                                EH.append(effectnum(content[i][ISAreaindex],digits))
                                EH.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="FL":                 
                                FL.append(effectnum(content[i][Areaindex],digits))
                                FL.append(effectnum(content[i][ISAreaindex],digits))
                                FL.append(new_round(float(content[i][AreaRatioindex]),3))
                            elif content[i][nameindex]=="FH":                 
                                FH.append(effectnum(content[i][Areaindex],digits))
                                FH.append(effectnum(content[i][ISAreaindex],digits))
                                FH.append(new_round(float(content[i][AreaRatioindex]),3))

                    group_Me["A"]=A
                    group_Me["B"]=B
                    group_Me["C"]=C
                    group_Me["D"]=D
                    group_Me["E"]=E
                    group_Me["F"]=F
                    group_Me["L"]=L
                    group_Me["H"]=H
                    group_Me["AL"]=AL
                    group_Me["AH"]=AH
                    group_Me["BL"]=BL
                    group_Me["BH"]=BH
                    group_Me["CL"]=CL
                    group_Me["CH"]=CH
                    group_Me["DL"]=DL
                    group_Me["DH"]=DH
                    group_Me["EL"]=EL
                    group_Me["EH"]=EH
                    group_Me["FL"]=FL
                    group_Me["FH"]=FH

                    Matrix_effect_dict[norm[j]]=group_Me
            
            elif manufacturers =="AB":
                # 定义化合物列表，列表统一命名为norm
                norm = normAB

                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = [] 

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                tables = file_data.tables  # 获取文件中的表格集

                # 循环定量表格的索引
                for k in range(len(tableindex)):
                    # 获取文件中的定量表格
                    tablequantify = tables[tableindex[k]] 

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表

                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    # AB厂家需自行计算Area Ratio值
                    nameindex=0
                    Areaindex=0
                    ISAreaindex=0

                    # 读取表格的第一行的单元格,判断实验号,峰面积和内标峰面积的索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name" :
                            nameindex=i
                        elif "Area" in rowdatagatherlist[0][i]  and "IS" not in rowdatagatherlist[0][i]:
                            Areaindex=i
                        elif "IS Area" in rowdatagatherlist[0][i]:
                            ISAreaindex=i

                    # 每个化合物的字典
                    normdict={}  

                    # 病人基质样本列表(PS-A-1) 
                    PSlist=[]
                    for i in range(len(rowdatagatherlist)): 
                        if "PS" in rowdatagatherlist[i][nameindex] and rowdatagatherlist[i][nameindex][0:4] not in PSlist:
                            PSlist.append(rowdatagatherlist[i][nameindex][0:4])

                    # 标准溶液列表
                    STDlist=[]
                    for i in range(len(rowdatagatherlist)): 
                        if "STD" in rowdatagatherlist[i][nameindex] and rowdatagatherlist[i][nameindex][0:5] not in STDlist:
                            STDlist.append(rowdatagatherlist[i][nameindex][0:5])

                    # 混合样本列表
                    MIXlist=[]
                    for i in range(len(rowdatagatherlist)): 
                        if "MIX" in rowdatagatherlist[i][nameindex] and rowdatagatherlist[i][nameindex][0:6] not in MIXlist:
                            MIXlist.append(rowdatagatherlist[i][nameindex][0:6])

                    # 循环病人基质样本列表
                    for j in PSlist:
                        middlelist=[] 
                        for i in range(len(rowdatagatherlist)):
                            if j in rowdatagatherlist[i][nameindex]:
                                middlelist.append(effectnum(rowdatagatherlist[i][Areaindex],digits))
                                middlelist.append(effectnum(rowdatagatherlist[i][ISAreaindex],digits))

                                divisor=float(effectnum(rowdatagatherlist[i][Areaindex],digits))
                                dividend=float(effectnum(rowdatagatherlist[i][ISAreaindex],digits))

                                middlelist.append(new_round(divisor/dividend,3))
                        normdict[j]=middlelist

                    # 循环标准溶液列表
                    for j in STDlist:
                        middlelist=[] 
                        for i in range(len(rowdatagatherlist)):
                            if j in rowdatagatherlist[i][nameindex]:
                                middlelist.append(effectnum(rowdatagatherlist[i][Areaindex],digits))
                                middlelist.append(effectnum(rowdatagatherlist[i][ISAreaindex],digits))

                                divisor=float(effectnum(rowdatagatherlist[i][Areaindex],digits))
                                dividend=float(effectnum(rowdatagatherlist[i][ISAreaindex],digits))

                                middlelist.append(new_round(divisor/dividend,3))

                        normdict[j]=middlelist

                    # 循环混合样本列表
                    for j in MIXlist:
                        middlelist=[] 
                        for i in range(len(rowdatagatherlist)):
                            if j in rowdatagatherlist[i][nameindex]:
                                middlelist.append(effectnum(rowdatagatherlist[i][Areaindex],digits))
                                middlelist.append(effectnum(rowdatagatherlist[i][ISAreaindex],digits))
                                
                                divisor=float(effectnum(rowdatagatherlist[i][Areaindex],digits))
                                dividend=float(effectnum(rowdatagatherlist[i][ISAreaindex],digits))

                                middlelist.append(new_round(divisor/dividend,3))
                        normdict[j]=middlelist

                    Matrix_effect_dict[norm[k]]=normdict

        elif platform=="液相":
            if manufacturers =="Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                norm=[] #化合物列表
                norm_row=[] #化合物所在行
                for j in range(nrows):
                    if file_data.row_values(j)[0] == "化合物:" :  #如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex=0  #实验号索引
                Areaindex=0
                ISAreaindex=0
                for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    if file_data.row_values(norm_row[0]+2)[i]=="样品名称":
                        nameindex=i
                    elif file_data.row_values(norm_row[0]+2)[i]=="Area":
                        Areaindex=i
                    elif file_data.row_values(norm_row[0]+2)[i]=="IS Area":
                        ISAreaindex=i

                for j in range(len(norm)):                       
                    A=[]
                    B=[]
                    C=[]
                    D=[]
                    E=[]
                    F=[]
                    L=[]
                    H=[]
                    AL=[]
                    AH=[]
                    BL=[]
                    BH=[]
                    CL=[]
                    CH=[]
                    DL=[]
                    DH=[]
                    EL=[]
                    EH=[]
                    FL=[]
                    FH=[]
                    group_Me={} #每个化合物的字典
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]): 
                            if file_data.row_values(i)[nameindex]=="A":
                                A.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                A.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                A.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="B":                 
                                B.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                B.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                B.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="C":                 
                                C.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                C.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                C.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="D":                 
                                D.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                D.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                D.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="E":                 
                                E.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                E.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                E.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="F":                 
                                F.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                F.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                F.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="L":                 
                                L.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                L.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                L.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="H":                 
                                H.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                H.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                H.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="AL":                 
                                AL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                AL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="AH":                 
                                AH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                AH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="BL":                 
                                BL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                BL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="BH":                 
                                BH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                BH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="CL":                 
                                CL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                CL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="CH":                 
                                CH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                CH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="DL":                 
                                DL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                DL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="DH":                 
                                DH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                DH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="EL":                 
                                EL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                EL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="EH":                 
                                EH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                EH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="FL":                 
                                FL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                FL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="FH":                 
                                FH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                FH.append(new_round(divisor/dividend,3))
                        
                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],nrows): 
                            if file_data.row_values(i)[nameindex]=="A":
                                A.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                A.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                A.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="B":                 
                                B.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                B.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                B.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="C":                 
                                C.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                C.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                C.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="D":                 
                                D.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                D.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                D.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="E":                 
                                E.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                E.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                E.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="F":                 
                                F.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                F.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                F.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="L":                 
                                L.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                L.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                L.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="H":                 
                                H.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                H.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                H.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="AL":                 
                                AL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                AL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="AH":                 
                                AH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                AH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                AH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="BL":                 
                                BL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                BL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="BH":                 
                                BH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                BH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                BH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="CL":                 
                                CL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                CL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="CH":                 
                                CH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                CH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                CH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="DL":                 
                                DL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                DL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="DH":                 
                                DH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                DH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                DH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="EL":                 
                                EL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                EL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="EH":                 
                                EH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                EH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                EH.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="FL":                 
                                FL.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FL.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                FL.append(new_round(divisor/dividend,3))
                            elif file_data.row_values(i)[nameindex]=="FH":                 
                                FH.append(effectnum(file_data.row_values(i)[Areaindex],digits))
                                FH.append(effectnum(file_data.row_values(i)[ISAreaindex],digits))

                                divisor=float(effectnum(file_data.row_values(i)[Areaindex],digits))
                                dividend=float(effectnum(file_data.row_values(i)[ISAreaindex],digits))
                                
                                FH.append(new_round(divisor/dividend,3))

                    group_Me["A"]=A
                    group_Me["B"]=B
                    group_Me["C"]=C
                    group_Me["D"]=D
                    group_Me["E"]=E
                    group_Me["F"]=F
                    group_Me["L"]=L
                    group_Me["H"]=H
                    group_Me["AL"]=AL
                    group_Me["AH"]=AH
                    group_Me["BL"]=BL
                    group_Me["BH"]=BH
                    group_Me["CL"]=CL
                    group_Me["CH"]=CH
                    group_Me["DL"]=DL
                    group_Me["DH"]=DH
                    group_Me["EL"]=EL
                    group_Me["EH"]=EH
                    group_Me["FL"]=FL
                    group_Me["FH"]=FH

                    Matrix_effect_dict[norm[j]]=group_Me
        
        ########文件读取完毕#######
                
    print("Matrix_effect_dict:%s" % (Matrix_effect_dict))
    #  第三步:文件读取完毕后的操作

    '''
    注释2:最终需要生成一个字典Matrix_effect_dict,数据格式如下：
    print(Matrix_effect_dict):
    {"化合物1":{"A":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3,Area/IS Area的平均值,混合样本Area/IS Area的理论值,偏差],
    "B":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3,Area/IS Area的平均值,混合样本Area/IS Area的理论值,偏差],...,
    "AL":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3,Area/IS Area的平均值,混合样本Area/IS Area的理论值,偏差],...}
     "化合物2":...
    }
    '''

    # 添加Area/IS Area的平均值
    for key,value in Matrix_effect_dict.items():
        for r,c in value.items():
            mean_Area_ISArea=new_round((float(c[2])+float(c[5])+float(c[8]))/3,3) #Area/IS Area的平均值
            value[r].append(mean_Area_ISArea)

    # 添加混合样本Area/IS Area的理论值
    for key,value in Matrix_effect_dict.items():
        for r,c in value.items():
            if "PS-" in r or "STD-" in r:  # 判断是否是混合样本
                value[r].append("/") #不是混合样本添加反斜杠
            else: # 混合样本MIX-AL,MIX-AH...,提取PS-A和STD-L,找到对应字典中的值
                mean1=float(value["PS-"+r[4]][9]) #提取第一个A,找到A字典中的Area/IS Area的平均值,即为第9个值
                mean2=float(value["STD-"+r[5]][9]) #提取第二个L,找到L字典中的Area/IS Area的平均值,即为第9个值
                value[r].append(new_round((mean1+mean2)/2,3)) #是混合样本添加mean1和mean2的均值

    # 添加偏差(%)
    Matrix_effect_judgenum=0
    for key,value in Matrix_effect_dict.items():
        for r,c in value.items():
            if "PS-" in r or "STD-" in r:  # 判断是否是混合样本
                for i in [2,5,8]:
                    value[r].append("/") #不是混合样本添加反斜杠
            else: # 混合样本AL,AH...,提取第一个A和第二个L,找到对应字典中的值
                for i in [2,5,8]: # 2,5,8为三个Area/IS Area值在字典中的索引
                    num1=float(value[r][i]) #提取Area/IS Area
                    num2=float(value[r][10]) #提取混合样本Area/IS Area的理论值,索引为10
                    bias=new_round(abs(num1-num2)/num2*100,2)
                    if float(bias)>maxBias:
                        Matrix_effect_judgenum+=1
                        bias=bias
                    else:
                        bias=bias
                    value[r].append(bias) #是混合样本添加mean1和mean2的均值

    #  第四步:数据存入数据库

    # 如果Matrix_effect_judgenum的值等于0才将数据存入数据库中
    if Matrix_effect_judgenum==0:
        insert_list =[]
        for key,value in Matrix_effect_dict.items():
            for r,c in value.items():
                insert_list.append(Matrixeffect(reportinfo=reportinfo,norm=key,samplename=r,Area_1=c[0],IS_Area_1=c[1],
                Area_IS_Area_1=c[2],Area_2=c[3],IS_Area_2=c[4],Area_IS_Area_2=c[5],Area_3=c[6],IS_Area_3=c[7],Area_IS_Area_3=c[8],
                singlemean=c[9],complexmean=c[10],bias1=c[11],bias2=c[12],bias3=c[13]))    
        Matrixeffect.objects.bulk_create(insert_list)
    else:
        insert_list =[]
        for key,value in Matrix_effect_dict.items():
            for r,c in value.items():
                insert_list.append(Matrixeffect(reportinfo=reportinfo,norm=key,samplename=r,Area_1=c[0],IS_Area_1=c[1],
                Area_IS_Area_1=c[2],Area_2=c[3],IS_Area_2=c[4],Area_IS_Area_2=c[5],Area_3=c[6],IS_Area_3=c[7],Area_IS_Area_3=c[8],
                singlemean=c[9],complexmean=c[10],bias1=c[11],bias2=c[12],bias3=c[13]))    
        Matrixeffect.objects.bulk_create(insert_list)

    return {"Matrix_effect_dict":Matrix_effect_dict,"maxBias":maxBias}


def related_Matrix_effect(id):
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    #特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = Matrixeffectspecial.objects.get(special=special_1)           
        if Matrixeffectspecialtexts.objects.filter(matrixeffectspecial=special_2).count()>0:
            text_special = Matrixeffectspecialtexts.objects.filter(matrixeffectspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = Matrixeffectgeneral.objects.get(general=general_1)
    text_general = Matrixeffectgeneraltexts.objects.filter(matrixeffectgeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits

    # 第二步：报告数据提取

    '''
    注释:需要生成一个字典Matrixeffect_dict,数据格式如下：
    print(Matrixeffect_dict):
    {"化合物1":{"A":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3,Area/IS Area的平均值,混合样本Area/IS Area的理论值,偏差],
    "B":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3,Area/IS Area的平均值,混合样本Area/IS Area的理论值,偏差],...,
    "AL":[Area1,IS Area1,Area/IS Area1,Area2,IS Area2,Area/IS Area2,Area3,IS Area3,Area/IS Area3,Area/IS Area的平均值,混合样本Area/IS Area的理论值,偏差],...}
     "化合物2":...
    }
    '''

    # 定义需要生成的字典
    Matrixeffect_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        Matrixeffect_data = Matrixeffect.objects.filter(reportinfo_id=id)

        Matrixeffect_norm = []  # 待测物质列表
        for i in Matrixeffect_data:
            if i.norm not in Matrixeffect_norm:
                Matrixeffect_norm.append(i.norm)

        for i in Matrixeffect_norm:
            norm_dict={} #各待测物质的字典
            middle_table = Matrixeffect.objects.filter(reportinfo_id=id,norm=i) #各待测物质的数据表        
            for item in middle_table:
                #没有为每个化合物单独设置有效位数，则调用通用性设置
                if Digitsdict=={} or list(Digitsdict.values())[0]==None:     
                    rowlist = [] 
                    rowlist.append(item.Area_1)
                    rowlist.append(item.IS_Area_1)
                    rowlist.append(item.Area_IS_Area_1)
                    rowlist.append(item.Area_2)
                    rowlist.append(item.IS_Area_2)
                    rowlist.append(item.Area_IS_Area_2)
                    rowlist.append(item.Area_3)
                    rowlist.append(item.IS_Area_3)
                    rowlist.append(item.Area_IS_Area_3)
                    rowlist.append(item.singlemean)
                    rowlist.append(item.complexmean)
                    rowlist.append(item.bias1)
                    rowlist.append(item.bias2)
                    rowlist.append(item.bias3)
                    norm_dict[item.samplename]=rowlist
                #为每个化合物单独设置了有效位数，则调用每个化合物的设置
                else:
                    rowlist=[] #各待测物质各samplename列表
                    rowlist.append(effectnum(item.Area_1,Digitsdict[i]))
                    rowlist.append(effectnum(item.IS_Area_1,Digitsdict[i]))
                    rowlist.append(item.Area_IS_Area_1)
                    rowlist.append(effectnum(item.Area_2,Digitsdict[i]))
                    rowlist.append(effectnum(item.IS_Area_2,Digitsdict[i]))
                    rowlist.append(item.Area_IS_Area_2)
                    rowlist.append(effectnum(item.Area_3,Digitsdict[i]))
                    rowlist.append(effectnum(item.IS_Area_3,Digitsdict[i]))
                    rowlist.append(item.Area_IS_Area_3)
                    rowlist.append(item.singlemean)
                    rowlist.append(item.complexmean)
                    rowlist.append(item.bias1)
                    rowlist.append(item.bias2)
                    rowlist.append(item.bias3)
                    norm_dict[item.samplename]=rowlist

            Matrixeffect_dict[i]=norm_dict
        
        Matrixeffect_endreport_conclusion="混合样本中" + "、" .join(list(Matrixeffect_dict.keys()))+"的相应值(a)与高低浓度样本和病人样本中"+"、" .join(list(Matrixeffect_dict.keys()))+"响应值的均值(b)差异均小于20%，说明无相对基质效应。"

        if len(textlist_special)!=0:
            return {"Matrixeffect_dict":Matrixeffect_dict,"textlist":textlist_special,"serial":len(textlist_special)+1,"Matrixeffect_endreport_conclusion":Matrixeffect_endreport_conclusion}

        else:
            return {"Matrixeffect_dict":Matrixeffect_dict,"textlist":textlist_general,"serial":len(textlist_general)+1,"Matrixeffect_endreport_conclusion":Matrixeffect_endreport_conclusion}

    except:
        pass