import xlrd
from report.models import *
from docx import Document
from report.effectnum import *
import numpy as np
import re

def fileread(files, Detectionplatform, reportinfo, project, platform,manufacturers, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB):

    # 第一步:后台数据抓取（最小样本数，最大允许CV）

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)


    #  第二步:开始文件读取
    '''
      数据读取完毕后,需要生成一个字典,分别对应不同化合物的原始数据。数据格式如下：
      Referenceinterval_dict = {norm1:[原始数据]，norm2:[原始数据]，norm3:[原始数据]}
    '''

    # 创新需要生成的结果容器
    Referenceinterval_dict = {}

    # 各仪器平台及各仪器厂家数据读取
    for file in files:
        if platform=="液质":
            if manufacturers =="Agilent":
                # 1 读取csv文件（Agilent）
                file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)): 
                    if len(lines[i])!=0:
                        lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                        # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                    else:
                        lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i] #最后一行如为空行，则删除该元素

                # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                norm=[] #化合物列表
                for j in range(len(lines[0])):
                    if "-Q Results" in lines[0][j]:
                        if lines[0][j].split("-Q")[0][0]!='"':  # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除  
                                norm.append(lines[0][j].split("-Q")[0])
                        else:
                                norm.append(lines[0][j].split("-Q")[0][1:])

                # 从第二行确定:实验号（Sample Name）,浓度（Final Conc.）的索引
                nameindex=0  #实验号索引
                concindex=[] #浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                for j in range(len(lines[1])):  #从第二行开始       
                    if lines[1][j] == "Sample Name" :
                        nameindex=j
                    elif lines[1][j]  == "Final Conc." :
                        concindex.append(j)

                # 添加原始数据
                for j in range(len(norm)):
                    normlist = []  # 每个化合物的结果列表
                    for i in range(len(lines)):  # 循环原始数据中的每一行
                        if "Reference_interval" in lines[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                            normlist.append(effectnum(lines[i][concindex[j]], digits)) 

                    Referenceinterval_dict[norm[j]] = normlist

            elif manufacturers == "岛津":
                # 读取txt文件
                content = []
                for line in file:
                    content.append(line.decode("UTF-8").replace("\r\n", "").split("\t")) # windows下
                    # content.append(line.decode("GB2312").replace("\r\n", "").split("\t")) # linux下

                nameindex = 0
                concindex = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行

                for i in range(len(content)):
                    if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称
                        norm.append(content[i][1])
                        norm_row.append(i)

                for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                    if content[2][i] == "数据文件名":
                        nameindex = i
                    elif content[2][i] == "浓度":
                        concindex = i

                # 添加原始数据
                for j in range(len(norm)):
                    normlist = []  # 每个化合物的结果列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            if "Reference_interval" in content[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                                normlist.append(effectnum(content[i][concindex], digits)) 

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], len(content)):
                            if "Reference_interval" in content[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                                normlist.append(effectnum(content[i][concindex], digits)) 

                    Referenceinterval_dict[norm[j]] = normlist

            elif manufacturers == "Waters":
                # 若是最新的 2.0.1 版本的xlrd包，只支持 .xls 文件，读取.xlsx文件会报错。若要正常读取，需安装旧版本的xlrd：pip3 install xlrd==1.2.0
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行
                for i in range(nrows):
                    if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):  # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                        norm.append(file_data.row_values(i)[0].split(":")[1].strip()) # strip()的作用是去除前后空格
                        norm_row.append(i)    

                print(norm_row)          

                nameindex = 0
                concindex = 0
                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "Name":
                        nameindex = i
                    elif "实际浓度" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                # 添加原始数据
                for j in range(len(norm)):
                    normlist = []  # 每个化合物的结果列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            if "Reference_interval" in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                                normlist.append(effectnum(file_data.row_values(i)[concindex], digits)) 

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], nrows):
                            if "Reference_interval" in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                                normlist.append(effectnum(file_data.row_values(i)[concindex], digits)) 

                    Referenceinterval_dict[norm[j]] = normlist

            # Thermo厂家需先在后台管理系统中设置本项目的化合物名称，以便查找上传文件中相应化合物的表格
            elif manufacturers == "Thermo":
                Thermo = Special.objects.get(project=project)
                pt_special = PTspecial.objects.get(special=Thermo)
                pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

                # 后台管理系统中设置的本项目化合物名称
                PTnorm = []  
                for i in pt_accept:
                    PTnorm.append(i.norm)

                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm = []  # Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                sheetindex = []  # 需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    if data.sheet_names()[index] in PTnorm:
                        norm.append(data.sheet_names()[index])
                        sheetindex.append(index)

                print(norm)
                # 循环读取每个sheet工作表,即为每个化合物的表
                for index in range(len(sheetindex)):
                    file_data = data.sheets()[sheetindex[index]]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    # 第一行确定samplename和浓度所在列
                    nameindex = 0
                    concindex = 0
                    for i in range(len(file_data.row_values(0))):
                        if file_data.row_values(0)[i] == "Compound":
                            nameindex = i
                        elif file_data.row_values(0)[i] == "Calculated Amt":
                            concindex = i

                    normlist = [] # 每个化合物的结果列表
                    for i in range(nrows):
                        if "Reference_interval" in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                            normlist.append(effectnum(file_data.row_values(i)[concindex], digits)) 

                    Referenceinterval_dict[norm[index]] = normlist

            # AB厂家需先在后台设置化合物和相应离子对数值，以便判断需要读取的表格（定量表格）和不需要读取的表格(定性表格)
            elif manufacturers == "AB":

                # 测试文件中的三个化合物（1  Clozapine定性 (327.2 / 192.2)，Clozapine定量 (327.2 / 270.2);
                #                     2  Sertraline定性 (306.1 / 275.1)，Sertraline定量 (306.1 / 159.1)）
                #                     3  Aripiprazole定性 (448.1 / 176.2),Aripiprazole定量 (448.1 / 285.2)

                # 定义化合物列表，列表统一命名为norm
                norm = normAB

                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = [] 

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                tables = file_data.tables  # 获取文件中的表格集

                # 循环定量表格的索引
                for k in range(len(tableindex)):
                    # 获取文件中的定量表格
                    tablequantify = tables[tableindex[k]]

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表
                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % 12 != 0 or i == 0:  # docx文件固定为12列
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    nameindex = 0
                    concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列
                    
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name":
                            nameindex = i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex = i

                    normlist = []  # 每个化合物的结果列表
                    for i in range(len(rowdatagatherlist)):
                        if "Reference_interval" in rowdatagatherlist[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                            normlist.append(effectnum(rowdatagatherlist[i][concindex], digits)) 

                    Referenceinterval_dict[norm[k]] = normlist
 
        elif platform == "液相":
            if manufacturers == "Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]  # 默认只读取第一个工作簿
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行
                for j in range(nrows):
                    # 如果某一行的第一个元素为'化合物'，则添加第三个元素进入化合物列表
                    if file_data.row_values(j)[0] == "化合物:":
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex = 0
                concindex = 0

                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "样品名称":
                        nameindex = i
                    elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                # 添加原始数据
                for j in range(len(norm)):
                    normlist = []  # 每个化合物的结果列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            if "Reference_interval" in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                                normlist.append(effectnum(file_data.row_values(i)[concindex], digits)) 

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], nrows):
                            if "Reference_interval" in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                                normlist.append(effectnum(file_data.row_values(i)[concindex], digits)) 

                    Referenceinterval_dict[norm[j]] = normlist

        elif platform == "ICP-MS":
            # ICP-MS平台Agilent厂家需先在后台管理系统中设置本项目的化合物名称，以便查找上传文件中相应化合物的表格
            if manufacturers == "Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                # 从第一行确定化合物名称
                norm = []
                for j in range(ncols):
                    for i in PTnorm:
                        if i in file_data.row_values(0)[j]:
                            norm.append(i)

                # 从第二行确定实验号（Sample Name）的索引和化合物浓度索引
                nameindex = 0  # 实验号索引
                concindex = []  # 浓度索引
                for j in range(ncols):
                    if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j
                    elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                        concindex.append(j)

                # 添加原始数据
                for j in range(len(norm)):
                    normlist = []  # 每个化合物的结果列表
                    for i in range(2, nrows):  # 循环原始数据中的每一行
                        if "Reference_interval" in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据          
                            normlist.append(effectnum(file_data.row_values(i)[concindex[j]], digits)) 

                    Referenceinterval_dict[norm[j]] = normlist

    print(Referenceinterval_dict)

    #  第二步:计算参考区间
    for key,value in Referenceinterval_dict.items():
        # 列表中的字符串转换为浮点数
        arry = np.array(list(map(float,value)))

        # 计算参考区间上下限
        lower_limit = new_round(np.percentile(arry, 2.5))
        upper_limit = new_round(np.percentile(arry, 97.5))

        # 添加参考区间上下限
        Referenceinterval_dict[key].append('~'.join([lower_limit,upper_limit]))
    
    print(Referenceinterval_dict)

    #  第三步:数据存入数据库

    insert_list = []
    for key, value in Referenceinterval_dict.items():
        forloopindex=0 # for循环索引，赋值给实验号
        for i in value:
            forloopindex+=1
            insert_list.append(Reference_Interval(reportinfo=reportinfo, norm=key,Experimentnum=forloopindex,Result=i))

    Reference_Interval.objects.bulk_create(insert_list)
    return {'Referenceinterval_dict': Referenceinterval_dict, "Unit": Unit}

def data_scrap(id):
    # 第一步：后台描述性内容数据提取
    # 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    # 特殊数据抓取

    #特殊参数设置描述性内容
    #   textlist_special = []
    #   try:
    #         special_1 = Special.objects.get(project=project) 
    #         prepared_Sample_Stability_special = Prepared_Sample_Stability_special.objects.get(special=special_1)           
    #         if Prepared_Sample_Stability_special_texts.objects.filter(prepared_Sample_Stability_special=prepared_Sample_Stability_special).count()>0:
    #               text_special = Prepared_Sample_Stability_special_texts.objects.filter(prepared_Sample_Stability_special=prepared_Sample_Stability_special)  
    #               for i in text_special:
    #                     textlist_special.append(i.text)
    #   except:
    #         pass
    
    # 通用数据抓取
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    stability_general = Stabilitygeneral.objects.get(general=general_1)
    text_general = Stabilitygeneraltexts.objects.filter(stabilitygeneral=stability_general)   

    # 描述性内容
    textlist_general = [] 
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project) 
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist=[] # 每个化合物有效位数列表
    Digitsdict={} # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist==[] or Digitslist[0]=="": #如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm]=i.digits

    # 第二步：数据库抓取数据

    '''
    1  需要生成几个字典，需要生成一个字典,分别对应不同化合物的原始数据。数据格式如下：
      Referenceinterval_dict = {norm1:[原始数据,参考区间]，norm2:[原始数据,参考区间]，norm3:[原始数据,参考区间]}
    2  总结性结论：实验结果如表xx-xx所示，表明处理后的样品在室温和冷藏条件下至少可以稳定存放72h；
    '''

    # 定义需要生成的字典
    Referenceinterval_dict = {}

    try:  
        Referenceinterval_data = Reference_Interval.objects.filter(reportinfo_id=id)

        norm = [] # 化合物列表
        for i in Referenceinterval_data:
            if i.norm not in norm:
                norm.append(i.norm)
        
        for j in norm:
            normlist = []  # 每个化合物的数据
            normtable = Reference_Interval.objects.filter(reportinfo_id=id,norm = j) # 每个化合物的数据表格                  
            for k in normtable:
                normlist.append(k.Result) # 添加基础数据
                
            Referenceinterval_dict[j]=normlist
        
        # 第三步:数据格式转换
        # 转换后需要生成的列表：[ [化合物1结果1，化合物2结果1，化合物3结果1],[化合物1结果2，化合物2结果2，化合物3结果2]]

        # 定义转换后需要生成的列表
        Referenceinterval_list = []
        
        for i in range(len(Referenceinterval_dict[norm[0]])):
            group = []
            for key,value in Referenceinterval_dict.items():     
                group.append(Referenceinterval_dict[key][i])
            Referenceinterval_list.append(group)
    
    except:
        pass
    
    return {"Referenceinterval_dict": Referenceinterval_dict,"Referenceinterval_list": Referenceinterval_list,"textlist": textlist_general,"serial": len(textlist_general)+1}