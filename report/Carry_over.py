from random import SystemRandom
import numpy as np
import xlrd
from docx import Document
import math
from report.models import *
from report.effectnum import *
from datetime import datetime


def Carryover_9sample_fileread(files,Detectionplatform,reportinfo,project,platform,manufacturers,Unit,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB,Number_of_compounds):

    # 第一步:后台数据抓取（回收率上下限）
    id1 = Special.objects.get(project=project).id  
    id2 = Carryoverspecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if Carryoverspecialmethod.objects.filter(carryoverspecial=id2): 
        maxaccept=Carryoverspecialmethod.objects.get(carryoverspecial=id2).accept #可接受标准
   
    else:
        general_1 = General.objects.get(name="通用性项目")
        carryover_general = Carryovergeneral.objects.get(general=general_1)
        maxaccept=Carryovergeneralmethod.objects.get(carryovergeneral=carryover_general).acceptable

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示
    if manufacturers=="AB":
        if len(normAB)!= Number_of_compounds:
            error_message="未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"  
            return {"error_message":error_message}

    #  第二步:开始文件读取

    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典Carryover_dict,数据格式如下：
    {"s1系统":{"化合物1":[C1,C2,C3,C1,C2,C3],"化合物2":[C1,C2,C3,C1,C2,C3]},
    "s2系统":{"化合物1":[C1,C2,C3,C1,C2,C3],"化合物2":[C1,C2,C3,C1,C2,C3]}}
    '''

    # 创建第二步需要生成的结果容器
    Carryover_dict={}
    Systermlist=["S1系统","S2系统","S3系统","S4系统"] #目前最多支持4个系统

    for index,file in enumerate(files):
        Carryover_dict[Systermlist[index]]={}
        if platform=="液质":
            if manufacturers =="Agilent":
                # 1 读取csv文件（Agilent）
                file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)): 
                    if len(lines[i])!=0:
                        lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                    else:
                        lines[i]=lines[i].split(',') 
                        del lines[i] #最后一行如为空行，则删除该元素

                # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                norm=[] #化合物列表
                for j in range(len(lines[0])):  #从第一行开始
                    if "-Q Results" in lines[0][j]:
                        norm.append(lines[0][j].split("-Q")[0])

                # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                nameindex=0  #实验号索引
                conindex=[] #浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                for j in range(len(lines[1])):  #从第二行开始       
                    if lines[1][j] == "Sample Name" :
                        nameindex=j
                    elif lines[1][j]  == "Final Conc." :
                        conindex.append(j)

                for j in range(len(norm)):
                    normlist=[]
                    
                    for i in range(len(lines)): # 循环原始数据中的每一行
                        if "Carryover-C1" in lines[i][nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                            normlist.append(effectnum(lines[i][conindex[j]],digits))
                        elif "Carryover-C2" in lines[i][nameindex]:
                            normlist.append(effectnum(lines[i][conindex[j]],digits))
                        elif "Carryover-C3" in lines[i][nameindex]:
                            normlist.append(effectnum(lines[i][conindex[j]],digits))
                    
                    Carryover_dict[Systermlist[index]][norm[j]]=normlist

            elif manufacturers =="Waters":
                # 内标标识
                ISlist=["D3","D4","D5","D6","D7","D8"]

                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                norm = []  # 化合物列表
                Compound_row =[] # 含有“Compound”关键词的所在行(包含内标)
                norm_row = []  # 实际化合物所在行(不包含内标)
                for i in range(nrows):
                    if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                        Compound_row.append(i)  

                    # 判断是否含有内标标识
                    if all(j not in str(file_data.row_values(i)[0]) for j in ISlist):
                        if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):  # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                            norm.append(file_data.row_values(i)[0].split(":")[1].strip()) # strip()的作用是去除前后空格
                            norm_row.append(i) 

                # 第一种情况，不含有内标
                if len(Compound_row) == len(norm_row):
                    pass

                # 第二种情况，含有内标
                else:
                    nrows = Compound_row[len(norm_row)]                

                nameindex = 0
                concindex = 0
                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "ID":
                        nameindex = i
                    elif "nmol/L" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                # 未准确设置表头列名,直接返回并提示!
                if nameindex==0 or concindex==0:
                    error_message="未准确设置表头列名!"
                    return {"error_message":error_message}

                for j in range(len(norm)):
                    normlist=[]
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]):
                            if "Carryover-C1" in file_data.row_values(i)[nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                normlist=[].append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C2" in file_data.row_values(i)[nameindex]:
                                normlist=[].append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C3" in file_data.row_values(i)[nameindex]:
                                normlist=[].append(effectnum(file_data.row_values(i)[conindex],digits))
                    
                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],nrows): 
                            if "Carryover-C1" in file_data.row_values(i)[nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                normlist=[].append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C2" in file_data.row_values(i)[nameindex]:
                                normlist=[].append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C3" in file_data.row_values(i)[nameindex]:
                                normlist=[].append(effectnum(file_data.row_values(i)[conindex],digits))
                    
                    Carryover_dict[Systermlist[index]][norm[j]]=normlist
            
            elif manufacturers =="Thermo":
                Thermo = Special.objects.get(project=project) 
                pt_special = PTspecial.objects.get(special=Thermo)
                pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
                PTnorm=[] # 待测物质列表
                for i in pt_accept:
                    PTnorm.append(i.norm)
                    
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm=[] #Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                sheetindex=[] #需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    if data.sheet_names()[index] in PTnorm:
                        norm.append(data.sheet_names()[index])
                        sheetindex.append(index)

                # 循环读取每个sheet工作表,即为每个化合物的表
                for index in range(len(sheetindex)):
                    file_data = data.sheets()[sheetindex[index]]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    #第一行确定samplename和浓度所在列
                    nameindex=0
                    conindex=0
                    for i in range(len(file_data.row_values(0))):  
                        if file_data.row_values(0)[i]=="Compound":
                            nameindex=i
                        elif file_data.row_values(0)[i]=="Calculated Amt":
                            conindex=i

                    C1=[] 
                    C2=[] 
                    C3=[]
                    group_Carryover=[]
                    for i in range(nrows):
                        if "Carryover-C1" in file_data.row_values(i)[nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                            C1.append(effectnum(file_data.row_values(i)[conindex],digits))
                        elif "Carryover-C2" in file_data.row_values(i)[nameindex]:
                            C2.append(effectnum(file_data.row_values(i)[conindex],digits))
                        elif "Carryover-C3" in file_data.row_values(i)[nameindex]:
                            C3.append(effectnum(file_data.row_values(i)[conindex],digits))
        
                    group_Carryover.append(C1)
                    group_Carryover.append(C2)
                    group_Carryover.append(C3)
                    Carryover_list.append(group_Carryover)

            elif manufacturers =="岛津":
                # 3 读取txt文件
                content= []
                for line in file:
                    content.append(line.decode("GB2312").replace("\r\n", "").split("\t"))

                nameindex=0
                conindex=0 #浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                norm=[] #化合物列表
                norm_row=[] #化合物所在行

                for i in range(len(content[2])):  #第二行确定samplename和浓度所在列
                    if content[2][i]=="数据文件名":
                        nameindex=i 
                    elif content[2][i]=="浓度":
                        conindex=i 

                for i in range(len(content)): 
                    if content[i][0]=="Name": #如果某一行第一列为"Name"，则该行第二列为化合物名称
                        norm.append(content[i][1])
                        norm_row.append(i)

                for j in range(len(norm)):
                    C1=[] 
                    C2=[] 
                    C3=[]
                    group_Carryover=[]
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]):
                            if "Carryover-C1" in content[i][nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                C1.append(effectnum(content[i][conindex],digits))
                            elif "Carryover-C2" in content[i][nameindex]:
                                C2.append(effectnum(content[i][conindex],digits))
                            elif "Carryover-C3" in content[i][nameindex]:
                                C3.append(effectnum(content[i][conindex],digits))
                    
                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],len(content)): 
                            if "Carryover-C1" in content[i][nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                C1.append(effectnum(content[i][conindex],digits))
                            elif "Carryover-C2" in content[i][nameindex]:
                                C2.append(effectnum(content[i][conindex],digits))
                            elif "Carryover-C3" in content[i][nameindex]:
                                C3.append(effectnum(content[i][conindex],digits))
                    
                    group_Carryover.append(C1)
                    group_Carryover.append(C2)
                    group_Carryover.append(C3)
                    Carryover_list.append(group_Carryover)

            elif manufacturers =="AB":

                # 定义化合物列表，列表统一命名为norm
                norm = normAB

                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = [] 

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                tables = file_data.tables  # 获取文件中的表格集

                for k in range(len(tableindex)):
                    # 获取文件中的定量表格
                    tablequantify = tables[tableindex[k]] 

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表

                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex=0 #实验号索引
                    concindex=0 #浓度索引

                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name" :
                            nameindex=i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex=i
                    
                    normlist=[]
                    for i in range(len(rowdatagatherlist)): 
                        if "Carryover-C1" in rowdatagatherlist[i][nameindex]:
                            normlist.append(effectnum(rowdatagatherlist[i][concindex],digits))
                        elif "Carryover-C2" in rowdatagatherlist[i][nameindex]:
                            normlist.append(effectnum(rowdatagatherlist[i][concindex],digits))
                        elif "Carryover-C3" in rowdatagatherlist[i][nameindex]:
                            normlist.append(effectnum(rowdatagatherlist[i][concindex],digits))

                    Carryover_dict[Systermlist[index]][norm[k]]=normlist

        elif platform=="液相":
            if manufacturers =="Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                norm=[] #化合物列表
                norm_row=[] #化合物所在行
                for j in range(nrows):
                    if file_data.row_values(j)[0] == "化合物:" :  #如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex=0
                conindex=0

                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+1是该化合物表格位于该化合物所在行的下一行
                for i in range(len(file_data.row_values(norm_row[0]+1))):
                    if "样品名称" in file_data.row_values(norm_row[0]+1)[i]:
                        nameindex = i
                    elif "含量" in file_data.row_values(norm_row[0]+1)[i]:
                        conindex = i

                for j in range(len(norm)):
                    normlist=[] 
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]):
                            if "Carryover-C1" in file_data.row_values(i)[nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                normlist.append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C2" in file_data.row_values(i)[nameindex]:
                                normlist.append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C3" in file_data.row_values(i)[nameindex]:
                                normlist.append(effectnum(file_data.row_values(i)[conindex],digits))
                    
                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],nrows): 
                            if "Carryover-C1" in file_data.row_values(i)[nameindex]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                normlist.append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C2" in file_data.row_values(i)[nameindex]:
                                normlist.append(effectnum(file_data.row_values(i)[conindex],digits))
                            elif "Carryover-C3" in file_data.row_values(i)[nameindex]:
                                normlist.append(effectnum(file_data.row_values(i)[conindex],digits))
                    
                    Carryover_dict[Systermlist[index]][norm[j]]=normlist

    ########文件读取完毕#######
                

    #  第三步:文件读取完毕后的操作(添加:C1mean,C3mean,(C3mean-C1mean)/C1mean)
    '''
    注释2:第三步,需要生成一个字典Carryover_enddict,数据格式如下：
    {"化合物1":{"s1系统":[C1,C2,C3,C1,C2,C3],"s2系统":[C1,C2,C3,C1,C2,C3]},
    "化合物2":{"s1系统":[C1,C2,C3,C1,C2,C3],"s2系统":[C1,C2,C3,C1,C2,C3]}
    }
    '''

    # 创建第三步需要生成的结果容器
    Carryover_enddict={}

    # Carryover_dict格式转换
    for i in norm:
        Carryover_enddict[i]={}
        for key,value in Carryover_dict.items():
            for r,c in value.items():
                if r==i:
                    Carryover_enddict[i][key]=c

    Carryover_judgenum=0

    for key,value in Carryover_enddict.items():
        for r,c in value.items():
            c2 = list(map(float,c)) # 列表中的字符串转换为浮点数
            C1mean=new_round((c2[0]+c2[3]+c2[6])/3,2)
            C3mean=new_round((c2[2]+c2[5]+c2[8])/3,2)
            bias=new_round(abs(float(C3mean)-float(C1mean))/float(C1mean)*100,2)
            value[r].append(C1mean)
            value[r].append(C3mean)
            value[r].append(bias)

            if float(bias)>maxaccept:
                Carryover_judgenum+=1

    #  第四步:数据存入数据库

    # 如果Carryover_judgenum的值等于0才将数据存入数据库中
    if Carryover_judgenum==0:
        if len(files)==1:
            insert_list =[]
            for key,value in Carryover_enddict.items():
                for r,c in value.items():
                    insert_list.append(Carryover(reportinfo=reportinfo,norm=key,systermnum=r,C1_1=c[0],C2_1=c[1],C3_1=c[2],C1_2=c[3],
                    C2_2=c[4],C3_2=c[5],C1_3=c[6],C2_3=c[7],C3_3=c[8],C1mean=c[9],C3mean=c[10],bias=c[11]))
            
            Carryover.objects.bulk_create(insert_list)

        else:
            insert_list =[]
            for key,value in Carryover_enddict.items():
                for r,c in value.items():
                    insert_list.append(Carryover(reportinfo=reportinfo,norm=key,systermnum=r,C1_1=c[0],C2_1=c[1],C3_1=c[2],C1_2=c[3],
                    C2_2=c[4],C3_2=c[5],C1_3=c[6],C2_3=c[7],C3_3=c[8],C1mean=c[9],C3mean=c[10],bias=c[11]))
            
            Carryover.objects.bulk_create(insert_list)

    print(Carryover_enddict)
    return {"Carryover_dict":Carryover_enddict,"Systermnum":len(files),"Unit":Unit,"maxaccept":maxaccept}

def Carryover_21sample_fileread(files,Detectionplatform,reportinfo,project,platform,manufacturers,Unit,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB):

    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典Carryover_dict,数据格式如下：
    print(Carryover_dict):
    {化合物1:[L01-L21(21个原始数据),X1,X2,SD1],化合物2:[L01-L21(21个原始数据),X1,X2,SD1]}
    '''
    
    Carryover_dict={}
    for file in files:
        if platform=="液质":
            if manufacturers =="Agilent":
                pass
                        
            elif manufacturers =="岛津":
                pass

            elif manufacturers =="AB":
                pass

        elif platform=="液相":
            pass

        elif platform=="ICP-MS":
            # 从后台管理系统PT数据表中找到化合物
            zqd = Special.objects.get(project=project) 
            pt_special = PTspecial.objects.get(special=zqd)
            pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
            PTnorm=[] # 化合物列表
            for i in pt_accept:
                PTnorm.append(i.norm)

            for i in range(len(PTnorm)):
                Carryover_dict[PTnorm[i]] =[]

            data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
            file_data = data.sheets()[0]
            nrows=file_data.nrows
            ncols=file_data.ncols
            conindex=[] #浓度索引

            # 从第一行确定化合物浓度对应的索引
            for j in range(ncols):
                for i in PTnorm:             
                    if i in file_data.row_values(0)[j]:
                        conindex.append(j)   

            # 从第二行确定实验号（Sample Name）的索引
            nameindex=0  #实验号索引
            for j in range(ncols):       
                if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j

            # 匹配原始数据中与携带效应相关(实验号前含有"Carryover-")的行  
            for j in range(len(conindex)):
                for i in range(nrows): # 循环原始数据中的每一行
                    if "Carryover-" in file_data.row_values(i)[nameindex]:                             
                        Carryover_dict[PTnorm[j]].append(effectnum(file_data.row_values(i)[conindex[j]],digits))

            Carryover_judgenum=0
            for key,value in Carryover_dict.items():  
                X1=[] #提前定义X1和X2两个列表,方便计算均值和SD
                X2=[] 
                X1.append(float(value[1]))   
                X1.append(float(value[2]))           
                X1.append(float(value[9]))
                X1.append(float(value[10]))
                X1.append(float(value[11]))

                X2.append(float(value[5]))   
                X2.append(float(value[8]))           
                X2.append(float(value[14]))
                X2.append(float(value[17]))
                X2.append(float(value[20]))

                Carryover_dict[key].append(new_round(np.mean(X1),2)) #添加X1(2、3、10、11、12样品的均值)
                Carryover_dict[key].append(new_round(np.mean(X2),2)) #添加X2(6、9、15、18、21样品的均值)
                Carryover_dict[key].append(new_round(new_round(np.mean(X2),2)-new_round(np.mean(X1),2),2)) #添加X2-X1
                Carryover_dict[key].append(new_round(np.std(X1,ddof=1),2)) #添加SD1(2、3、10、11、12样品的标准差)
                Carryover_dict[key].append(new_round(3*new_round(np.std(X1,ddof=1),2),2)) #添加3SD1(2、3、10、11、12样品的标准差)

                if new_round(np.mean(X2),2)-new_round(np.mean(X1),2)<3*new_round(np.std(X1,ddof=1),2):  #判断是否通过(X2-X1<3SD1)
                    Carryover_dict[key].append("通过!")
                else:
                    Carryover_dict[key].append("不通过!")
                    Carryover_judgenum+=1      

    # 如果Carryover_judgenum的值等于0才将数据存入数据库中
    print(Carryover_dict)

    if Carryover_judgenum==0:
        insert_list =[]
        for key,value in Carryover_dict.items():
            insert_list.append(Carryover2(reportinfo=reportinfo,norm=key,L01=value[0],L02=value[1],L03=value[2],H04=value[3],H05=value[4],
            L06=value[5],H07=value[6],H08=value[7],L09=value[8],L10=value[9],L11=value[10],L12=value[11],H13=value[12],H14=value[13],L15=value[14],
            H16=value[15],H17=value[16],L18=value[17],H19=value[18],H20=value[19],L21=value[20],X1=value[21],X2=value[22],SD1=value[24]))
        
        Carryover2.objects.bulk_create(insert_list)
    
    else:
        insert_list =[]
        for key,value in Carryover_dict.items():
            insert_list.append(Carryover2(reportinfo=reportinfo,norm=key,L01=value[0],L02=value[1],L03=value[2],H04=value[3],H05=value[4],
            L06=value[5],H07=value[6],H08=value[7],L09=value[8],L10=value[9],L11=value[10],L12=value[11],H13=value[12],H14=value[13],L15=value[14],
            H16=value[15],H17=value[16],L18=value[17],H19=value[18],H20=value[19],L21=value[20],X1=value[21],X2=value[22],SD1=value[24]))
        
        Carryover2.objects.bulk_create(insert_list)
    
    return {"Carryover_dict":Carryover_dict,"Unit":Unit}


# 携带效应数据关联进入最终报告
def related_Carryover(id): 
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    #特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = Carryoverspecial.objects.get(special=special_1)           
        if Carryoverspecialtexts.objects.filter(carryoverspecial=special_2).count()>0:
            text_special = Carryoverspecialtexts.objects.filter(carryoverspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = Carryovergeneral.objects.get(general=general_1)
    text_general = Carryovergeneraltexts.objects.filter(carryovergeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)
    
    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project) 
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist=[] # 每个化合物有效位数列表
    Digitsdict={} # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist==[] or Digitslist[0]=="": #如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm]=i.digits
    

    # 第二步：报告数据提取

    '''
    注释:需要生成一个字典Carryover_dict,数据格式如下：
    {"化合物1":{"s1系统":[C1,C2,C3,C1,C2,C3],"s2系统":[C1,C2,C3,C1,C2,C3]},
    "化合物2":{"s1系统":[C1,C2,C3,C1,C2,C3],"s2系统":[C1,C2,C3,C1,C2,C3]}
    }
    '''

    # 1 9个样本的做法

    # 定义需要生成的字典
    Carryover_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        Carryover_data = Carryover.objects.filter(reportinfo_id=id)

        # 化合物列表
        Carryover_norm=[]
        for i in Carryover_data:
            if i.norm not in Carryover_norm:
                Carryover_norm.append(i.norm)

        # 携带效应系统编号列表
        Carryover_systermnum=[]
        for i in Carryover_data:
            if i.systermnum not in Carryover_systermnum:
                Carryover_systermnum.append(i.systermnum)
        

        for i in Carryover_norm:
            middle_dict = {}  # 每个化合物的数据字典
            for j in Carryover_systermnum:
                middle_list=[] # 每个化合物下每个系统编号的数据列表
                middle_table = Carryover.objects.filter(reportinfo_id=id,norm=i,systermnum=j)
                for item in middle_table:
                    # 没有为每个化合物单独设置有效位数，则调用通用性设置
                    if Digitsdict=={} or list(Digitsdict.values())[0]==None: 
                        middle_list.append(item.C1_1)
                        middle_list.append(item.C2_1)
                        middle_list.append(item.C3_1)
                        middle_list.append(item.C1_2)
                        middle_list.append(item.C2_2)
                        middle_list.append(item.C3_2)
                        middle_list.append(item.C1_3)
                        middle_list.append(item.C2_3)
                        middle_list.append(item.C3_3)
                        middle_list.append(item.C1mean)
                        middle_list.append(item.C3mean)
                        middle_list.append(new_round(item.bias,2))
                    #为每个化合物单独设置了有效位数，则调用每个化合物的设置
                    else:
                        middle_list.append(effectnum(item.C1_1,Digitsdict[i]))
                        middle_list.append(effectnum(item.C2_1,Digitsdict[i]))
                        middle_list.append(effectnum(item.C3_1,Digitsdict[i]))
                        middle_list.append(effectnum(item.C1_2,Digitsdict[i]))
                        middle_list.append(effectnum(item.C2_2,Digitsdict[i]))
                        middle_list.append(effectnum(item.C3_2,Digitsdict[i]))
                        middle_list.append(effectnum(item.C1_3,Digitsdict[i]))
                        middle_list.append(effectnum(item.C2_3,Digitsdict[i]))
                        middle_list.append(effectnum(item.C3_3,Digitsdict[i]))
                        middle_list.append(item.C1mean)
                        middle_list.append(item.C3mean)
                        middle_list.append(new_round(item.bias,2))
                
                middle_dict[j]=middle_list
            Carryover_dict[i]=middle_dict

            
        Carryover_conclusion="(C3mean-C1mean)/C1mean均小于20%，说明系统无携带效应，满足检测要求。"

        # 单系统,每行最多排列7个化合物，如超过7个，需拆分表格
        if len(Carryover_systermnum)==1:
            d=Carryover_dict
            if len(Carryover_norm)<=7:
                d_reshape={}
                d_reshape["table1"]=d

            elif len(d)>7 and len(d)<=14: #7-14，拆分为两个table
                d_reshape={}
                if len(d)%2==0: #如果是奇数，第一个表格多一个
                    d_reshape["table1"]=dict(list(d.items())[:len(d)//2+1])
                    d_reshape["table2"]=dict(list(d.items())[len(d)//2+1:])            
                else: #如果是偶数，两个表格平分
                    d_reshape["table1"]=dict(list(d.items())[:len(d)//2])
                    d_reshape["table2"]=dict(list(d.items())[len(d)//2:])
                    
            elif len(d)>14 and len(d)<=21: #14-21，拆分为三个table
                d_reshape={}
                if len(d)==15 or len(d)==18 or len(d)==21:
                    d_reshape["table1"]=dict(list(d.items())[:len(d)//3]) 
                    d_reshape["table2"]=dict(list(d.items())[len(d)//3:len(d)//3*2])      
                    d_reshape["table3"]=dict(list(d.items())[len(d)//3*2:])         
                elif len(d)==16 or len(d)==19: 
                    d_reshape["table1"]=dict(list(d.items())[:len(d)//3+1])
                    d_reshape["table2"]=dict(list(d.items())[len(d)//3+1:len(d)//3*2+1])
                    d_reshape["table3"]=dict(list(d.items())[len(d)//3*2+1:]) 
                elif len(d)==17 or len(d)==20: 
                    d_reshape["table1"]=dict(list(d.items())[:len(d)//3+1])
                    d_reshape["table2"]=dict(list(d.items())[len(d)//3+1:len(d)//3*2+2])
                    d_reshape["table3"]=dict(list(d.items())[len(d)//3*2+2:])  

            elif len(d)>21 and len(d)<=28: #21-28，拆分为四个table
                d_reshape={}
                if len(d)==24 or len(d)==28:
                    d_reshape["table1"]=dict(list(d.items())[:len(d)//4]) 
                    d_reshape["table2"]=dict(list(d.items())[len(d)//4:len(d)//4*2])      
                    d_reshape["table3"]=dict(list(d.items())[len(d)//4*2:len(d)//4*3]) 
                    d_reshape["table4"]=dict(list(d.items())[len(d)//4*3:])        

            colspan_num=[] 
            for key,value in d_reshape.items():
                colspan_num.append(len(value))
            
            print(Carryover_dict)
            print(d_reshape)

            if len(textlist_special)!=0:
                return {"Carryover_dict":Carryover_dict,"d_reshape":d_reshape,"Systermnum":len(Carryover_systermnum),"textlist":textlist_special,"serial":len(textlist_special)+1,
                "Carryover_conclusion":Carryover_conclusion,"colspan_num":colspan_num}
            else:
                return {"Carryover_dict":Carryover_dict,"d_reshape":d_reshape,"Systermnum":len(Carryover_systermnum),"textlist":textlist_general,"serial":len(textlist_general)+1,
                "Carryover_conclusion":Carryover_conclusion,"colspan_num":colspan_num}

        # 多系统,先不做判断
        else:
            if len(textlist_special)!=0:
                return {"Carryover_dict":Carryover_dict,"normnum":len(Carryover_norm),"Systermnum":len(Carryover_systermnum),"textlist":textlist_special,"serial":len(textlist_special)+1,
                "Carryover_conclusion":Carryover_conclusion}
            else:
                return {"Carryover_dict":Carryover_dict,"normnum":len(Carryover_norm),"Systermnum":len(Carryover_systermnum),"textlist":textlist_general,"serial":len(textlist_general)+1,
                "Carryover_conclusion":Carryover_conclusion}


    except:
        pass


    # 2 21个样本的做法

    '''
    注释:需要生成一个字典Carryover_dict,数据格式如下：
    print(Carryover_dict):
    {化合物1:[L01-L21(21个原始数据),X1,X2,SD1],化合物2:[L01-L21(21个原始数据),X1,X2,SD1]}
    '''

    # dataCarryover = Carryover2.objects.filter(reportinfo_id=id)

    # if dataCarryover:
    #     Carryover_dict={}  #最终需要的字典  

    #     for item in dataCarryover:
    #         group=[]
    #         group.append(item.L01)
    #         group.append(item.L02)
    #         group.append(item.L03)
    #         group.append(item.H04)
    #         group.append(item.H05)
    #         group.append(item.L06)
    #         group.append(item.H07)
    #         group.append(item.H08)
    #         group.append(item.L09)
    #         group.append(item.L10)
    #         group.append(item.L11)
    #         group.append(item.L12)
    #         group.append(item.H13)
    #         group.append(item.H14)
    #         group.append(item.L15)
    #         group.append(item.H16)
    #         group.append(item.H17)
    #         group.append(item.L18)
    #         group.append(item.H19)
    #         group.append(item.H20)
    #         group.append(item.L21)
    #         group.append(item.X1)
    #         group.append(item.X2)
    #         group.append(new_round(float(item.X2)-float(item.X1),2))
    #         group.append(item.SD1)
    #         group.append(new_round(3*float(item.SD1),2))
    #         Carryover_dict[item.norm]=group

    #     Carryover_conclusion="X2-X1均小于3SD1，说明系统无携带效应，满足检测要求。"


    #     if len(textlist_special)!=0:
    #         return {"Carryover_dict":Carryover_dict,"textlist":textlist_special,"serial":len(textlist_special)+1,
    #         "Carryover_conclusion":Carryover_conclusion}
    #     else:
    #         return {"Carryover_dict":Carryover_dict,"textlist":textlist_general,"serial":len(textlist_general)+1,
    #         "Carryover_conclusion":Carryover_conclusion}