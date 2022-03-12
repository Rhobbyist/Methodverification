from django.http import HttpResponse, HttpResponseRedirect
import numpy as np
import xlrd
from docx import Document
import math
from report.models import *
from report.effectnum import *
from report.layout import *
from datetime import datetime
import re

# 重复性精密度数据读取
def IntraP_fileread(files, reportinfo, namejmd, project, platform, manufacturers, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB):

    # 第一步:后台数据抓取（最小样本数，最大允许CV）
    id1 = Special.objects.get(project=project).id
    id2 = Repeatprecisionspecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if Repeatprecisionspecialmethod.objects.filter(repeatprecisionspecial=id2):
        lownumber = Repeatprecisionspecialmethod.objects.get(repeatprecisionspecial=id2).minSample  # 最小样本数
        maxCV = Repeatprecisionspecialmethod.objects.get(repeatprecisionspecial=id2).maxCV  # 最大允许CV

    else:
        general_1 = General.objects.get(name="通用性项目")
        jmd_general = Repeatprecisiongeneral.objects.get(general=general_1)
        lownumber = Repeatprecisiongeneralmethod.objects.get(repeatprecisiongeneral=jmd_general).minSample  # 最小样本数
        maxCV = Repeatprecisiongeneralmethod.objects.get(repeatprecisiongeneral=jmd_general).maxCV  # 最大允许CV

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)

    #  第二步:开始文件读取
    '''
    数据读取完毕后,需要生成一个列表jmdone,数据格式如下：
    [ [[化合物1低浓度结果列表],[化合物1中浓度结果列表],[化合物1高浓度结果列表]], [[化合物2低浓度结果列表],[化合物2中浓度结果列表],[化合物2高浓度结果列表]] ]
    '''

    # 设定实验号1，2，3...
    indexone = []
    for i in range(1, 50):
        indexone.append(str(i))

    # 创新第二步需要生成的结果容器
    jmdone = []  # 需要生成的列表
    jmdnum = 0  # 确定一共有多少数据，方便后面判断数据量是否大于20(lownumber)个
    lownum = 0  # 如果低中高中的任一列表为空集，则数字加1，最后html中判断如果数字存在（不为0），则代表该浓度为空列表，无数据，则HTML中不显示
    mediannum = 0
    highnum = 0

    # 各仪器平台及各仪器厂家数据读取
    for file in files:
        if platform == "液质":
            if manufacturers == "Agilent":
                # 1 读取csv文件（Agilent）
                file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)):
                    if len(lines[i]) != 0:
                        # 以逗号分隔字符串,但忽略双引号内的逗号
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                    else:
                        lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i]  # 最后一行如为空行，则删除该元素

                # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                norm = []  # 化合物列表
                for j in range(len(lines[0])):  # 从第一行开始
                    if "-Q Results" in lines[0][j]:
                        # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                        if lines[0][j].split("-Q")[0][0] != '"':
                            norm.append(lines[0][j].split("-Q")[0])
                        else:
                            norm.append(lines[0][j].split("-Q")[0][1:])

                # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                nameindex = 0  # 实验号索引
                concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                for j in range(len(lines[1])):  # 从第二行开始
                    if lines[1][j] == "Sample Name":
                        nameindex = j
                    elif lines[1][j] == "Final Conc.":
                        concindex.append(j)

                # 匹配原始数据中与重复性精密度相关(实验号前含有"IntraP-)的行
                for j in range(len(norm)):
                    low = []  # 低浓度列表
                    median = []  # 中浓度列表
                    high = []  # 高浓度列表
                    normlist = []  # 一个化合物的低中高三个浓度合并列表
                    for i in range(len(lines)):  # 循环原始数据中的每一行
                        if "IntraP-L" in lines[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                            if j < 1:  # 如有多个化合物，只循环添加第一个化合物的样本量，否则样本量数目会重复添加
                                jmdnum += 1  # 样本量加1,最终需要将此数目与设定的最小样本数相比          
                            low.append(effectnum(lines[i][concindex[j]], digits)) 
                        elif "IntraP-M" in lines[i][nameindex]:
                            median.append(effectnum(lines[i][concindex[j]], digits))
                        elif "IntraP-H" in lines[i][nameindex]:
                            high.append(effectnum(lines[i][concindex[j]], digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

            elif manufacturers == "岛津":
                # 读取txt文件
                content = []
                for line in file:
                    content.append(line.decode("UTF-8").replace("\r\n", "").split("\t")) # windows下
                    # content.append(line.decode("GB2312").replace("\r\n", "").split("\t")) # linux下

                nameindex = 0
                concindex = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行

                for i in range(len(content)):
                    if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称
                        norm.append(content[i][1])
                        norm_row.append(i)

                for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                    if content[2][i] == "数据文件名":
                        nameindex = i
                    elif content[2][i] == "浓度":
                        concindex = i

                # print(norm_row) [1, 120, 239, 358, 477, 596, 715, 834, 953, 1072, 1191]

                for j in range(len(norm)):
                    low = []  # 低浓度列表
                    median = []  # 中浓度列表
                    high = []  # 高浓度列表
                    normlist = []  # 一个化合物的低中高三个浓度合并列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            if "JMD-L" in content[i][nameindex]:
                                if j < 1:  # 第一个化合物的样本量即为每个化合物的样本量
                                    jmdnum += 1
                                low.append(effectnum(content[i][concindex], digits))
                            elif "JMD-M" in content[i][nameindex]:
                                median.append(effectnum(content[i][concindex], digits))
                            elif "JMD-H" in content[i][nameindex]:
                                high.append(effectnum(content[i][concindex], digits))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], len(content)):
                            if "JMD-L" in content[i][nameindex]:
                                low.append(effectnum(content[i][concindex], digits))
                            elif "JMD-M" in content[i][nameindex]:
                                median.append(effectnum(content[i][concindex], digits))
                            elif "JMD-H" in content[i][nameindex]:
                                high.append(effectnum(content[i][concindex], digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

            elif manufacturers == "Waters":
                # 若是最新的 2.0.1 版本的xlrd包，只支持 .xls 文件，读取.xlsx文件会报错。若要正常读取，需安装旧版本的xlrd：pip3 install xlrd==1.2.0
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行
                for i in range(nrows):
                    if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):  # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                        norm.append(file_data.row_values(i)[0].split(":")[1].strip()) # strip()的作用是去除前后空格
                        norm_row.append(i)    

                print(norm_row)          

                nameindex = 0
                concindex = 0
                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "Name":
                        nameindex = i
                    elif "实际浓度" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                for j in range(len(norm)):
                    low = []
                    median = []
                    high = []
                    normlist = []
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            if "JMD-L" in file_data.row_values(i)[nameindex]:
                                if j < 1:  # 第一个化合物的样本量即为每个化合物的样本量
                                    jmdnum += 1
                                low.append(effectnum(file_data.row_values(i)[concindex], digits))
                            elif "JMD-M" in file_data.row_values(i)[nameindex]:
                                median.append(effectnum(file_data.row_values(i)[concindex], digits))
                            elif "JMD-H" in file_data.row_values(i)[nameindex]:
                                high.append(effectnum(file_data.row_values(i)[concindex], digits))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], nrows):
                            if "JMD-L" in file_data.row_values(i)[nameindex]:
                                low.append(effectnum(file_data.row_values(i)[concindex], digits))
                            elif "JMD-M" in file_data.row_values(i)[nameindex]:
                                median.append(effectnum(file_data.row_values(i)[concindex], digits))
                            elif "JMD-H" in file_data.row_values(i)[nameindex]:
                                high.append(effectnum(file_data.row_values(i)[concindex], digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

            # Thermo厂家需先在后台管理系统中设置本项目的化合物名称，以便查找上传文件中相应化合物的表格
            elif manufacturers == "Thermo":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm = []  # Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                sheetindex = []  # 需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    if data.sheet_names()[index] in PTnorm:
                        norm.append(data.sheet_names()[index])
                        sheetindex.append(index)

                print(norm)
                # 循环读取每个sheet工作表,即为每个化合物的表
                for index in range(len(sheetindex)):
                    file_data = data.sheets()[sheetindex[index]]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    # 第一行确定samplename和浓度所在列
                    nameindex = 0
                    concindex = 0
                    for i in range(len(file_data.row_values(0))):
                        if file_data.row_values(0)[i] == "Compound":
                            nameindex = i
                        elif file_data.row_values(0)[i] == "Calculated Amt":
                            concindex = i

                    low = []
                    median = []
                    high = []
                    normlist = []
                    for i in range(nrows):
                        if "JMD-L" in file_data.row_values(i)[nameindex]:
                            if index < 1:  # 如有多个化合物，只循环添加第一个化合物的样本量，否则样本量数目会重复添加
                                jmdnum += 1  # 样本量加1
                            low.append(effectnum(file_data.row_values(i)[concindex],digits))
                        elif "JMD-M" in file_data.row_values(i)[nameindex]:
                            median.append(effectnum(file_data.row_values(i)[concindex],digits))
                        elif "JMD-H" in file_data.row_values(i)[nameindex]:
                            high.append(effectnum(file_data.row_values(i)[concindex],digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

            # AB厂家需先在后台设置化合物和相应离子对数值，以便判断需要读取的表格（定量表格）和不需要读取的表格(定性表格)
            # 后台管理系统中进行设置时，化合物的设置顺序需与文件中的化合物排列顺序一致，即从上向下进行设置
            elif manufacturers == "AB":

                # 测试文件中的三个化合物（1  Clozapine定性 (327.2 / 192.2)，Clozapine定量 (327.2 / 270.2);
                #                     2  Sertraline定性 (306.1 / 275.1)，Sertraline定量 (306.1 / 159.1)）
                #                     3  Aripiprazole定性 (448.1 / 176.2),Aripiprazole定量 (448.1 / 285.2)

                # 定义化合物列表，列表统一命名为norm
                norm = normAB

                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = [] 

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                tables = file_data.tables  # 获取文件中的表格集

                # 循环定量表格的索引
                for k in range(len(tableindex)):
                     # 获取文件中的定量表格
                    tablequantify = tables[tableindex[k]] 

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表

                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex = 0
                    concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列
                    
                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name":
                            nameindex = i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex = i

                    low = []  # 低浓度列表
                    median = []  # 中浓度列表
                    high = []  # 高浓度列表
                    normlist = []  # 一个化合物的低中高三个浓度合并列表
                    for i in range(len(rowdatagatherlist)):
                        if "IntraP-L" in rowdatagatherlist[i][nameindex]:
                            if k < 1:
                                jmdnum += 1
                            low.append(effectnum(rowdatagatherlist[i][concindex], digits))
                        elif "IntraP-M" in rowdatagatherlist[i][nameindex]:
                            median.append(effectnum(rowdatagatherlist[i][concindex], digits))
                        elif "IntraP-H" in rowdatagatherlist[i][nameindex]:
                            high.append(effectnum(rowdatagatherlist[i][concindex], digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

                print(jmdone)
                print(jmdnum)

        elif platform == "液相":
            if manufacturers == "Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]  # 默认只读取第一个工作簿
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm = []  # 化合物列表
                norm_row = []  # 化合物所在行
                for j in range(nrows):
                    # 如果某一行的第一个元素为'化合物'，则添加第三个元素进入化合物列表
                    if file_data.row_values(j)[0] == "化合物:":
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex = 0
                concindex = 0

                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "样品名称":
                        nameindex = i
                    elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                for j in range(len(norm)):
                    low = []  # 低浓度列表
                    median = []  # 中浓度列表
                    high = []  # 高浓度列表
                    normlist = []  # 一个化合物的低中高三个浓度合并列表
                    if j < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j], norm_row[j+1]):
                            if "JMD-L" in file_data.row_values(i)[nameindex]:
                                if j < 1:  # 第一个化合物的样本量即为每个化合物的样本量
                                    jmdnum += 1
                                low.append(effectnum(file_data.row_values(i)[concindex],digits))
                            elif "JMD-M" in file_data.row_values(i)[nameindex]:
                                median.append(effectnum(file_data.row_values(i)[concindex],digits))
                            elif "JMD-H" in file_data.row_values(i)[nameindex]:
                                high.append(effectnum(file_data.row_values(i)[concindex],digits))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j], nrows):
                            if "JMD-L" in file_data.row_values(i)[nameindex]:
                                low.append(effectnum(file_data.row_values(i)[concindex],digits))
                            elif "JMD-M" in file_data.row_values(i)[nameindex]:
                                median.append(effectnum(file_data.row_values(i)[concindex],digits))
                            elif "JMD-H" in file_data.row_values(i)[nameindex]:
                                high.append(effectnum(file_data.row_values(i)[concindex],digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

        elif platform == "ICP-MS":
            # ICP-MS平台Agilent厂家需先在后台管理系统中设置本项目的化合物名称，以便查找上传文件中相应化合物的表格
            if manufacturers == "Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                # 从第一行确定化合物名称
                norm = []
                for j in range(ncols):
                    for i in PTnorm:
                        if i in file_data.row_values(0)[j]:
                            norm.append(i)

                # 从第二行确定实验号（Sample Name）的索引和化合物浓度索引
                nameindex = 0  # 实验号索引
                concindex = []  # 浓度索引
                for j in range(ncols):
                    if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j
                    elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                        concindex.append(j)

                # 匹配原始数据中与精密度相关(实验号前含有"JMD-L-"或"JMD-L-"或"JMD-H")的行
                for j in range(len(concindex)):
                    low = []  # 低浓度列表
                    median = []  # 中浓度列表
                    high = []  # 高浓度列表
                    normlist = []  # 一个化合物的低中高三个浓度合并列表
                    for i in range(2, nrows):
                        if "JMD-L" in file_data.row_values(i)[nameindex]:
                            if j < 1:
                                jmdnum += 1
                            low.append(effectnum(file_data.row_values(i)[concindex[j]],digits))
                        elif "JMD-M" in file_data.row_values(i)[nameindex]:
                            median.append(effectnum(file_data.row_values(i)[concindex[j]],digits))
                        elif "JMD-H" in file_data.row_values(i)[nameindex]:
                            high.append(effectnum(file_data.row_values(i)[concindex[j]],digits))

                    normlist.append(low)
                    normlist.append(median)
                    normlist.append(high)
                    jmdone.append(normlist)

        ########文件读取完毕#######
    print(jmdone)

    #  第三步:计算平均值，标准差，CV

    '''
    注释2:最终需要生成一个字典jmd_dict,分别对应不同化合物的数据，数据格式如下：
    {化合物1:[[SAMPLE1结果列表],[SAMPLE2结果列表],[SAMPLE3结果列表]...],化合物2:[[SAMPLE1结果列表],[SAMPLE2结果列表],[SAMPLE3结果列表]...]}
    '''
    # 创新第三步需要生成的结果容器
    jmd_dict = {} 

    # 先判断低中高三个浓度是否都有数据。如果低中高中的任一列表为空集，则数字加1，最后html中判断如果数字存在（不为0），则代表该浓度为空列表，无数据，则HTML中不显示
    for i in jmdone:
        if i[0] == []:  # 按第二步添加列表的顺序,i[0]为低浓度,i[1]为中浓度,i[2]为高浓度
            lownum += 1
        elif i[1] == []:
            mediannum += 1
        elif i[2] == []:
            highnum += 1

    # 在jmdone列表里添加实验号
    jmdtwo = []  
    for i in jmdone:
        normlist2 = [] 
        for j in range(len(i[0])):
            group = []
            group.append(indexone[j])
            for k in i:
                if k != []:  # 可能没有低中高浓度的其中一个或几个,所以需要进行判断
                    group.append(k[j])
            normlist2.append(group)
        jmdtwo.append(normlist2)
    
    print(jmdtwo)

    # 计算均值，标准差，CV
    meanlist = []  # 均值列表,可能不止一个化合物,格式为[ ["均值",低浓度均值,中浓度均值,高浓度均值], ["均值",低浓度均值,中浓度均值,高浓度均值] ]
    sdlist = []  # 标准差列表
    cvlist = []  # cv列表
    cvjudgenum = 0  # 判断超过cv上限的值有多少个，不管有几个化合物，只要此值大于1，都不将数据保存到数据库
    for i in jmdone:
        meangroup = ["均值"]
        sdgroup = ["标准差"]
        cvgroup = ["CV(%)"]
        for j in i:
            if j!= []:
                j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                meangroup.append(mean(j2))
                sdgroup.append(sd(j2))
                cvgroup.append(cv(j2))
        meanlist.append(meangroup)
        sdlist.append(sdgroup)
        cvlist.append(cvgroup)

    # 列表末尾添加均值,标准差,CV
    for i in range(len(jmdtwo)):
        jmdtwo[i].append(meanlist[i])
        jmdtwo[i].append(sdlist[i])
        jmdtwo[i].append(cvlist[i])

    # 添加键到字典中
    for i in range(len(norm)):
        jmd_dict[norm[i]] = jmdtwo[i]
    
    print(jmd_dict)

    #  第四步:数据存入数据库
    # 样本量大于最小样本量，并且超过cv上限的值等于0才将数据存入数据库中
    if cvjudgenum == 0 and jmdnum >= lownumber:
        insert_list = []
        for key, value in jmd_dict.items():
            for i in value:
                if lownum != 0 and mediannum == 0 and highnum == 0:  # 低浓度无数据，中高浓度有数据，只存中高浓度
                    insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,
                                       norm=key, Experimentnum=i[0], median=i[1], high=i[2]))
                elif lownum == 0 and mediannum != 0 and highnum == 0:  # 中浓度无数据，低高浓度有数据，只存低高浓度
                    insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,
                                       norm=key, Experimentnum=i[0], low=i[1], high=i[2]))
                elif lownum == 0 and mediannum == 0 and highnum != 0:  # 高浓度无数据，低中浓度有数据，只存低中浓度
                    insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,
                                       norm=key, Experimentnum=i[0], low=i[1], median=i[2]))
                else:
                    insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,
                                       norm=key, Experimentnum=i[0], low=i[1], median=i[2], high=i[3]))

        JMD.objects.bulk_create(insert_list)

    return {'jmd_dict': jmd_dict, "namejmd": namejmd, "nrows": jmdnum, "lownumber": int(lownumber), "maxCV": maxCV, "lownum": lownum, "mediannum": mediannum, "highnum": highnum, "Unit": Unit}

# 中间精密度数据读取
def InterP_fileread(files, reportinfo, namejmd, project, platform, manufacturers, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB):

    # 第一步:后台数据抓取（最小样本数，最大允许CV,化合物个数）
    id1 = Special.objects.get(project=project).id
    id2 = Repeatprecisionspecial.objects.get(special_id=id1).id
    norm_num =  Special.objects.get(project=project).Number_of_compounds
    
    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if Interprecisionspecialmethod.objects.filter(interprecisionspecial=id2):
        lownumber = Interprecisionspecialmethod.objects.get(interprecisionspecial=id2).minSample  # 最小样本数
        maxCV = Interprecisionspecialmethod.objects.get(interprecisionspecial=id2).maxCV  # 最大允许CV

    else:
        general_1 = General.objects.get(name="通用性项目")
        jmd_general = Interprecisiongeneral.objects.get(general=general_1)
        lownumber = Interprecisiongeneralmethod.objects.get(interprecisiongeneral=jmd_general).minSample  # 最小样本数
        maxCV = Interprecisiongeneralmethod.objects.get(interprecisiongeneral=jmd_general).maxCV  # 最大允许CV

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)


    #  第二步:开始文件读取

    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个列表jmdone,数据格式如下：
    print(jmdone):
    [ [[D3low],[D3median],[D3high]], [[D2low],[D2median],[D2high]] ]
    '''

    # 自己设定实验号1，2，3...
    indexone = []
    for i in range(1, 50):
        indexone.append(str(i))

    # 头部定义相关需要提取生成的结果
    jmdone = []  # 定义需要生成的列表
    jmdnum = 0  # 确定一共有多少数据，方便后面判断数据量是否大于20(number)个
    lownum = 0  # 低浓度数据量，如果低浓度数据量为0，则在前端模板不显示，也不将数据存入数据库，下同
    mediannum = 0
    highnum = 0
    norm = []  # 指标列表

    for k in range(norm_num):
        low = []  # 低浓度列表
        median = []  # 中浓度列表
        high = []  # 高浓度列表
        normlist = []  # 一个化合物的低中高三个浓度合并列表
        for index, file in enumerate(files):
            if platform == "液质":
                if manufacturers == "Agilent":
                    # 1 读取csv文件（Agilent）https://www.jianshu.com/p/0d15ed85df2b
                    file.seek(0)
                    file_data = file.read().decode('utf-8')
                    lines = file_data.split('\r\n')
                    for i in range(len(lines)):
                        if len(lines[i]) != 0:
                            # 以逗号分隔字符串,但忽略双引号内的逗号
                            lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                            # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                        else:
                            lines[i] = re.split(r',\s*(?![^"]*\"\,)', lines[i])
                            del lines[i]  # 最后一行如为空行，则删除该元素

                    # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                    if k < 1 and index < 1:   # 第一个循环即可确定，无需重复循环占用内存空间
                        for j in range(len(lines[0])):  # 从第一行开始
                            if "-Q Results" in lines[0][j]:
                                # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                                if lines[0][j].split("-Q")[0][0] != '"':
                                    norm.append(lines[0][j].split("-Q")[0])
                                else:
                                    norm.append(lines[0][j].split("-Q")[0][1:])

                    # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                    nameindex = 0  # 实验号索引
                    concindex = []  # 浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                    for j in range(len(lines[1])):  # 从第二行开始
                        if lines[1][j] == "Sample Name":
                            nameindex = j
                        elif lines[1][j] == "Final Conc.":
                            concindex.append(j)

                    # 匹配原始数据中与重复性精密度相关(实验号前含有"InterP-)的行
                    for i in range(len(lines)):  # 循环原始数据中的每一行
                        if "InterP-L" in lines[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                            if k < 1:  # 如有多个化合物，只循环添加第一个化合物的样本量，否则样本量数目会重复添加
                                jmdnum += 1  # 样本量加1,最终需要将此数目与设定的最小样本数相比          
                            low.append(effectnum(lines[i][concindex[k]], digits)) 
                        elif "InterP-M" in lines[i][nameindex]:
                            median.append(effectnum(lines[i][concindex[k]], digits))
                        elif "InterP-H" in lines[i][nameindex]:
                            high.append(effectnum(lines[i][concindex[k]], digits))

                elif manufacturers == "Waters":
                    data = xlrd.open_workbook(
                        filename=None, file_contents=file.read())  # 读取表格
                    file.seek(0, 0)  # 循环读取同一个文件两遍，需加此句代码移动文件读取指针到开头，否则会报错
                    file_data = data.sheets()[0]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    norm_row = []  # 化合物所在行
                    for j in range(nrows):
                        for i in PTnorm:
                            if i in str(file_data.row_values(j)[0]):
                                norm_row.append(j)

                    nameindex = 0
                    concindex = 0
                    # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    for i in range(len(file_data.row_values(norm_row[0]+2))):
                        if file_data.row_values(norm_row[0]+2)[i] == "Name":
                            nameindex = i
                        elif "实际浓度" in file_data.row_values(norm_row[0]+2)[i]:
                            concindex = i

                    if k < norm_num-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[k], norm_row[k+1]):
                            if "L-" in file_data.row_values(i)[nameindex]:
                                if k < 1:  # 第一个化合物的样本量即为每个化合物的样本量
                                    jmdnum += 1
                                low.append(
                                    float(file_data.row_values(i)[concindex]))
                            elif "M-" in file_data.row_values(i)[nameindex]:
                                median.append(
                                    float(file_data.row_values(i)[concindex]))
                            elif "H-" in file_data.row_values(i)[nameindex]:
                                high.append(
                                    float(file_data.row_values(i)[concindex]))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[k], nrows):
                            if "L-" in file_data.row_values(i)[nameindex]:
                                low.append(
                                    float(file_data.row_values(i)[concindex]))
                            elif "M-" in file_data.row_values(i)[nameindex]:
                                median.append(
                                    float(file_data.row_values(i)[concindex]))
                            elif "H-" in file_data.row_values(i)[nameindex]:
                                high.append(
                                    float(file_data.row_values(i)[concindex]))

                elif manufacturers == "Thermo":
                    data = xlrd.open_workbook(
                        filename=None, file_contents=file.read())  # 读取表格
                    file.seek(0, 0)  # 循环读取同一个文件两遍，需加此句代码移动文件读取指针到开头，否则会报错

                    # Thermo需要依据在后台管理系统里设置的化合物名称判断需要抓取的化合物表格
                    # 后台管理系统-各项目参数设置-PT指标设置里找到化合物名称
                    if k < 1 and index < 1:
                        Thermo = Special.objects.get(project=project)
                        pt_special = PTspecial.objects.get(special=Thermo)
                        pt_accept = PTspecialaccept.objects.filter(
                            pTspecial=pt_special)
                        PTnorm = []  # 待测物质列表
                        for i in pt_accept:
                            PTnorm.append(i.norm)

                    if k < 1 and index < 1:  # 第一个化合物确定sheetindex
                        sheetindex = []  # 需要的化合物所在sheet索引列表
                        for index in range(len(data.sheet_names())):
                            if data.sheet_names()[index] in PTnorm:
                                norm.append(data.sheet_names()[index])
                                sheetindex.append(index)

                    file_data = data.sheets()[sheetindex[k]]
                    nrows = file_data.nrows
                    ncols = file_data.ncols

                    # 第一行确定samplename和浓度所在列
                    nameindex = 0
                    concindex = 0
                    for i in range(len(file_data.row_values(0))):
                        if file_data.row_values(0)[i] == "Compound":
                            nameindex = i
                        elif file_data.row_values(0)[i] == "Calculated Amt":
                            concindex = i

                    for i in range(nrows):
                        if "L-" in file_data.row_values(i)[nameindex]:
                            if k < 1:  # 如有多个化合物，只循环添加第一个化合物的样本量，否则样本量数目会重复添加
                                jmdnum += 1  # 样本量加1
                            low.append(
                                float(file_data.row_values(i)[concindex]))
                        elif "M-" in file_data.row_values(i)[nameindex]:
                            median.append(
                                float(file_data.row_values(i)[concindex]))
                        elif "H-" in file_data.row_values(i)[nameindex]:
                            high.append(
                                float(file_data.row_values(i)[concindex]))

                elif manufacturers == "岛津":
                    # 读取txt文件
                    content = []
                    for line in file:
                        content.append(line.decode("GB2312").replace(
                            "\r\n", "").split("\t"))

                    nameindex = 0
                    concindex = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                    # norm=[] #化合物列表
                    norm_row = []  # 化合物所在行

                    for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                        if content[2][i] == "数据文件名":
                            nameindex = i
                        elif content[2][i] == "浓度":
                            concindex = i

                    for i in range(len(content)):
                        if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称
                            norm.append(content[i][1])
                            norm_row.append(i)

                    # print(norm_row) [1, 120, 239, 358, 477, 596, 715, 834, 953, 1072, 1191]

                    if k < norm_num-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[k], norm_row[k+1]):
                            if "L-" in content[i][nameindex]:
                                if k < 1:  # 第一个化合物的样本量即为每个化合物的样本量
                                    jmdnum += 1
                                low.append(float(content[i][concindex]))
                            elif "M-" in content[i][nameindex]:
                                median.append(float(content[i][concindex]))
                            elif "H-" in content[i][nameindex]:
                                high.append(float(content[i][concindex]))

                    else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[k], len(content)):
                            if "L-" in content[i][nameindex]:
                                low.append(float(content[i][concindex]))
                            elif "M-" in content[i][nameindex]:
                                median.append(float(content[i][concindex]))
                            elif "H-" in content[i][nameindex]:
                                high.append(float(content[i][concindex]))

                elif manufacturers == "AB":
                    if k < 1 and index < 1:
                        norm = normAB

                    # 获取上传的文件
                    file_data = Document(file)

                    # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                    paragraphs = [] 

                    # 若标题长度为0或为换行等，不添加进入标题内容列表
                    for p in file_data.paragraphs:
                        if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                            paragraphs.append(p.text.strip())

                    # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                    tableindex = []
                    for i in range(len(paragraphs)):
                        for j in range(len(ZP_Method_precursor_ion)):
                            if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                tableindex.append(2*i+1)

                    tables = file_data.tables  # 获取文件中的表格集

                    tablequantify = tables[tableindex[k]]  # 获取文件中的相关表格

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表

                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex = 0
                    concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列

                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name":
                            nameindex = i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex = i

                    for i in range(len(rowdatagatherlist)):
                        if "InterP-L" in rowdatagatherlist[i][nameindex]:
                            if k < 1:
                                jmdnum += 1
                            low.append(effectnum(rowdatagatherlist[i][concindex], digits))
                        elif "InterP-M" in rowdatagatherlist[i][nameindex]:
                            median.append(effectnum(rowdatagatherlist[i][concindex], digits))
                        elif "InterP-H" in rowdatagatherlist[i][nameindex]:
                            high.append(effectnum(rowdatagatherlist[i][concindex], digits))

            elif platform == "液相":
                data = xlrd.open_workbook(
                    filename=None, file_contents=file.read())  # 读取表格
                file.seek(0, 0)  # 循环读取同一个文件两遍，需加此句代码移动文件读取指针到开头，否则会报错
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                norm_row = []  # 化合物所在行
                for j in range(nrows):
                    # 如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                    if file_data.row_values(j)[0] == "化合物:":
                        if file_data.row_values(j)[2] not in norm:
                            norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex = 0
                concindex = 0
                # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                for i in range(len(file_data.row_values(norm_row[0]+2))):
                    if file_data.row_values(norm_row[0]+2)[i] == "样品名称":
                        nameindex = i
                    elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                        concindex = i

                if k < norm_num-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                    for i in range(norm_row[k], norm_row[k+1]):
                        if "L-" in file_data.row_values(i)[nameindex]:
                            if k < 1:  # 第一个化合物的样本量即为每个化合物的样本量
                                jmdnum += 1
                            low.append(
                                float(file_data.row_values(i)[concindex]))
                        elif "M-" in file_data.row_values(i)[nameindex]:
                            median.append(
                                float(file_data.row_values(i)[concindex]))
                        elif "H-" in file_data.row_values(i)[nameindex]:
                            high.append(
                                float(file_data.row_values(i)[concindex]))

                else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                    for i in range(norm_row[k], nrows):
                        if "L-" in file_data.row_values(i)[nameindex]:
                            low.append(
                                float(file_data.row_values(i)[concindex]))
                        elif "M-" in file_data.row_values(i)[nameindex]:
                            median.append(
                                float(file_data.row_values(i)[concindex]))
                        elif "H-" in file_data.row_values(i)[nameindex]:
                            high.append(
                                float(file_data.row_values(i)[concindex]))

            elif platform == "ICP-MS":
                data = xlrd.open_workbook(
                    filename=None, file_contents=file.read())  # 读取表格
                file.seek(0, 0)  # 循环读取同一个文件两遍，需加此句代码移动文件读取指针到开头，否则会报错
                file_data = data.sheets()[0]
                nrows = file_data.nrows
                ncols = file_data.ncols

                # 从第一行确定化合物名称
                for j in range(ncols):
                    for i in PTnorm:
                        if i in file_data.row_values(0)[j] and i not in norm:
                            norm.append(i)

                # 从第二行确定实验号和化合物浓度索引
                nameindex = 0  # 实验号索引
                concindex = []  # 浓度索引
                for j in range(ncols):
                    if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j
                    elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                        conindex.append(j)

                # 匹配原始数据中与精密度相关(实验号前含有"L-"或"M-"或"H-")的行
                for i in range(2, nrows):
                    if "L-" in file_data.row_values(i)[nameindex]:
                        if k < 1:
                            jmdnum += 1
                        low.append(file_data.row_values(i)[conindex[k]])
                    elif "M" in file_data.row_values(i)[nameindex]:
                        median.append(file_data.row_values(i)[conindex[k]])
                    elif "H" in file_data.row_values(i)[nameindex]:
                        high.append(file_data.row_values(i)[conindex[k]])

        normlist.append(low)
        normlist.append(median)
        normlist.append(high)
        jmdone.append(normlist)

    ####文件读写完毕####


    #  第三步:文件读取完毕后的操作
    # norm=PTnorm
    '''
    注释2:最终需要生成一个字典jmd_dict,数据格式如下：
    print(jmd_dict):
    {"D3":[[SAMPLE1],[SAMPLE2],[SAMPLE3]...],"D2":[[SAMPLE1],[SAMPLE2],[SAMPLE3]...]}
    '''

    # 创新第三步需要生成的结果容器
    jmd_dict = {} 

    # 先判断低中高三个浓度是否都有数据。如果低中高中的任一列表为空集，则数字加1，最后html中判断如果数字存在（不为0），则代表该浓度为空列表，无数据，则HTML中不显示
    for i in jmdone:
        if i[0] == []:
            lownum += 1
        elif i[1] == []:
            mediannum += 1
        elif i[2] == []:
            highnum += 1
 
    # 查找数据库中已经提交的数据数量,以便递增实验号
    count = JMD.objects.filter(reportinfo=reportinfo, namejmd=namejmd, norm=norm[0]).count()

    # 在jmdone列表里添加实验号
    jmdtwo = [] 
    for i in jmdone:
        normlist2 = []  # 加入实验号后的临时列表
        for j in range(len(i[0])):  # 列表长度(也是样本量)
            group = []
            group.append(indexone[j+count])
            for k in i:
                if k != []:  # 可能没有低中高浓度的其中一个或几个，如果没有append会报错
                    group.append(k[j])
            normlist2.append(group)
        jmdtwo.append(normlist2)

    # 计算均值，标准差，CV
    meanlist = []  # 均值列表,可能不止一个化合物,格式为[ ["均值",低浓度均值,中浓度均值,高浓度均值], ["均值",低浓度均值,中浓度均值,高浓度均值] ]
    sdlist = []  # 标准差列表
    cvlist = []  # cv列表
    cvjudgenum = 0  # 判断超过cv上限的值有多少个，不管有几个化合物，只有此值大于1，都不将数据保存到数据库
    for i in jmdone:
        meangroup = ["均值"]
        sdgroup = ["标准差"]
        cvgroup = ["CV(%)"]
        for j in i:
            if j!= []:
                j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                meangroup.append(mean(j2))
                sdgroup.append(sd(j2))
                cvgroup.append(cv(j2))
        meanlist.append(meangroup)
        sdlist.append(sdgroup)
        cvlist.append(cvgroup)

    # 列表末尾添加均值,标准差,CV
    for i in range(len(jmdtwo)):
        jmdtwo[i].append(meanlist[i])
        jmdtwo[i].append(sdlist[i])
        jmdtwo[i].append(cvlist[i])

    # 添加键到字典中
    for i in range(len(norm)):
        jmd_dict[norm[i]] = jmdtwo[i]
    
    print(jmd_dict)

    #  第四步:数据存入数据库
    # 如果超过cv上限的值等于0才将数据存入数据库中
    if cvjudgenum == 0 and jmdnum >= lownumber:
        insert_list = []

        for key, value in jmd_dict.items():
            for i in value:
                if lownum != 0 and mediannum == 0 and highnum == 0:  # 低浓度无数据，中高浓度有数据，只存中高浓度
                    if i[0]!="均值" and i[0]!="标准差" and i[0]!="CV(%)":  # 中间精密度不在数据库中保存计算结果，报告预览时重新计算
                        insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,norm=key, Experimentnum=i[0], median=i[1], high=i[2]))
                elif lownum == 0 and mediannum != 0 and highnum == 0:  # 中浓度无数据，低高浓度有数据，只存低高浓度
                    if i[0]!="均值" and i[0]!="标准差" and i[0]!="CV(%)":  # 中间精密度不在数据库中保存计算结果，报告预览时重新计算
                        insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,norm=key, Experimentnum=i[0], low=i[1], high=i[2]))
                elif lownum == 0 and mediannum == 0 and highnum != 0:  # 高浓度无数据，低中浓度有数据，只存低中浓度
                    if i[0]!="均值" and i[0]!="标准差" and i[0]!="CV(%)":  # 中间精密度不在数据库中保存计算结果，报告预览时重新计算
                        insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,norm=key, Experimentnum=i[0], low=i[1], median=i[2]))
                else:
                    if i[0]!="均值" and i[0]!="标准差" and i[0]!="CV(%)":  # 中间精密度不在数据库中保存计算结果，报告预览时重新计算
                        insert_list.append(JMD(reportinfo=reportinfo, namejmd=namejmd,norm=key, Experimentnum=i[0], low=i[1], median=i[2], high=i[3]))

        JMD.objects.bulk_create(insert_list)

    return {'jmd_dict': jmd_dict, "namejmd": namejmd, "nrows": jmdnum, "lownumber": int(lownumber), "maxCV": maxCV, "lownum": lownum, "mediannum": mediannum, "highnum": highnum, "Unit": Unit}

# 重复性精密度数据抓取，关联到最终报告预览界面
def related_PNjmd(id):
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    #特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = Repeatprecisionspecial.objects.get(special=special_1)           
        if Repeatprecisionspecialtexts.objects.filter(repeatprecisionspecial=special_2).count()>0:
            text_special = Repeatprecisionspecialtexts.objects.filter(repeatprecisionspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = Repeatprecisiongeneral.objects.get(general=general_1)
    text_general = Repeatprecisiongeneraltexts.objects.filter(repeatprecisiongeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits

    # 第二步：报告数据提取

    '''
    注释:需要生成一个字典JMD_endreport_dict和一句话JMD_CONCLUSION,数据格式如下:
    JMD_endreport_dict:{"D3":[[SAMPLE1],[SAMPLE2],[SAMPLE3]...],"D2":[[SAMPLE1],[SAMPLE2],[SAMPLE3]...]}
    JMD_CONCLUSION:结果表明25OHD3、25OHD2的变异系数CV分别在x1%-x2%,y1%-y2%范围内,均小于20%,满足检测要求。
    '''

    # 定义需要生成的字典
    JMD_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        PNJMD_data = JMD.objects.filter(reportinfo_id=id, namejmd="重复性精密度")

        JMD_norm = []  # 待测物质列表
        for i in PNJMD_data:
            if i.norm not in JMD_norm:
                JMD_norm.append(i.norm)

        JMD_CV = {}  # CV字典，方便提取CV范围到JMD_CONCLUSION中
        lownum = 0  # 低浓度数据量，如果低浓度数据量为0，则在前端模板不显示，也不将数据存入数据库，下同
        mediannum = 0
        highnum = 0

        for i in JMD_norm:
            middle_list = []  # 每个化合物的数据列表
            JMD_CV[i] = []
            middle_table = JMD.objects.filter(reportinfo_id=id, namejmd="重复性精密度", norm=i)
            for j in middle_table:
                rowlist = []  # 每一行的小列表
                rowlist.append(j.Experimentnum)

                # 没有为每个化合物单独设置有效位数，则调用通用性设置
                if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                    # 判断低中高三个浓度是否都有数据，同文件读取步骤
                    if j.low == "" and j.median != '' and j.high != '':
                        lownum += 1
                        rowlist.append(j.median)
                        rowlist.append(j.high)
                    elif j.low != "" and j.median == '' and j.high != '':
                        mediannum += 1
                        rowlist.append(j.low)
                        rowlist.append(j.high)
                    elif j.low != "" and j.median != '' and j.high == '':
                        highnum += 1
                        rowlist.append(j.low)
                        rowlist.append(j.median)
                    else:
                        rowlist.append(j.low)
                        rowlist.append(j.median)
                        rowlist.append(j.high)
                    middle_list.append(rowlist)

                    # 实验号为CV，添加结果进入CV字典
                    if j.Experimentnum == "CV(%)":
                        if j.low == "" and j.median != '' and j.high != '':
                            JMD_CV[i].append(j.median)
                            JMD_CV[i].append(j.high)
                        elif j.low != "" and j.median == '' and j.high != '':
                            JMD_CV[i].append(j.low)
                            JMD_CV[i].append(j.high)
                        elif j.low != "" and j.median != '' and j.high == '':
                            JMD_CV[i].append(j.low)
                            JMD_CV[i].append(j.median)
                        else:
                            JMD_CV[i].append(j.low)
                            JMD_CV[i].append(j.median)
                            JMD_CV[i].append(j.high)

                else:  # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                    # 原始数据才取有效位数，计算结果(均值，标准差，cv)有效位数保持不变
                    if j.Experimentnum != "CV(%)" and j.Experimentnum != "均值" and j.Experimentnum != "标准差":
                        # 判断低中高三个浓度是否都有数据，同文件读取步骤
                        if j.low == "" and j.median != '' and j.high != '':
                            lownum += 1
                            rowlist.append(effectnum(j.median, Digitsdict[i]))
                            rowlist.append(effectnum(j.high, Digitsdict[i]))
                        elif j.low != "" and j.median == '' and j.high != '':
                            mediannum += 1
                            rowlist.append(effectnum(j.low, Digitsdict[i]))
                            rowlist.append(effectnum(j.high, Digitsdict[i]))
                        elif j.low != "" and j.median != '' and j.high == '':
                            highnum += 1
                            rowlist.append(effectnum(j.low, Digitsdict[i]))
                            rowlist.append(effectnum(j.median, Digitsdict[i]))
                        else:
                            rowlist.append(effectnum(j.low, Digitsdict[i]))
                            rowlist.append(effectnum(j.median, Digitsdict[i]))
                            rowlist.append(effectnum(j.high, Digitsdict[i]))
                        middle_list.append(rowlist)

                    else:
                        # 计算结果(均值，标准差，cv)有效位数保持不变
                        if j.low == "" and j.median != '' and j.high != '':
                            lownum += 1
                            rowlist.append(j.median)
                            rowlist.append(j.high)
                        elif j.low != "" and j.median == '' and j.high != '':
                            mediannum += 1
                            rowlist.append(j.low)
                            rowlist.append(j.high)
                        elif j.low != "" and j.median != '' and j.high == '':
                            highnum += 1
                            rowlist.append(j.low)
                            rowlist.append(j.median)
                        else:
                            rowlist.append(j.low)
                            rowlist.append(j.median)
                            rowlist.append(j.high)
                        middle_list.append(rowlist)

                    if j.Experimentnum == "CV(%)":
                        if j.low == "" and j.median != '' and j.high != '':
                            JMD_CV[i].append(j.median)
                            JMD_CV[i].append(j.high)
                        elif j.low != "" and j.median == '' and j.high != '':
                            JMD_CV[i].append(j.low)
                            JMD_CV[i].append(j.high)
                        elif j.low != "" and j.median != '' and j.high == '':
                            JMD_CV[i].append(j.low)
                            JMD_CV[i].append(j.median)
                        else:
                            JMD_CV[i].append(j.low)
                            JMD_CV[i].append(j.median)
                            JMD_CV[i].append(j.high)

            JMD_dict[i] = middle_list

        JMD_CONCLUSION = "结果表明" + "、" .join(list(JMD_CV.keys()))+"的变异系数CV分别在 "  # 最终需要的结论

        JMD_CVrange = []
        for value in JMD_CV.values():
            JMD_CVrange.append(str(min(value))+"%"+"-"+str(max(value))+"%")

        string = "," .join(list(JMD_CVrange))+"范围内,均小于20%,满足检测要求。"

        for i in string:
            JMD_CONCLUSION = JMD_CONCLUSION + i

        if len(textlist_special) != 0:
            return {"JMD_dict":JMD_dict, "textlist": textlist_special, "serial": len(textlist_special)+1,
                    "JMD_CONCLUSION": JMD_CONCLUSION, "lownum": lownum, "mediannum": mediannum, "highnum": highnum}
        else:
            return {"JMD_dict":JMD_dict, "textlist": textlist_general, "serial": len(textlist_general)+1,
                    "JMD_CONCLUSION": JMD_CONCLUSION, "lownum": lownum, "mediannum": mediannum, "highnum": highnum}

    except:
        pass

# 中间精密度数据抓取，关联到最终报告预览界面
def related_PJjmd(id):
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    #特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = Interprecisionspecial.objects.get(special=special_1)           
        if Interprecisionspecialtexts.objects.filter(interprecisionspecial=special_2).count()>0:
            text_special = Interprecisionspecialtexts.objects.filter(interprecisionspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = Interprecisiongeneral.objects.get(general=general_1)
    text_general = Interprecisiongeneraltexts.objects.filter(interprecisiongeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits

    # 第二步：报告数据提取
    '''
    注释:需要生成一个字典JMD_endreport_dict和结论JMD_CONCLUSION,数据格式如下：
    JMD_endreport_dict:{"D3":[[SAMPLE1],[SAMPLE2],[SAMPLE3]...],"D2":[[SAMPLE1],[SAMPLE2],[SAMPLE3]...]}
    JMD_CONCLUSION:结果表明25OHD3、25OHD2的变异系数CV分别在x1%-x2%,y1%-y2%范围内，均小于20%，满足检测要求。
    '''

    # 定义需要生成的字典
    JMD_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        PJJMD_data = JMD.objects.filter(reportinfo_id=id, namejmd="中间精密度")

        JMD_norm = []  # 待测物质列表
        for i in PJJMD_data:
            if i.norm not in JMD_norm:
                JMD_norm.append(i.norm)

        JMD_CV = {}  # CV字典，方便提取CV范围到JMD_CONCLUSION中
        lownum = 0  # 低浓度数据量，如果低浓度数据量为0，则在前端模板不显示，也不将数据存入数据库，下同
        mediannum = 0
        highnum = 0

        # 中间精密度进行数据抓取时，与重复性精密度不同的一点：需要重新计算均值，标准差和cv
        for i in JMD_norm:
            middle_list = []
            JMD_CV[i] = []

            # 低中高原始数据的单独数据列表，方便重新计算均值，标准差和cv。这里和重复性精密度不同，因此重复性精密度无需定义这三个列表
            lowdata = []
            mediandata = []
            highdata = []

            middle_table = JMD.objects.filter(reportinfo_id=id, namejmd="中间精密度", norm=i)
            for j in middle_table:
                rowlist = []  # 每一行的小列表
                rowlist.append(j.Experimentnum)

                # 没有为每个化合物单独设置有效位数，则调用通用性设置
                if Digitsdict == {} or list(Digitsdict.values())[0] == None:
                    # 判断低中高三个浓度是否都有数据，同文件读取步骤
                    if j.low == "" and j.median != '' and j.high != '':
                        lownum += 1
                        rowlist.append(j.median)
                        rowlist.append(j.high)

                        # 单独数据列表添加数据，因为后续要执行计算，因此需要转换为浮点数
                        mediandata.append(float(j.median))
                        highdata.append(float(j.high))
                    elif j.low != "" and j.median == '' and j.high != '':
                        mediannum += 1
                        rowlist.append(j.low)
                        rowlist.append(j.high)

                        # 单独数据列表添加数据，因为后续要执行计算，因此需要转换为浮点数
                        lowdata.append(float(j.low))
                        highdata.append(float(j.high))
                    elif j.low != "" and j.median != '' and j.high == '':
                        highnum += 1
                        rowlist.append(j.low)
                        rowlist.append(j.median)

                        # 单独数据列表添加数据，因为后续要执行计算，因此需要转换为浮点数
                        lowdata.append(float(j.low))
                        mediandata.append(float(j.median))
                    else:
                        rowlist.append(j.low)
                        rowlist.append(j.median)
                        rowlist.append(j.high)

                        # 单独数据列表添加数据，因为后续要执行计算，因此需要转换为浮点数
                        lowdata.append(float(j.low))
                        mediandata.append(float(j.median))
                        highdata.append(float(j.high))

                else:  # 为每个化合物单独设置了有效位数，则调用每个化合物的设置
                    # 判断低中高三个浓度是否都有数据，同文件读取步骤
                    if j.low == "" and j.median != '' and j.high != '':
                        lownum += 1
                        rowlist.append(effectnum(j.median, Digitsdict[i]))
                        rowlist.append(effectnum(j.high, Digitsdict[i]))

                        mediandata.append(floateffectnum(j.median, Digitsdict[i]))
                        highdata.append(floateffectnum(j.high, Digitsdict[i]))
                    elif j.low != "" and j.median == '' and j.high != '':
                        mediannum += 1
                        rowlist.append(effectnum(j.low, Digitsdict[i]))
                        rowlist.append(effectnum(j.high, Digitsdict[i]))

                        lowdata.append(floateffectnum(j.low, Digitsdict[i]))
                        highdata.append(floateffectnum(j.high, Digitsdict[i]))
                    elif j.low != "" and j.median != '' and j.high == '':
                        highnum += 1
                        rowlist.append(effectnum(j.low, Digitsdict[i]))
                        rowlist.append(effectnum(j.median, Digitsdict[i]))

                        lowdata.append(floateffectnum(j.low, Digitsdict[i]))
                        mediandata.append(floateffectnum(j.median, Digitsdict[i]))
                    else:
                        rowlist.append(effectnum(j.low, Digitsdict[i]))
                        rowlist.append(effectnum(j.median, Digitsdict[i]))
                        rowlist.append(effectnum(j.high, Digitsdict[i]))

                        lowdata.append(floateffectnum(j.low, Digitsdict[i]))
                        mediandata.append(floateffectnum(j.median, Digitsdict[i]))
                        highdata.append(floateffectnum(j.high, Digitsdict[i]))
                middle_list.append(rowlist)

            # 重新计算均值，标准差和cv

            # 1 三个浓度水平全覆盖
            if lownum == 0 and mediannum == 0 and highnum == 0:
                middle_list.append(['均值', mean(lowdata), mean(mediandata), mean(highdata)])
                middle_list.append(['标准差', sd(lowdata), sd(mediandata), sd(highdata)])
                middle_list.append(['CV(%)', cv(lowdata), cv(mediandata), cv(highdata)])

                JMD_CV[i]=[cv(lowdata), cv(mediandata), cv(highdata)]

            # 2 中高浓度
            elif lownum != 0 and mediannum == 0 and highnum == 0:
                middle_list.append(['均值',mean(mediandata), mean(highdata)])
                middle_list.append(['标准差', sd(mediandata), sd(highdata)])
                middle_list.append(['CV(%)', cv(mediandata), cv(highdata)])

                JMD_CV[i]=[cv(mediandata), cv(highdata)]

            # 3 低高浓度
            elif lownum == 0 and mediannum != 0 and highnum == 0:
                middle_list.append(['均值',mean(lowdata), mean(highdata)])
                middle_list.append(['标准差', sd(lowdata), sd(highdata)])
                middle_list.append(['CV(%)', cv(lowdata), cv(highdata)])

                JMD_CV[i]=[cv(lowdata), cv(highdata)]

            # 低中浓度
            elif lownum == 0 and mediannum == 0 and highnum != 0:
                middle_list.append(['均值', mean(lowdata), mean(mediandata)])
                middle_list.append(['标准差', sd(lowdata), sd(mediandata)])
                middle_list.append(['CV(%)', cv(lowdata), cv(mediandata)])

                JMD_CV[i]=[cv(lowdata), cv(mediandata)]

            JMD_dict[i] = middle_list

        JMD_CONCLUSION = "结果表明" + "、" .join(list(JMD_CV.keys()))+"的变异系数CV分别在 "  # 最终需要的结论

        JMD_CVrange = []
        for value in JMD_CV.values():
            JMD_CVrange.append(str(min(value))+"%"+"-"+str(max(value))+"%")

        string = "," .join(list(JMD_CVrange))+"范围内,均小于20%,满足检测要求。"

        for i in string:
            JMD_CONCLUSION = JMD_CONCLUSION + i

        if len(textlist_special) != 0:
            return {"JMD_dict": JMD_dict, "textlist": textlist_special, "serial": len(textlist_special)+1,
                    "JMD_CONCLUSION": JMD_CONCLUSION, "lownum": lownum, "mediannum": mediannum, "highnum": highnum}
        else:
            return {"JMD_dict": JMD_dict, "textlist": textlist_general, "serial": len(textlist_general)+1,
                    "JMD_CONCLUSION": JMD_CONCLUSION, "lownum": lownum, "mediannum": mediannum, "highnum": highnum}

    except:
        pass

# 精密度最终结论数据提取，关联到最终报告预览界面
def related_jmdendconclusion(id):
    '''
    注释:需要生成一个字典JMD_CONCLUSION_table和结论JMD_CONCLUSION,数据格式如下：
    JMD_CONCLUSION_table:{ 化合物1:{ 重复性精密度:[L测定次数,M测定次数,H测定次数,L均值,M均值,H均值,LCV,MCV,HCV]},
    化合物2:{ 重复性精密度:[L测定次数,M测定次数,H测定次数,L均值,M均值,H均值,LCV,MCV,HCV]} }
    JMD_CONCLUSION:结果表明25OHD3、25OHD2的变异系数CV分别在x1%-x2%,y1%-y2%范围内，均小于20%，满足检测要求。
    '''

    # 定义需要生成的字典
    JMD_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        JMD_data = JMD.objects.filter(reportinfo_id=id)

        JMD_norm = []  # 待测物质列表
        for i in JMD_data:
            if i.norm not in JMD_norm:
                JMD_norm.append(i.norm)

        for i in JMD_norm:
            middle_list = []  # 每个化合物的数据列表
            
            # 重复性精密度
            middle_table_PN = JMD.objects.filter(reportinfo_id=id, norm=i, namejmd="重复性精密度")
            if middle_table_PN:

                # 低中高浓度测定次数
                L_times = 0  
                M_times = 0
                H_times = 0
                for j in middle_table_PN:
                    if j.Experimentnum != "均值" and j.Experimentnum != "标准差" and j.Experimentnum != "CV(%)" and j.low != "":
                        L_times += 1
                    if j.Experimentnum != "均值" and j.Experimentnum != "标准差" and j.Experimentnum != "CV(%)" and j.median != "":
                        M_times += 1
                    if j.Experimentnum != "均值" and j.Experimentnum != "标准差" and j.Experimentnum != "CV(%)" and j.high != "":
                        H_times += 1
                middle_list.append(L_times)
                middle_list.append(M_times)
                middle_list.append(H_times)

                for k in middle_table_PN:
                    if k.Experimentnum == "均值":
                        middle_list.append(k.low)
                        middle_list.append(k.median)
                        middle_list.append(k.high)
                    elif k.Experimentnum == "CV(%)":
                        middle_list.append(k.low)
                        middle_list.append(k.median)
                        middle_list.append(k.high)

            # 中间精密度
            middle_table_PJ = JMD.objects.filter(reportinfo_id=id, norm=i, namejmd="中间精密度")
            if middle_table_PJ:
                L_times = 0
                M_times = 0
                H_times = 0

                # 中间精密度需重新计算均值，标准差和cv
                lowdata = []
                mediandata = []
                highdata = []
                for j in middle_table_PJ:
                    if j.Experimentnum != "均值" and j.Experimentnum != "标准差" and j.Experimentnum != "CV(%)" and j.low != "":
                        L_times += 1
                        lowdata.append(float(j.low))
                    if j.Experimentnum != "均值" and j.Experimentnum != "标准差" and j.Experimentnum != "CV(%)" and j.median != "":
                        M_times += 1
                        mediandata.append(float(j.median))
                    if j.Experimentnum != "均值" and j.Experimentnum != "标准差" and j.Experimentnum != "CV(%)" and j.high != "":
                        H_times += 1
                        highdata.append(float(j.high))
                middle_list.append(L_times)
                middle_list.append(M_times)
                middle_list.append(H_times)

                if lowdata != []:
                    lowmean = mean(lowdata)
                    lowsd = sd(lowdata)
                    lowcv = cv(lowdata)
                else:
                    lowmean = ''
                    lowcv = ''

                if mediandata != []:
                    medianmean = mean(mediandata)
                    mediansd = sd(mediandata)
                    mediancv = cv(mediandata)
                else:
                    medianmean = ''
                    mediancv = ''

                if highdata != []:
                    highmean = mean(highdata)
                    highsd = sd(highdata)
                    highcv = cv(highdata)
                else:
                    highmean = ''
                    highcv = ''

                middle_list.append(lowmean)
                middle_list.append(medianmean)
                middle_list.append(highmean)

                middle_list.append(lowcv)
                middle_list.append(mediancv)
                middle_list.append(highcv)

            JMD_dict[i] = middle_list

        JMD_CONCLUSION = "、" .join(list(JMD_dict.keys())) + "的重复性精密度与中间精密度的结果分析，CV均小于20%，满足检测要求。"

    except:
        pass
    return {"JMD_dict": JMD_dict, "JMD_CONCLUSION": JMD_CONCLUSION}
