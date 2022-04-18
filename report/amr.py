from django.http import HttpResponse,HttpResponseRedirect
from django.shortcuts import render,redirect
from report import models
from report.models import *
import xlrd
import numpy as np
import math
from report.effectnum import *
import csv
from docx import Document
from datetime import datetime
import re

# LOQ数据抓取
# 1 上传单个数据文件
def LOQfileread(files,reportinfo,project,platform,manufacturers,Unit,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB,Number_of_compounds):

    # 第一步:后台数据抓取（回收率上下限，最大允许CV）
    id1 = Special.objects.get(project=project).id  
    id2 = AMRspecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if AMRspecialmethod.objects.filter(aMRspecial=id2): 
        lowvalue=AMRspecialmethod.objects.get(aMRspecial=id2).lowvalue #回收率下限
        upvalue=AMRspecialmethod.objects.get(aMRspecial=id2).upvalue #回收率上限
        maxCV=AMRspecialmethod.objects.get(aMRspecial=id2).cv #最大允许CV
        
    else:
        general_1 = General.objects.get(name="通用性项目")
        amr_general = AMRgeneral.objects.get(general=general_1)
        lowvalue=AMRgeneralmethod.objects.get(aMRgeneral=amr_general).lowvalue #回收率下限
        upvalue=AMRgeneralmethod.objects.get(aMRgeneral=amr_general).upvalue #回收率上限
        maxCV=AMRgeneralmethod.objects.get(aMRgeneral=amr_general).cv #最大允许CV

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示
    if manufacturers=="AB":
        if len(normAB)!=Number_of_compounds:
            error_message="未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"  
            return {"error_message":error_message}

    #  第二步:开始文件读取
    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典AMR_dict,数据格式如下：
    print(AMR_dict):
    {"化合物1":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,]},
    {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,]},
    "化合物2":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,]},
    {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,]}
    '''

    #  创新第二步需要生成的结果容器
    AMR_dict={} # 需要生成的字典
    Accuracyjudge=[] #每个化合物超过回收率范围的个数列表
    CVjudge=[] #每个化合物超过CV范围的个数列表
    S=["S1","S2","S3","S4","S5","S6","S7","S8","S9","S10","S11","S12","S13","S14","S15"] # 预定义浓度序号列表

    picturelist=[] # 图片文件列表
    picturenum=0 # 上传文件中的图片个数

    # 各仪器平台及各仪器厂家数据读取
    for file in files:
        print(file.name)
        # 文件为图片(.png或.JPG)
        if '.png' in file.name or ".JPG" in file.name or ".PNG" in file.name:   
            picturenum+=1
            picturelist.append(file)
            if AMRpicture.objects.filter(reportinfo = reportinfo,img = "img/"+file.name):
                pass
            else:
                AMRpicture.objects.create(reportinfo = reportinfo,img = file,name="")

            AMRpicture_table = AMRpicture.objects.filter(reportinfo = reportinfo)
            id=[]
            for item in AMRpicture_table:
                id.append(item.reportinfo_id)

        else:
            # 一 液质平台(理论浓度从原始文件中读取)
            if platform=="液质":
                if manufacturers =="Agilent":                             
                    # 1 读取csv文件（Agilent）
                    file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                    file_data = file.read().decode('utf-8')
                    lines = file_data.split('\r\n')
                    for i in range(len(lines)): 
                        if len(lines[i])!=0:
                            lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                        else:
                            lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                            del lines[i] #最后一行如为空行，则删除该元素

                    # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                    norm=[] #化合物列表
                    for j in range(len(lines[0])):  #从第一行开始
                        if "-Q Results" in lines[0][j]:
                            # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除
                            if lines[0][j].split("-Q")[0][0]!='"':  
                                norm.append(lines[0][j].split("-Q")[0])
                            else:
                                norm.append(lines[0][j].split("-Q")[0][1:])
                
                    # 从第二行确定实验号（Sample Name）,理论浓度（Exp. Conc.）,实际浓度（Calc. Conc.）和回收率（Accuracy）的索引
                    nameindex=0  #实验号索引
                    theoryconindex=[]   #理论浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里，下同）
                    concindex=[]  #实际浓度索引列表
                    accuracyindex=[]    #回收率索引列表
                    for j in range(len(lines[1])):  #从第二行开始  

                        # 实验号
                        if lines[1][j] == "Sample Name" :
                            nameindex=j
                        
                        # 理论浓度
                        elif lines[1][j]  == "Exp. Conc." :
                            theoryconindex.append(j)

                        # 实际浓度
                        elif lines[1][j]  == "Final Conc." :
                            concindex.append(j)

                        # 回收率
                        elif lines[1][j]  == "Accuracy" :
                            accuracyindex.append(j)
                
                    # 确认原始数据中与AMR相关(实验号前含有"AMR-")的sample name名，即曲线点个数，放进一个列表(目前只能用于个各化合物曲线点个数一致的情况)

                    # AMR实验号命名格式:"AMR-S1-B1",lines[j][nameindex][0:6]即为"AMR-S1"
                    AMR_STD=[] 
                    for j in range(len(lines)):
                        if "AMR-" in lines[j][nameindex] and lines[j][nameindex][0:6] not in AMR_STD:
                            AMR_STD.append(lines[j][nameindex][0:6])

                    # 按顺序重新排列AMR_STD
                    AMR_STD_sort=[]
                    for i in S:
                        for j in AMR_STD:
                            if i in j:
                                AMR_STD_sort.append(j)

                    # 从原始数据表格中抓取数据
                    for k in range(len(norm)): # 循环化合物列表
                        normdict={} #每个化合物数据字典
                        for j in range(len(AMR_STD)): 
                            calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                            theoryconc=[] # 每个化合物内各曲线点的理论值列表,会有重复                    
                            Accuracy=[] # 回收率列表
                                        
                            for i in range(len(lines)): # 循环原始数据中的每一行
                                if lines[i][nameindex][0:6] == AMR_STD[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                    calconc.append(effectnum(lines[i][concindex[k]],digits)) #添加实际浓度              
                                    theoryconc.append(effectnum(lines[i][theoryconindex[k]],digits)) # 添加理论浓度
                                    Accuracy.append(new_round(lines[i][accuracyindex[k]],2)) #添加回收率

                            normdict[AMR_STD[j]]=[]
                            normdict[AMR_STD[j]].append(theoryconc[0]) #理论浓度列表有重复，只添加第一个值
                            for i in range(len(calconc)):
                                normdict[AMR_STD[j]].append(calconc[i])
                                normdict[AMR_STD[j]].append(Accuracy[i])

                        AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物        
                    print(AMR_dict)
                    
                elif manufacturers =="Waters": 
                    # 内标标识
                    ISlist=["D3","D4","D5","D6","D7","D8"]

                    # 若是最新的 2.0.1 版本的xlrd包，只支持 .xls 文件，读取.xlsx文件会报错。若要正常读取，需安装旧版本的xlrd：pip3 install xlrd==1.2.0     
                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    file_data = data.sheets()[0]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    norm = []  # 化合物列表
                    Compound_row =[] # 含有“Compound”关键词的所在行(包含内标)
                    norm_row = []  # 实际化合物所在行(不包含内标)
                    for i in range(nrows):
                        if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):
                            Compound_row.append(i)  

                        # 判断是否含有内标标识
                        if all(j not in str(file_data.row_values(i)[0]) for j in ISlist):
                            if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):  # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                                norm.append(file_data.row_values(i)[0].split(":")[1].strip()) # strip()的作用是去除前后空格
                                norm_row.append(i) 

                    # 第一种情况，不含有内标
                    if len(Compound_row) == len(norm_row):
                        pass

                    # 第二种情况，含有内标
                    else:
                        nrows = Compound_row[len(norm_row)]                   

                    nameindex=0  #实验号索引
                    conindex=0  #实际浓度索引
                    theoryconindex=0   #理论浓度索引
                    accuracyindex=0   #回收率索引
                    for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                        if file_data.row_values(norm_row[0]+2)[i]=="ID":
                            nameindex=i
                        elif "nmol/L" in file_data.row_values(norm_row[0]+2)[i]:
                            conindex=i
                        elif "Std. Conc" in file_data.row_values(norm_row[0]+2)[i]:
                            theoryconindex=i
                        elif "%Rec" in file_data.row_values(norm_row[0]+2)[i]:
                            accuracyindex=i
                    
                    # 未准确设置表头列名,直接返回并提示!
                    if theoryconindex==0 or conindex==0 or accuracyindex==0:
                        error_message="未准确设置表头列名!"
                        return {"error_message":error_message}

                    for k in range(len(norm)):
                        AMR_STD=[]
                        if k<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[k],norm_row[k+1]): 
                                if "AMR-" in file_data.row_values(i)[nameindex] and file_data.row_values(i)[nameindex][0:6] not in AMR_STD:
                                    AMR_STD.append(file_data.row_values(i)[nameindex][0:6])

                            # 按顺序重新排列AMR_STD
                            AMR_STD_sort=[]
                            for i in S:
                                for j in AMR_STD:
                                    if i in j:
                                        AMR_STD_sort.append(j)

                            normdict={} #每个化合物数据字典
                            for j in range(len(AMR_STD_sort)): 
                                calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                                theoryconc=[]                 
                                Accuracy=[] 
                                            
                                for i in range(norm_row[k],norm_row[k+1]): 
                                    if file_data.row_values(i)[nameindex][0:6] == AMR_STD_sort[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                        calconc.append(effectnum(file_data.row_values(i)[conindex],digits)) #添加实际浓度              
                                        theoryconc.append(effectnum(file_data.row_values(i)[theoryconindex],digits)) # 添加理论浓度
                                        Accuracy.append(new_round(file_data.row_values(i)[accuracyindex],2)) #添加回收率

                                normdict[AMR_STD_sort[j]]=[]
                                normdict[AMR_STD_sort[j]].append(theoryconc[0]) #理论浓度列表有重复，只添加第一个值
                                for i in range(len(calconc)):
                                    normdict[AMR_STD_sort[j]].append(calconc[i])
                                    normdict[AMR_STD_sort[j]].append(Accuracy[i])

                            AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物        
                        
                        else:
                            for i in range(norm_row[k],nrows):  
                                if "AMR-" in file_data.row_values(i)[nameindex] and file_data.row_values(i)[nameindex][0:6] not in AMR_STD:
                                    AMR_STD.append(file_data.row_values(i)[nameindex][0:6])

                            # 按顺序重新排列AMR_STD
                            AMR_STD_sort=[]
                            for i in S:
                                for j in AMR_STD:
                                    if i in j:
                                        AMR_STD_sort.append(j)

                            normdict={} #每个化合物数据字典
                            for j in range(len(AMR_STD_sort)): 
                                calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                                theoryconc=[]                 
                                Accuracy=[] 
                                            
                                for i in range(norm_row[k],nrows): 
                                    if file_data.row_values(i)[nameindex][0:6] == AMR_STD_sort[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                        calconc.append(effectnum(file_data.row_values(i)[conindex],digits)) #添加实际浓度              
                                        theoryconc.append(effectnum(file_data.row_values(i)[theoryconindex],digits)) # 添加理论浓度
                                        Accuracy.append(new_round(file_data.row_values(i)[accuracyindex],2)) #添加回收率

                                normdict[AMR_STD_sort[j]]=[]
                                normdict[AMR_STD_sort[j]].append(theoryconc[0]) #理论浓度列表有重复，只添加第一个值
                                for i in range(len(calconc)):
                                    normdict[AMR_STD_sort[j]].append(calconc[i])
                                    normdict[AMR_STD_sort[j]].append(Accuracy[i])

                            AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物

                elif manufacturers =="Thermo":
                    if '.png' not in file.name and ".JPG" not in file.name: 
                        Thermo = Special.objects.get(project=project) 
                        pt_special = PTspecial.objects.get(special=Thermo)
                        pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
                        PTnorm=[] # 待测物质列表
                        for i in pt_accept:
                            PTnorm.append(i.norm)

                        data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                        norm=[] #Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                        sheetindex=[] #需要的化合物所在sheet索引列表
                        for index in range(len(data.sheet_names())):
                            if data.sheet_names()[index] in PTnorm:
                                norm.append(data.sheet_names()[index])
                                sheetindex.append(index)

                        # 循环读取每个sheet工作表,即为每个化合物的表
                        for index in range(len(sheetindex)):
                            file_data = data.sheets()[sheetindex[index]]
                            nrows=file_data.nrows
                            ncols=file_data.ncols

                            nameindex=0  #实验号索引
                            conindex=0  #实际浓度索引
                            theoryconindex=0   #理论浓度索引
                            accuracyindex=0   #回收率索引
                            for i in range(len(file_data.row_values(0))):  
                                if file_data.row_values(0)[i]=="Compound":
                                    nameindex=i
                                elif file_data.row_values(0)[i]=="Calculated Amt":
                                    conindex=i
                                elif file_data.row_values(0)[i]=="Theoretical Amt":
                                    theoryconindex=i
                                elif file_data.row_values(0)[i]=="回收率换算":
                                    accuracyindex=i

                            AMR_STD=[]
                            AMR_STD_distict=[] 
                            for i in range(nrows): 
                                if "AMR-" in file_data.row_values(i)[nameindex]:
                                    AMR_STD.append(file_data.row_values(i)[nameindex])
                                            
                            for i in AMR_STD:
                                if i not in AMR_STD_distict: # AMR_STD列表去重
                                    AMR_STD_distict.append(i)

                            normdict={} #每个化合物数据字典
                            for j in range(len(AMR_STD_distict)): 
                                calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                                theoryconc=[]                 
                                Accuracy=[] 
                                            
                                for i in range(nrows): 
                                    if file_data.row_values(i)[nameindex] == AMR_STD_distict[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                        calconc.append(effectnum(file_data.row_values(i)[conindex],digits)) #添加实际浓度              
                                        theoryconc.append(effectnum(file_data.row_values(i)[theoryconindex],digits)) # 添加理论浓度

                                        accuracyconvert = float(file_data.row_values(i)[accuracyindex])+100 # Thermo数据需对回收率进行换算
                                        Accuracy.append(new_round(accuracyconvert,1)) #添加回收率

                                # # 第一个化合物的第一个曲线点列表calconc循环完成，放入normdict中，开始循环该化合物的下一个曲线点
                                normdict[AMR_STD_distict[j]]=[]
                                normdict[AMR_STD_distict[j]].append(theoryconc[0]) #理论浓度列表有重复，只添加第一个值
                                for i in calconc:
                                    normdict[AMR_STD_distict[j]].append(i)
                                for i in Accuracy:
                                    normdict[AMR_STD_distict[j]].append(i)

                            AMR_dict[norm[index]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物

                elif manufacturers =="岛津":
                    if '.png' not in file.name and ".JPG" not in file.name: 
                        content= []
                        for line in file:
                            content.append(line.decode("GB2312").replace("\r\n", "").split("\t"))

                        nameindex=0
                        norm=[] #化合物列表
                        norm_row=[] #化合物所在行
                        theoryconindex=0  #理论浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列，下同
                        concindex=0 #实际浓度索引
                        accuracyindex=0    #回收率索引
                        
                        for i in range(len(content[2])):  #第二行确定samplename和浓度所在列
                            if content[2][i]=="数据文件名":
                                nameindex=i 
                            elif content[2][i]=="浓度":
                                concindex=i 
                            elif content[2][i]=="理论浓度":
                                theoryconindex=i 
                            elif content[2][i]=="回收率":
                                accuracyindex=i 

                        for i in range(len(content)): 
                            if content[i][0]=="Name": #如果某一行第一列为"Name"，则该行第二列为化合物名称
                                norm.append(content[i][1])
                                norm_row.append(i)

                        # 确定曲线点数，实验号前含有"AMR-"(以第一个化合物为准确定曲线点数)
                        AMR_STD=[] 
                        if len(norm)==1: #如果只有一个化合物      
                            for i in range(norm_row[0],len(content)):                    
                                if "AMR-" in content[i][nameindex]: 
                                    AMR_STD.append(content[i][nameindex])
                        else:
                            for i in range(norm_row[0],norm_row[1]):                    
                                if "AMR-" in content[i][nameindex]: 
                                    AMR_STD.append(content[i][nameindex])
                    
                        AMR_STD_distict=[] 
                        for i in AMR_STD:
                            if i not in AMR_STD_distict: # AMR_STD列表去重
                                AMR_STD_distict.append(i)
                
                        # print(AMR_STD_distict) : ['AMR-STD-1', 'AMR-STD-2', 'AMR-STD-3', 'AMR-STD-4', 'AMR-STD-5',...]

                        for k in range(len(norm)): # 循环化合物列表
                            normdict={} # 每个化合物数据字典
                            for j in range(len(AMR_STD_distict)): 
                                calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                                theoryconc=[] # 每个化合物内各曲线点的理论值列表,会有重复                    
                                Accuracy=[] # 回收率列表

                                if k<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行              
                                    for i in range(norm_row[k],norm_row[k+1]):
                                        if content[i][nameindex] == AMR_STD_distict[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                            calconc.append(effectnum(content[i][concindex],digits)) #添加检测值                        
                                            theoryconc.append(effectnum(content[i][theoryconindex],digits)) # 添加理论值
                                            Accuracy.append(new_round(content[i][accuracyindex],1)) #添加回收率
                                else:
                                    for i in range(norm_row[k],len(content)): 
                                        if content[i][nameindex] == AMR_STD_distict[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据 
                                            calconc.append(effectnum(content[i][concindex],digits)) #添加检测值                        
                                            theoryconc.append(effectnum(content[i][theoryconindex],digits)) # 添加理论值
                                            Accuracy.append(new_round(content[i][accuracyindex],1)) #添加回收率

                                normdict[AMR_STD_distict[j]]=[]
                                normdict[AMR_STD_distict[j]].append(theoryconc[0]) #理论浓度列表有重复，只添加第一个值
                                for i in calconc:
                                    normdict[AMR_STD_distict[j]].append(i)
                                for i in Accuracy:
                                    normdict[AMR_STD_distict[j]].append(i)

                            AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物        

                elif manufacturers =="AB":           
                    # 定义化合物列表，列表统一命名为norm
                    norm = normAB

                    # 获取上传的文件
                    file_data = Document(file)

                    # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                    paragraphs = [] 

                    # 若标题长度为0或为换行等，不添加进入标题内容列表
                    for p in file_data.paragraphs:
                        if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                            paragraphs.append(p.text.strip())

                    # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                    tableindex=[]
                    for i in range(len(paragraphs)):
                        for j in range(len(ZP_Method_precursor_ion)):
                            if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                tableindex.append(2*i+1)

                    tables = file_data.tables #获取文件中的表格集
                    
                    # 循环定量表格的索引
                    for k in range(len(tableindex)):
                        # 获取文件中的定量表格
                        tablequantify = tables[tableindex[k]] 

                        # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                        cells=tablequantify._cells
                        ROWS=len(tablequantify.rows)
                        COLUMNS=len(tablequantify.columns)
                        rowdatalist=[] #每一行的数据
                        rowdatagatherlist=[] #大列表，包含每一行的数据

                        for i in range(ROWS*COLUMNS):
                            text=cells[i].text.replace("\n","")
                            text=text.strip() #去除空白符
                            if i % COLUMNS != 0 or i == 0:
                                rowdatalist.append(text)
                            else:
                                rowdatagatherlist.append(rowdatalist)
                                rowdatalist=[]
                                rowdatalist.append(text)
                        rowdatagatherlist.append(rowdatalist)

                        nameindex=0  #实验号索引
                        theoryconindex=0   #理论浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里，下同）
                        concindex=0  #实际浓度索引列表
                        accuracyindex=0    #回收率索引列表

                        # 读取表格的第一行的单元格,判断实验号和浓度索引
                        for i in range(len(rowdatagatherlist[0])):
                            if rowdatagatherlist[0][i] == "Sample Name" :
                                nameindex=i
                            elif "Target" in rowdatagatherlist[0][i]:
                                theoryconindex=i
                            elif "Calculated Conc" in rowdatagatherlist[0][i]:
                                concindex=i
                            elif "Accuracy" in rowdatagatherlist[0][i]:
                                accuracyindex=i
            
                        # 未准确设置表头列名,直接返回并提示!
                        if theoryconindex==0 or concindex==0 or accuracyindex==0:
                            error_message="未准确设置表头列名!"
                            return {"error_message":error_message}

                        # 确认原始数据中与AMR相关(实验号前含有"AMR-")的sample name名，放进一个列表
                        AMR_STD=[] 
                        for i in range(len(rowdatagatherlist)): 
                            if "AMR-" in rowdatagatherlist[i][nameindex] and rowdatagatherlist[i][nameindex][0:6] not in AMR_STD:
                                AMR_STD.append(rowdatagatherlist[i][nameindex][0:6])
                        
                        # 按顺序重新排列AMR_STD
                        AMR_STD_sort=[]
                        for i in S:
                            for j in AMR_STD:
                                if i in j:
                                    AMR_STD_sort.append(j)

                        # 从原始数据表格中抓取数据(耗时较久,需优化)
                        normdict={} # 每个化合物数据字典
                        for j in range(len(AMR_STD_sort)): 
                            calconc=[]  # 实际浓度列表
                            theoryconc=[] # 理论浓度列表
                            Accuracy=[] # 回收率列表

                            for i in range(len(rowdatagatherlist)): # 循环原始数据中的每一行       
                                if rowdatagatherlist[i][nameindex][0:6] == AMR_STD_sort[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                    calconc.append(effectnum(rowdatagatherlist[i][concindex],digits)) #添加检测值           
                                    theoryconc.append(effectnum(rowdatagatherlist[i][theoryconindex],digits)) # 添加理论值
                                    Accuracy.append(new_round(rowdatagatherlist[i][accuracyindex],2)) #添加回收率

                            normdict[AMR_STD_sort[j]]=[]
                            normdict[AMR_STD_sort[j]].append(theoryconc[0]) #理论浓度列表有重复，只添加第一个值
                            for i in range(len(calconc)):
                                normdict[AMR_STD_sort[j]].append(calconc[i])
                                normdict[AMR_STD_sort[j]].append(Accuracy[i])

                        AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物

            # 二 非液质平台(理论浓度由用户自行输入)
            elif platform=="液相":
                if manufacturers =="Agilent":   

                    # .xlsx格式        
                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    file_data = data.sheets()[0]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    norm=[] #化合物列表
                    norm_row=[] #化合物所在行
                    for j in range(nrows):
                        if file_data.row_values(j)[0] == "化合物:" :  #如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                            norm.append(file_data.row_values(j)[2])
                            norm_row.append(j)

                    nameindex=0  #实验号索引
                    conindex=0  #实际浓度索引

                    # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+1是该化合物表格位于该化合物所在行的下一行
                    for i in range(len(file_data.row_values(norm_row[0]+1))):
                        if "样品名称" in file_data.row_values(norm_row[0]+1)[i]:
                            nameindex = i
                        elif "含量" in file_data.row_values(norm_row[0]+1)[i]:
                            conindex = i

                    # 未准确设置表头列名,直接返回并提示!
                    if conindex==0:
                        error_message="未准确设置表头列名!"
                        return {"error_message":error_message}

                    for k in range(len(norm)):
                        AMR_STD=[]
                        if k<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                            for i in range(norm_row[k],norm_row[k+1]): 
                                if "AMR-" in file_data.row_values(i)[nameindex] and file_data.row_values(i)[nameindex][0:6] not in AMR_STD:
                                    AMR_STD.append(file_data.row_values(i)[nameindex][0:6])

                            # 按顺序重新排列AMR_STD(以防原始数据曲线点不按照顺序排列的情况)
                            AMR_STD_sort=[]
                            for i in S:
                                for j in AMR_STD:
                                    if i in j:
                                        AMR_STD_sort.append(j)

                            normdict={} #每个化合物数据字典
                            for j in range(len(AMR_STD_sort)): 
                                # 液相平台原始数据不包含理论浓度和回收率
                                calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                                            
                                for i in range(norm_row[k],norm_row[k+1]): 
                                    if file_data.row_values(i)[nameindex][0:6] == AMR_STD_sort[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                        calconc.append(effectnum(file_data.row_values(i)[conindex],digits)) #添加实际浓度              

                                normdict[AMR_STD_sort[j]]=[]
                                for i in range(len(calconc)):
                                    normdict[AMR_STD_sort[j]].append(calconc[i])

                            AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物        
                        
                        else:
                            for i in range(norm_row[k],nrows):  
                                if "AMR-" in file_data.row_values(i)[nameindex] and file_data.row_values(i)[nameindex][0:6] not in AMR_STD:
                                    AMR_STD.append(file_data.row_values(i)[nameindex][0:6])

                            # 按顺序重新排列AMR_STD
                            AMR_STD_sort=[]
                            for i in S:
                                for j in AMR_STD:
                                    if i in j:
                                        AMR_STD_sort.append(j)

                            normdict={} #每个化合物数据字典
                            for j in range(len(AMR_STD_sort)): 
                                calconc=[]  # 每个化合物内各曲线点（S1,S2,S3...）的数据列表列表：['S1检测值1','S1检测值2',...] 
                                            
                                for i in range(norm_row[k],nrows): 
                                    if file_data.row_values(i)[nameindex][0:6] == AMR_STD_sort[j]: # 如果实验号命名方式匹配上，则在相应列表中添加相应数据  
                                        calconc.append(effectnum(file_data.row_values(i)[conindex],digits)) #添加实际浓度              

                                normdict[AMR_STD_sort[j]]=[]
                                for i in range(len(calconc)):
                                    normdict[AMR_STD_sort[j]].append(calconc[i])

                            AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物

            print(AMR_dict)

    # 判断是否一共6条曲线
    print(len(AMR_dict[norm[k]][AMR_STD_sort[0]]))

    # 一 液质平台判断
    if platform=="液质":
        if len(AMR_dict[norm[k]][AMR_STD_sort[0]])!=13:
            error_message="LOQ验证曲线不为6条,请规范操作后再进行数据验证!"  
            return {"error_message":error_message}
    else:
        if len(AMR_dict[norm[k]][AMR_STD_sort[0]])!=6:
            error_message="LOQ验证曲线不为6条,请规范操作后再进行数据验证!"  
            return {"error_message":error_message}

    ########文件读取完毕#######
    #  第三步:文件读取完毕后的操作(添加平均回收率和检测值CV)

    # 曲线点个数
    LOQ_num = len(AMR_dict[norm[0]])
    if platform=="液质":
        
        '''
        注释2:最终需要生成一个字典AMR_dict,数据格式如下：
        print(AMR_dict):
        {"化合物1":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,'平均回收率','检测值CV']},
        {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,'平均回收率','检测值CV']},
        "化合物2":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,'平均回收率','检测值CV']},
        {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,'平均回收率','检测值CV']}
        ''' 
        for key,value in AMR_dict.items():
            for r in value.keys():   
                sumcalconc=[] ##检测值列表，方便计算检测值CV(固定六个点)
                sumcalconc.append(float(value[r][1]))
                sumcalconc.append(float(value[r][3]))
                sumcalconc.append(float(value[r][5]))
                sumcalconc.append(float(value[r][7]))
                sumcalconc.append(float(value[r][9]))
                sumcalconc.append(float(value[r][11])) 
                cvcalconc=new_round(np.std(sumcalconc,ddof=1)/np.mean(sumcalconc)*100,2) #检测值CV

                sumrecycle=float(value[r][2])+float(value[r][4])+float(value[r][6])+float(value[r][8])+float(value[r][10])+float(value[r][12])   #回收率总和，方便计算平均回收率
                meanrecycle=new_round(sumrecycle/6,2) #平均回收率

                value[r].append(meanrecycle) #添加平均回收率
                value[r].append(cvcalconc) #添加检测值CV

        print(AMR_dict)        
        
        #  第四步:数据存入数据库

        # judgenum(回收率及检测值CV超过范围的个数,此数字为0才将数据存入数据库)
        judgenum=0 #每个化合物超过CV范围的个数

        # 需要进行判断的数据索引
        indexlist=[2,4,6,8,10,12,13,14]
        for key,value in AMR_dict.items():     
            for i in value.values():
                for index,j in enumerate(indexlist): # 前7个值为回收率和平均回收率,最后一个为检测值cv
                    if index<7:
                        if float(i[j])<lowvalue or float(i[j])>upvalue: # 如回收率不通过，添加不通过提示
                            judgenum+=1
                    else:
                        if float(i[j])>maxCV: 
                            judgenum+=1

        # 判断judgenum是否为0,为0才能将数据存入数据库              
        if judgenum==0:
            insert_list =[]
            for key,value in AMR_dict.items():
                for r,c in value.items():
                    insert_list.append(AMR(reportinfo=reportinfo,Experimentnum=r,norm=key,therory_conc=c[0],test_conc1=c[1],test_conc2=c[3],
                    test_conc3=c[5],test_conc4=c[7],test_conc5=c[9],test_conc6=c[11],recycle1=c[2],recycle2=c[4],recycle3=c[6],recycle4=c[8],
                    recycle5=c[10],recycle6=c[12],meanrecycle=c[13],cvtest_conc=c[14]))

            AMR.objects.bulk_create(insert_list)

        else: 
            insert_list =[]
            for key,value in AMR_dict.items():
                for r,c in value.items():
                    insert_list.append(AMR(reportinfo=reportinfo,Experimentnum=r,norm=key,therory_conc=c[0],test_conc1=c[1],test_conc2=c[3],
                    test_conc3=c[5],test_conc4=c[7],test_conc5=c[9],test_conc6=c[11],recycle1=c[2],recycle2=c[4],recycle3=c[6],recycle4=c[8],
                    recycle5=c[10],recycle6=c[12],meanrecycle=c[13],cvtest_conc=c[14]))

            AMR.objects.bulk_create(insert_list)

        if picturenum==0:
            return {"AMR_dict":AMR_dict,"Unit":Unit,"judgenum":judgenum,"picturenum":picturenum,"lowvalue":lowvalue,
            "upvalue":upvalue,"maxCV":maxCV,"platform":platform,"LOQ_num":LOQ_num}      
        else:  
            return {"AMR_dict":AMR_dict,"AMRpicture_table":AMRpicture_table,"id":id[0],"Unit":Unit,"judgenum":judgenum,
            "picturenum":picturenum,"lowvalue":lowvalue,"upvalue":upvalue,"maxCV":maxCV,"platform":platform,"LOQ_num":LOQ_num}

    elif platform=="液相": 

        if picturenum==0:
            return {"AMR_dict":AMR_dict,"Unit":Unit,"picturenum":picturenum,"lowvalue":lowvalue,"upvalue":upvalue,
            "maxCV":maxCV,"platform":platform,"LOQ_num":LOQ_num}
        else:
            return {"AMR_dict":AMR_dict,"AMRpicture_table":AMRpicture_table,"id":id[0],"Unit":Unit,"picturelist":picturelist,
            "picturenum":picturenum,"lowvalue":lowvalue,"upvalue":upvalue,"maxCV":maxCV,"platform":platform,"LOQ_num":LOQ_num}

# 2 上传多个数据文件
def LOQgeneral_multiplefileread(files,reportinfo,project,platform,manufacturers,Unit,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB,Number_of_compounds):

    # 第一步:后台数据抓取（最小样本数，最大允许CV,化合物个数）
    norm_num = Special.objects.get(project=project).Number_of_compounds
    id1 = Special.objects.get(project=project).id  
    id2 = AMRspecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if AMRspecialmethod.objects.filter(aMRspecial=id2): 
        lowvalue=AMRspecialmethod.objects.get(aMRspecial=id2).lowvalue #回收率下限
        upvalue=AMRspecialmethod.objects.get(aMRspecial=id2).upvalue #回收率上限
        maxCV=AMRspecialmethod.objects.get(aMRspecial=id2).cv #最大允许CV
        
    else:
        general_1 = General.objects.get(name="通用性项目")
        amr_general = AMRgeneral.objects.get(general=general_1)
        lowvalue=AMRgeneralmethod.objects.get(aMRgeneral=amr_general).lowvalue #回收率下限
        upvalue=AMRgeneralmethod.objects.get(aMRgeneral=amr_general).upvalue #回收率上限
        maxCV=AMRgeneralmethod.objects.get(aMRgeneral=amr_general).cv #最大允许CV

    # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
    special = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=special)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

    # 后台管理系统中设置的本项目化合物名称
    PTnorm = []  
    for i in pt_accept:
        PTnorm.append(i.norm)

    # 后台管理系统中设置的本项目每个化合物单位
    Unitlist = []
    for i in pt_accept:
        Unitlist.append(i.unit)

    # 未在后台管理系统里准确设置离子对数值,直接返回并提示
    if len(normAB)!=Number_of_compounds:
        error_message="未在后台管理系统里准确设置离子对数值!"  
        return {"error_message":error_message}

    #  第二步:开始文件读取

    '''
    注释1:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典AMR_dict,数据格式如下：
    print(AMR_dict):
    {"化合物1":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,]},
    {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,]},
    "化合物2":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,]},
    {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,]}
    '''

    #  1 定义最终需要的列表和字典
    AMR_dict={} #最终需要生成的字典
    Accuracyjudge=[] #每个化合物超过回收率范围的个数列表
    CVjudge=[] #每个化合物超过CV范围的个数列表
    norm=[] #指标列表
    S=["S1","S2","S3","S4","S5","S6","S7","S8","S9","S10"] # 预定义浓度序号列表

    picturelist=[] # 图片文件列表
    picturenum=0 # 上传文件中的图片个数

    for k in range(norm_num):  #多个文件上传时，先循环化合物，在循环文件
        normdict={} #每个化合物数据字典
        for index,file in enumerate(files):
            # 文件为图片(.png或.JPG)
            if '.png' in file.name or ".JPG" in file.name:    
                picturenum+=1
                picturelist.append(file)
                if AMRpicture.objects.filter(reportinfo = reportinfo,img = "img/"+file.name):
                    pass
                else:
                    AMRpicture.objects.create(reportinfo = reportinfo,img = file,name="")

                AMRpicture_table = AMRpicture.objects.filter(reportinfo = reportinfo)
                id=[]
                for item in AMRpicture_table:
                    id.append(item.reportinfo_id)

            else:
                if manufacturers =="AB":

                    # 定义化合物列表，列表统一命名为norm(第一个化合物才需要设置)
                    if index==0:
                        norm = normAB

                    # 获取上传的文件
                    file_data = Document(file)

                    # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                    paragraphs = [] 

                    # 若标题长度为0或为换行等，不添加进入标题内容列表
                    for p in file_data.paragraphs:
                        if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                            paragraphs.append(p.text.strip())

                    # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格(第一个化合物才需要设置)
                    if index==0:
                        tableindex=[]

                        # 1 每个化合物都含有定性和定量表格的计算方式
                        for i in range(len(paragraphs)):
                            for j in range(len(ZP_Method_precursor_ion)):
                                if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                    tableindex.append(2*i+1) 

                    tables = file_data.tables #获取文件中的表格集

                    # 无需再循环定量表格的索引，依据最外层for循环确定循环中的文件中的定量表格
                    # 获取文件中的定量表格
                    tablequantify = tables[tableindex[k]] 

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells=tablequantify._cells
                    ROWS=len(tablequantify.rows)
                    COLUMNS=len(tablequantify.columns)
                    rowdatalist=[] #每一行的数据
                    rowdatagatherlist=[] #大列表，包含每一行的数据

                    for i in range(ROWS*COLUMNS):
                        text=cells[i].text.replace("\n","")
                        text=text.strip() #去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist=[]
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex=0  #实验号索引
                    theoryconindex=0   #理论浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里，下同）
                    concindex=0  #实际浓度索引列表
                    accuracyindex=0    #回收率索引列表

                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name" :
                            nameindex=i
                        elif "Target" in rowdatagatherlist[0][i]:
                            theoryconindex=i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex=i
                        elif "Accuracy" in rowdatagatherlist[0][i]:
                            accuracyindex=i
        
                    # 未准确设置表头列名,直接返回并提示!
                    if theoryconindex==0 or concindex==0 or accuracyindex==0:
                        error_message="未准确设置表头列名!"
                        return {"error_message":error_message}

                    # 确认原始数据中与AMR相关(实验号前含有"AMR-")的sample name名，放进一个列表(第一个化合物才需要设置)
                    if index==0:
                        AMR_STD=[] 
                        for i in range(len(rowdatagatherlist)): 
                            if "AMR-" in rowdatagatherlist[i][nameindex] and rowdatagatherlist[i][nameindex][0:6] not in AMR_STD:
                                AMR_STD.append(rowdatagatherlist[i][nameindex][0:6])
                        
                        # 按顺序重新排列AMR_STD
                        AMR_STD_sort=[]
                        for i in S:
                            for j in AMR_STD:
                                if i in j:
                                    AMR_STD_sort.append(j)

                    for j in range(len(AMR_STD_sort)): 
                        # 第一个化合物才需要设置
                        if index==0:
                            normdict[AMR_STD_sort[j]]=[]

                        for i in range(len(rowdatagatherlist)): # 循环原始数据中的每一行       
                            if rowdatagatherlist[i][nameindex][0:6] == AMR_STD_sort[j]: # 如果实验号命名方式匹配上，则在相应列表中添加数据 
                                # 添加理论浓度(需避免重复添加)
                                if effectnum(rowdatagatherlist[i][theoryconindex],digits) not in normdict[AMR_STD_sort[j]]:
                                    normdict[AMR_STD_sort[j]].append(effectnum(rowdatagatherlist[i][theoryconindex],digits))

                                # 添加检测值
                                normdict[AMR_STD_sort[j]].append(effectnum(rowdatagatherlist[i][concindex],digits))

                                # 添加回收率
                                normdict[AMR_STD_sort[j]].append(new_round(rowdatagatherlist[i][accuracyindex],2))


        AMR_dict[norm[k]]=normdict # 第一个化合物的数据列表normdict循环完成，放入最终字典AMR_dict中，开始循环下一个化合物

    # 判断是否一共6条曲线
    if len(AMR_dict[norm[k]][AMR_STD_sort[0]])!=13:
        error_message="LOQ验证曲线不为6条,请规范操作后再进行数据验证!"  
        return {"error_message":error_message}

    #  第三步:文件读取完毕后的操作(添加平均回收率和检测值CV)

    if platform=="液质":
        '''
        注释2:最终需要生成一个字典AMR_dict,数据格式如下：
        print(AMR_dict):
        {"化合物1":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,'平均回收率','检测值CV']},
        {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,'平均回收率','检测值CV']},
        "化合物2":{'S1':['S1理论浓度','S1检测值1','S1检测值2',...'S1回收率1','S1回收率2',...,'平均回收率','检测值CV']},
        {'S2':['S2理论浓度','S2检测值1','S2检测值2',...'S2回收率1','S2回收率2',...,'平均回收率','检测值CV']}
        ''' 
        for key,value in AMR_dict.items():
            for r in value.keys():   
                sumcalconc=[] ##检测值列表，方便计算检测值CV(固定六个点)
                sumcalconc.append(float(value[r][1]))
                sumcalconc.append(float(value[r][3]))
                sumcalconc.append(float(value[r][5]))
                sumcalconc.append(float(value[r][7]))
                sumcalconc.append(float(value[r][9]))
                sumcalconc.append(float(value[r][11])) 
                cvcalconc=new_round(np.std(sumcalconc,ddof=1)/np.mean(sumcalconc)*100,2) #检测值CV

                sumrecycle=float(value[r][2])+float(value[r][4])+float(value[r][6])+float(value[r][8])+float(value[r][10])+float(value[r][12])   #回收率总和，方便计算平均回收率
                meanrecycle=new_round(sumrecycle/6,2) #平均回收率

                value[r].append(meanrecycle) #添加平均回收率
                value[r].append(cvcalconc) #添加检测值CV

        #  第四步:数据存入数据库

        # judgenum(回收率及检测值CV超过范围的个数,此数字为0才将数据存入数据库)
        judgenum=0 #每个化合物超过CV范围的个数

        # 需要进行判断的数据索引
        indexlist=[2,4,6,8,10,12,13,14]
        for key,value in AMR_dict.items():     
            for i in value.values():
                for index,j in enumerate(indexlist): # 前7个值为回收率和平均回收率,最后一个为检测值cv
                    if index<7:
                        if float(i[j])<lowvalue or float(i[j])>upvalue: # 如回收率不通过，添加不通过提示
                            judgenum+=1
                    else:
                        if float(i[j])>maxCV: 
                            judgenum+=1


        # 判断judgenum是否为0,为0才能将数据存入数据库              
        # if judgenum==0:
        #     insert_list =[]
        #     for key,value in AMR_dict.items():
        #         for r,c in value.items():
        #             insert_list.append(AMR(reportinfo=reportinfo,Experimentnum=r,norm=key,therory_conc=c[0],test_conc1=c[1],test_conc2=c[3],
        #             test_conc3=c[5],test_conc4=c[7],test_conc5=c[9],test_conc6=c[11],recycle1=c[2],recycle2=c[4],recycle3=c[6],recycle4=c[8],
        #             recycle5=c[10],recycle6=c[12],meanrecycle=c[13],cvtest_conc=c[14]))

        #     AMR.objects.bulk_create(insert_list)

        # else: 
        #     insert_list =[]
        #     for key,value in AMR_dict.items():
        #         for r,c in value.items():
        #             insert_list.append(AMR(reportinfo=reportinfo,Experimentnum=r,norm=key,therory_conc=c[0],test_conc1=c[1],test_conc2=c[3],
        #             test_conc3=c[5],test_conc4=c[7],test_conc5=c[9],test_conc6=c[11],recycle1=c[2],recycle2=c[4],recycle3=c[6],recycle4=c[8],
        #             recycle5=c[10],recycle6=c[12],meanrecycle=c[13],cvtest_conc=c[14]))

        #     AMR.objects.bulk_create(insert_list)

        if picturenum==0:
            return {"AMR_dict":AMR_dict,"Unit":Unit,"judgenum":judgenum,"picturenum":picturenum,"lowvalue":lowvalue,
            "upvalue":upvalue,"maxCV":maxCV}      
        else:  
            return {"AMR_dict":AMR_dict,"AMRpicture_table":AMRpicture_table,"id":id[0],"Unit":Unit,"judgenum":judgenum,
            "picturenum":picturenum,"lowvalue":lowvalue,"upvalue":upvalue,"maxCV":maxCV}


## 检出限数据上传读取
def LODfileread(files,reportinfo,project,platform,manufacturers,Unit,digits,ZP_Method_precursor_ion,ZP_Method_product_ion):
    LOD_dict={}

    if platform=="ICP-MS":
        if manufacturers =="Agilent":
            for file in files:
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                # 后台管理系统-各项目参数设置-PT指标设置里找到化合物名称
                zqd = Special.objects.get(project=project) 
                pt_special = PTspecial.objects.get(special=zqd)
                pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
                PTnorm=[] # 待测物质列表

                for i in pt_accept:
                    PTnorm.append(i.norm)

                # 从第一行确定化合物名称
                norm=[]
                for j in range(ncols):
                    for i in PTnorm:             
                        if i in file_data.row_values(0)[j]:
                            norm.append(i) 

                # 从第二行确定实验号（Sample Name）的索引和化合物浓度索引
                nameindex=0  #实验号索引
                conindex=[] #浓度索引
                for j in range(ncols):       
                    if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j
                    elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                        conindex.append(j)

                # 匹配原始数据中与精密度相关(实验号前含有"L-"或"M-"或"H-")的行  
                for j in range(len(conindex)):
                    LOD_dict[norm[j]]=[]
                    for i in range(2,nrows): 
                        if "JCX-" in file_data.row_values(i)[nameindex]:
                            LOD_dict[norm[j]].append(float(effectnum(file_data.row_values(i)[conindex[j]],digits)))

            ### 数据文件读取完毕后的操作          
            for key in LOD_dict.keys():
                mean=new_round(np.mean(LOD_dict[key]),5)
                sd=new_round(np.std(LOD_dict[key],ddof=1),5)
                lod_3sd=new_round(3*sd,5)
                lod_10sd=new_round(10*sd,5)
                LOD_dict[key].append(str(mean))
                LOD_dict[key].append(str(sd))
                LOD_dict[key].append(str(lod_3sd))
                LOD_dict[key].append(str(lod_10sd))

            for key,value in LOD_dict.items():
                for i in range(len(LOD_dict[key])):
                    if isinstance(LOD_dict[key][i], str)!=True:
                        LOD_dict[key][i]=effectnum(LOD_dict[key][i],digits)

            Experimentnum=["1","2","3","4","5","6","7","8","9","10","11","12","13","14","15","16","17","18","19","20","X","SD","LOD(3 SD)","LOD(10 SD)"]

            insert_list =[]
            for key,value in LOD_dict.items():
                for i in range(len(value)):
                    insert_list.append(LOD(reportinfo=reportinfo,Experimentnum=Experimentnum[i],norm=key,result=value[i]))

            LOD.objects.bulk_create(insert_list)

        return {"LOD_dict":LOD_dict,"Unit":Unit}

    else:
        for file in files:
            if '.png' in file.name:
                if LODpicture.objects.filter(reportinfo = reportinfo,img = "img2/"+file.name):
                    pass
                else:
                    LODpicture.objects.create(reportinfo = reportinfo,img = file,name="")

                objs_verify = LODpicture.objects.filter(reportinfo = reportinfo)
                id=[]
                for item in objs_verify:
                    id.append(item.reportinfo_id)

        return {"objs_verify":objs_verify,"id":id[0]}

                    
# AMR数据关联进入最终报告
def related_AMR(id,unit): 
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    #特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = AMRspecial.objects.get(special=special_1)           
        if AMRspecialtexts.objects.filter(aMRspecial=special_2).count()>0:
            text_special = AMRspecialtexts.objects.filter(aMRspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = AMRgeneral.objects.get(general=general_1)
    text_general = AMRgeneraltexts.objects.filter(aMRgeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project) 
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist=[] # 每个化合物有效位数列表
    Digitsdict={} # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist==[] or Digitslist[0]=="": #如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm]=i.digits

    # 第二步：报告数据提取
    

    '''
    此数据关联步骤需要生成一个字典AMR_dict,数据格式如下:
    print(AMR_dict):
    {"化合物1":[['S1','S1理论浓度','S1检测值1','S1回收率1','S1检测值2','S1回收率2'...,'平均回收率','检测值CV'],
    ['S2','S2理论浓度','S2检测值1','S2回收率1','S2检测值2','S2回收率2'...,'平均回收率','检测值CV'],...],
    "化合物2":[['S1','S1理论浓度','S1检测值1','S1回收率1','S1检测值2','S1回收率2'...,'平均回收率','检测值CV'],
    ['S2','S2理论浓度','S2检测值1','S2回收率1','S2检测值2','S2回收率2'...,'平均回收率','检测值CV'],...]]}
    '''

    # 定义需要生成的字典
    AMR_dict = {}  # 最终需要的字典
    middle_dict={} # 每个化合物的线性范围数据字典

    try:
        # 1 基础数据抓取
        AMR_data = AMR.objects.filter(reportinfo_id=id)

        AMR_norm=[] #化合物列表
        for i in AMR_data:
            if i.norm not in AMR_norm:
                AMR_norm.append(i.norm)
                           
        for i in AMR_norm:
            middle_table = AMR.objects.filter(reportinfo_id=id,norm=i)
            AMR_dict[i]=[] 
            middle_list = []  # 每个化合物的数据列表
            for item in middle_table: # 依次按顺序从数据库中抓取数据
                #没有为每个化合物单独设置有效位数，则调用通用性设置
                if Digitsdict=={} or list(Digitsdict.values())[0]==None:     
                    rowlist = []  # 每一行的小列表
                    rowlist.append(item.Experimentnum)
                    rowlist.append(item.therory_conc)
                    rowlist.append(item.test_conc1)
                    rowlist.append(item.recycle1)
                    rowlist.append(item.test_conc2)
                    rowlist.append(item.recycle2)
                    rowlist.append(item.test_conc3)
                    rowlist.append(item.recycle3)
                    rowlist.append(item.test_conc4)
                    rowlist.append(item.recycle4)
                    rowlist.append(item.test_conc5)
                    rowlist.append(item.recycle5)
                    rowlist.append(item.test_conc6)
                    rowlist.append(item.recycle6)
                    rowlist.append(item.meanrecycle)
                    rowlist.append(item.cvtest_conc)
                    AMR_dict[i].append(rowlist)
                    middle_list.append(float(item.therory_conc)) #要计算线性范围最大最小值，得把理论浓度添加进来

                #为每个化合物单独设置了有效位数，则调用每个化合物的设置
                else:
                    rowlist=[]
                    rowlist.append(item.Experimentnum)
                    rowlist.append(effectnum(item.therory_conc,Digitsdict[i]))
                    rowlist.append(effectnum(item.test_conc1,Digitsdict[i]))
                    rowlist.append(item.recycle1)
                    rowlist.append(effectnum(item.test_conc2,Digitsdict[i]))
                    rowlist.append(item.recycle2)
                    rowlist.append(effectnum(item.test_conc3,Digitsdict[i]))
                    rowlist.append(item.recycle3)
                    rowlist.append(effectnum(item.test_conc4,Digitsdict[i]))
                    rowlist.append(item.recycle4)
                    rowlist.append(effectnum(item.test_conc5,Digitsdict[i]))
                    rowlist.append(item.recycle5)
                    rowlist.append(effectnum(item.test_conc6,Digitsdict[i]))
                    rowlist.append(item.recycle6)
                    rowlist.append(item.meanrecycle)
                    rowlist.append(item.cvtest_conc)
                    AMR_dict[i].append(rowlist)
                    middle_list.append(float(item.therory_conc)) #要计算线性范围最大最小值，得把理论浓度添加进来
            middle_dict[i]=middle_list

        # 描述性内容最后一句结论格式：结果数据如表10和图1~2，18-OHF线性范围为{{resultAMR.min}}~{{resultAMR.max}}pmol/L, 定量限为{{resultAMR.min}}pmol/L。
        AMR_textlist_end='' # 预先定义一个空字符串，用来存放AMR描述性内容的最后一句话
        for key,value in middle_dict.items():
            AMR_textlist_end+=key
            AMR_textlist_end+="线性范围为"
            AMR_textlist_end+=str(min(value))+'~'+str(max(value))+unit+'，'
            AMR_textlist_end+='定量限为'+str(min(value))+unit+'；'

        AMR_textlist_end=AMR_textlist_end[:-1] #去除最后一个分号
        AMR_textlist_end=AMR_textlist_end+"。"
        
        objs = AMRpicture.objects.filter(reportinfo_id=id) #图片数据表

        if len(textlist_special)!=0:
            return {"AMR_dict":AMR_dict,"AMR_textlist_end":AMR_textlist_end,"textlist":textlist_special,"serial":len(textlist_special)+1,"id":id,"objs":objs}

        else:
            return {"AMR_dict":AMR_dict,"AMR_textlist_end":AMR_textlist_end,"textlist":textlist_general,"serial":len(textlist_general)+1,"id":id,"objs":objs}

    except:
        pass


# 方法检出限数据关联进入最终报告
def related_LOD(id): 
    # 第一步：后台描述性内容数据提取

    # 根据id找到项目
    project=ReportInfo.objects.get(id=id).project

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    # 特殊数据抓取
    JCX_special = Special.objects.get(project=project)   
    jcx_special = JCXspecial.objects.get(special=JCX_special) 
    textlist_special = [] #特殊参数设置描述性内容
    if JCXspecialtexts.objects.filter(jCXspecial=jcx_special).count()>0: 
        text_special = JCXspecialtexts.objects.filter(jCXspecial=jcx_special)   
        for i in text_special:
            textlist_special.append(i.text)
    
    # 通用数据抓取
    JCX_general = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    jcx_general = JCXgeneral.objects.get(general=JCX_general)
    text_general = JCXgeneraltexts.objects.filter(jCXgeneral=jcx_general)   
    textlist_general = [] 
    for i in text_general:
        textlist_general.append(i.text) #AMR通用参数设置描述性内容添加完毕

    # 第二步：报告数据提取
    dataLOD = LOD.objects.filter(reportinfo_id=id)
    dataLOD_picture = LODpicture.objects.filter(reportinfo_id=id)

    if dataLOD:
        LOD_endreport_dict={}

        LOD_endreport_norm=[] #待测物质列表

        for i in dataLOD:
            if i not in LOD_endreport_norm:
                LOD_endreport_norm.append(i.norm)

        for i in LOD_endreport_norm:
            dataLOD_group = LOD.objects.filter(reportinfo_id=id,norm=i)
            group=[]
            for j in dataLOD_group:    
                group.append(j.result)
            LOD_endreport_dict[i]=group
        
        print(LOD_endreport_dict)
        
        if len(textlist_special)!=0:
            return {"LOD_endreport_dict":LOD_endreport_dict,"textlist":textlist_special,"serial":len(textlist_special)+1}
        else:
            return {"LOD_endreport_dict":LOD_endreport_dict,"textlist":textlist_general,"serial":len(textlist_general)+1}

    if dataLOD_picture:
        LOD_conclusion=[] 
        for i in dataLOD_picture:
            LOD_conclusion.append(i.conclusion)

        if len(textlist_special)!=0:
            return {"objs":dataLOD_picture,"LOD_conclusion":LOD_conclusion[0],"textlist":textlist_special,"serial":len(textlist_special)+1}
        else:
            return {"objs":dataLOD_picture,"LOD_conclusion":LOD_conclusion[0],"textlist":textlist_general,"serial":len(textlist_general)+1}

# AMR最终结论表格进入最终报告
def related_AMRconclusion(id): 
    # 定义需要生成的字典
    AMRconclusion_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        AMRconclusion_data = AMRconsluion.objects.filter(reportinfo_id=id)

        for i in AMRconclusion_data:
            AMRconclusion_dict[i.name]=[]
            AMRconclusion_dict[i.name].append(i.lodconclusion)
            AMRconclusion_dict[i.name].append(i.loqconclusion)
            AMRconclusion_dict[i.name].append(i.amrconclusion)
    
        return {"AMRconclusion_dict":AMRconclusion_dict}

    except:
        pass