import xlrd
from report.models import *
from docx import Document
from report.effectnum import *
import numpy as np
import re


# 样品储存稳定性
def store_fileread(files, Detectionplatform, reportinfo, project, platform,manufacturers, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB):
      pass

# 样品处理后稳定性
def handle_fileread(files, Detectionplatform,reportinfo, project, platform, manufacturers, Unit, digits, ZP_Method_precursor_ion, ZP_Method_product_ion, normAB):

      # 第一步:后台数据抓取（最小样本数，最大允许CV）
      # id1 = Special.objects.get(project=project).id
      # id2 = Stabilityspecial.objects.get(special_id=id1).id

      # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
      # if Repeatprecisionspecialmethod.objects.filter(repeatprecisionspecial=id2):
      #     number = Repeatprecisionspecialmethod.objects.get(
      #         repeatprecisionspecial=id2).minSample  # 最小样本数
      #     CV = Repeatprecisionspecialmethod.objects.get(
      #         repeatprecisionspecial=id2).maxCV  # 最大允许CV

      # else:

      general_1 = General.objects.get(name="通用性项目")
      stability_general = Stabilitygeneral.objects.get(general=general_1)
      lowrecycle = Stabilitygeneralmethod.objects.get(stabilitygeneral=stability_general).lowvalue  # 回收率下限
      uprecycle = Stabilitygeneralmethod.objects.get(stabilitygeneral=stability_general).upvalue  # 回收率上限

      # 后台管理系统-各项目参数设置-PT指标设置里找到是否设置了每个化合物的单位
      special = Special.objects.get(project=project)
      pt_special = PTspecial.objects.get(special=special)
      pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

      # 后台管理系统中设置的本项目化合物名称
      PTnorm = []  
      for i in pt_accept:
            PTnorm.append(i.norm)

      # 后台管理系统中设置的本项目每个化合物单位
      Unitlist = []
      for i in pt_accept:
            Unitlist.append(i.unit)


      #  第二步:开始文件读取
      '''
      数据读取完毕后,需要生成几个字典,分别对应不同的储藏温度。数据格式如下：
      Room_tem_dict = {norm1:{'0h':[低浓度值1,低浓度值2,低浓度值3,中浓度值1,中浓度值2,中浓度值3,高浓度值1,高浓度值2,高浓度值3],'4h':[],'8h':[]},...}
      Refrigerate_tem_dict = {norm1:{'0h':[低浓度值1,低浓度值2,低浓度值3,中浓度值1,中浓度值2,中浓度值3,高浓度值1,高浓度值2,高浓度值3]},{'4h':[]},...}
      '''

      # 创新第二步需要生成的结果容器
      Room_tem_dict = {} # 室温字典
      Refrigerate_tem_dict = {} # 冷藏字典
      Freeze_tem_dict = {} # 冷藏字典

      # 创建中英文浓度水平对应字典
      Conclevel_dict = {"L":"低","M":"中","H":"高"}

      # 各仪器平台及各仪器厂家数据读取
      for file in files:
            if platform=="液质":
                  if manufacturers =="Agilent":
                        # 1 读取csv文件（Agilent）
                        file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                        file_data = file.read().decode('utf-8')
                        lines = file_data.split('\r\n')
                        for i in range(len(lines)): 
                              if len(lines[i])!=0:
                                    lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                                    # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                              else:
                                    lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                                    del lines[i] #最后一行如为空行，则删除该元素

                        # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                        norm=[] #化合物列表
                        for j in range(len(lines[0])):
                              if "-Q Results" in lines[0][j]:
                                    if lines[0][j].split("-Q")[0][0]!='"':  # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除  
                                          norm.append(lines[0][j].split("-Q")[0])
                                    else:
                                          norm.append(lines[0][j].split("-Q")[0][1:])

                        # 从第二行确定:实验号（Sample Name）,浓度（Final Conc.）的索引
                        nameindex=0  #实验号索引
                        concindex=[] #浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                        for j in range(len(lines[1])):  #从第二行开始       
                              if lines[1][j] == "Sample Name" :
                                    nameindex=j
                              elif lines[1][j]  == "Final Conc." :
                                    concindex.append(j)

                        # 定义各验证条件下(室温，冷藏，冷冻)的时间列表
                        Room_tem_list = []
                        Refrigerate_tem_list = []
                        Freeze_tem_list = []

                        # 定义各验证条件下(室温，冷藏，冷冻)的浓度水平(低中高)列表
                        Room_conclevel_list = []
                        Refrigerate_conclevel_list = []
                        Freeze_conclevel_list = []

                        # 确定每个验证条件下（室温，冷藏，冷冻）的验证时间 
                        for i in range(len(lines)): # 循环原始数据中的每一行                   
                              if "Stability" in lines[i][nameindex]: # 稳定性，命名前必须含有关键字段“Stability”

                                    # 1 含有关键字段"RT",说明为室温
                                    if "RT" in lines[i][nameindex]:  
                                          # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间，“L”为浓度水平
                                          
                                          # 1.1 添加浓度水平
                                          # 不论samplename长度为多少，浓度水平固定为第11位
                                          if Conclevel_dict[lines[i][nameindex][10:11]] not in Room_conclevel_list: # 避免重复加入列表
                                                Room_conclevel_list.append(Conclevel_dict[lines[i][nameindex][10:11]])  # 取出浓度水平位放入列表中
                                          
                                          # 1.2 添加时间
                                          if len(lines[i][nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                if lines[i][nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                      Room_tem_list.append(lines[i][nameindex][11:12])  # 取出时间位放入列表中                                   
                                          elif len(lines[i][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                if lines[i][nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                      Room_tem_list.append(lines[i][nameindex][11:13])  # 取出时间位放入列表中

                                    # 2 含有关键字段"Refrigerate",说明为冷藏
                                    elif "Refrigerate" in lines[i][nameindex]:  
                                          # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间

                                          # 2.1 添加浓度水平
                                          # 不论samplename长度为多少，浓度水平固定为第11位
                                          if Conclevel_dict[lines[i][nameindex][10:11]] not in Refrigerate_conclevel_list: # 避免重复加入列表
                                                Refrigerate_conclevel_list.append(Conclevel_dict[lines[i][nameindex][10:11]])  # 取出浓度水平位放入列表中

                                          # 2.2 添加时间
                                          if len(lines[i][nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                if lines[i][nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                      Refrigerate_tem_list.append(lines[i][nameindex][11:12])  # 取出时间位放入列表中
                                          elif len(lines[i][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                if lines[i][nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                      Refrigerate_tem_list.append(lines[i][nameindex][11:13])  # 取出时间位放入列表中

                                    # 3 含有关键字段"Freeze",说明为冷冻
                                    elif "Freeze" in lines[i][nameindex]:  
                                          # samplename格式固定为“Stability-L5h-Freeze01”,其中“5”为时间

                                          # 2.1 添加浓度水平
                                          # 不论samplename长度为多少，浓度水平固定为第11位
                                          if Conclevel_dict[lines[i][nameindex][10:11]] not in Freeze_conclevel_list: # 避免重复加入列表
                                                Freeze_conclevel_list.append(Conclevel_dict[lines[i][nameindex][10:11]])  # 取出浓度水平位放入列表中

                                          # 2.2 添加时间
                                          if len(lines[i][nameindex]) == 22: # 依据上述格式，长度为22，说明时间仅有一位
                                                if lines[i][nameindex][11:12] not in Freeze_tem_list: # 避免重复加入列表
                                                      Freeze_tem_list.append(lines[i][nameindex][11:12])  # 取出时间位放入列表中
                                          elif len(lines[i][nameindex]) == 23: # 依据上述格式，长度为23，说明时间有两位
                                                if lines[i][nameindex][11:13] not in Freeze_tem_list: # 避免重复加入列表
                                                      Freeze_tem_list.append(lines[i][nameindex][11:13])  # 取出时间位放入列表中

                        print(Room_conclevel_list)
                        print(Refrigerate_conclevel_list)
                        print(Freeze_conclevel_list)

                        # 生成第二步的字典         
                        for m in range(len(norm)):  # 循环每个化合物
                              # 1 添加室温数据
                              norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典  
                              for n in Room_tem_list: # 先循环室温条件下的时间列表
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(len(lines)):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in lines[q][nameindex] and "RT" in lines[q][nameindex]: # 判断是否为室温条件                                                                     
                                                if len(lines[q][nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                      if lines[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                            time_list.append(effectnum(lines[q][concindex[m]],digits))
                                                elif len(lines[q][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                      if lines[q][nameindex][11:13] == n: 
                                                            time_list.append(effectnum(lines[q][concindex[m]],digits))
                                          
                                    norm_Room_tem_dict[n] = time_list
                              Room_tem_dict[norm[m]] = norm_Room_tem_dict
                              
                              # 2 添加冷藏数据
                              norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                              for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(len(lines)):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in lines[q][nameindex] and "Refrigerate" in lines[q][nameindex]: # 判断是否为冷藏条件                                                                     
                                                if len(lines[q][nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                      if lines[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                            time_list.append(effectnum(lines[q][concindex[m]],digits))
                                                elif len(lines[q][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                      if lines[q][nameindex][11:13] == n: 
                                                            time_list.append(effectnum(lines[q][concindex[m]],digits))
                                          
                                    norm_Refrigerate_tem_dict[n] = time_list
                              Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

                              # 3 添加冷冻数据
                              norm_Freeze_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                              for n in Freeze_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(len(lines)):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in lines[q][nameindex] and "Freeze" in lines[q][nameindex]: # 判断是否为冷藏条件                                                                     
                                                if len(lines[q][nameindex]) == 22: # 长度为22，说明时间仅有一位,且第12位为时间
                                                      if lines[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                            time_list.append(effectnum(lines[q][concindex[m]],digits))
                                                elif len(lines[q][nameindex]) == 23: # 依据上述格式，长度为23，说明时间有两位,且第12-13位为时间
                                                      if lines[q][nameindex][11:13] == n: 
                                                            time_list.append(effectnum(lines[q][concindex[m]],digits))
                                          
                                    norm_Freeze_tem_dict[n] = time_list
                              Freeze_tem_dict[norm[m]] = norm_Freeze_tem_dict
                        

                  elif manufacturers == "岛津":
                        # 读取txt文件
                        content = []
                        for line in file:
                              content.append(line.decode("UTF-8").replace("\r\n", "").split("\t")) # windows下
                              # content.append(line.decode("GB2312").replace("\r\n", "").split("\t")) # linux下

                        nameindex = 0
                        concindex = 0  # 浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                        norm = []  # 化合物列表
                        norm_row = []  # 化合物所在行

                        # txt数据抓取，content[行索引][列索引]
                        for i in range(len(content)):
                              if content[i][0] == "Name":  # 如果某一行第一列为"Name"，则该行第二列为化合物名称
                                    norm.append(content[i][1])
                                    norm_row.append(i)
                        
                        print(norm_row)
                        print(norm)

                        for i in range(len(content[2])):  # 第二行确定samplename和浓度所在列
                              if content[2][i] == "数据文件名":
                                    nameindex = i
                              elif content[2][i] == "浓度":
                                    concindex = i

                        # 生成第二步的字典(岛津的数据格式与Agilent不同，每个化合物不共用同一个samplename，因此稳定性下的验证时间可能不同，需要分开定义时间列表)         
                        for m in range(len(norm)):  # 循环每个化合物
                              # 定义室温和冷藏条件下的时间列表
                              Room_tem_list = []
                              Refrigerate_tem_list = []

                              # 分别添加每一个化合物的验证时间列表
                              if m < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                    for i in range(norm_row[m], norm_row[m+1]):
                                          if "Stability" in content[i][nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                                # 含有关键字段"RT",说明为室温
                                                if "RT" in content[i][nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                      if len(content[i][nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                            if content[i][nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(content[i][nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(content[i][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                            if content[i][nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(content[i][nameindex][11:13])  # 取出时间位放入列表中
                                                # 含有关键字段"Refrigerate",说明为冷藏
                                                elif "Refrigerate" in content[i][nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                      if len(content[i][nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                            if content[i][nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(content[i][nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(content[i][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                            if content[i][nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(content[i][nameindex][11:13])  # 取出时间位放入列表中

                                    norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                                    for n in Room_tem_list: # 先循环室温条件下的时间列表
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                                if "Stability" in content[q][nameindex] and "RT" in content[q][nameindex]: # 判断是否为室温条件                                                                     
                                                      if len(content[q][nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                            if content[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                                  time_list.append(effectnum(content[q][concindex],digits))
                                                      elif len(content[q][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                            if content[q][nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(content[q][concindex],digits))
                                                
                                          norm_Room_tem_dict[n] = time_list
                                    Room_tem_dict[norm[m]] = norm_Room_tem_dict

                                    norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                                    for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                                if "Stability" in content[q][nameindex] and "Refrigerate" in content[q][nameindex]: # 判断是否为冷藏条件                                                                     
                                                      if len(content[q][nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                            if content[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                                  time_list.append(effectnum(content[q][concindex],digits))
                                                      elif len(content[q][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                            if content[q][nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(content[q][concindex],digits)) 
                                                
                                          norm_Refrigerate_tem_dict[n] = time_list
                                    Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

                              else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                                    for i in range(norm_row[m], len(content)):
                                          if "Stability" in content[i][nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                                # 含有关键字段"RT",说明为室温
                                                if "RT" in content[i][nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                      if len(content[i][nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                            if content[i][nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(content[i][nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(content[i][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                            if content[i][nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(content[i][nameindex][11:13])  # 取出时间位放入列表中
                                                # 含有关键字段"Refrigerate",说明为冷藏
                                                elif "Refrigerate" in content[i][nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                      if len(content[i][nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                            if content[i][nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(content[i][nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(content[i][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                            if content[i][nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(content[i][nameindex][11:13])  # 取出时间位放入列表中

                                    norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                                    for n in Room_tem_list: # 先循环室温条件下的时间列表
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], len(content)):  # 循环每一行，如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                                if "Stability" in content[q][nameindex] and "RT" in content[q][nameindex]: # 判断是否为室温条件                                                                     
                                                      if len(content[q][nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                            if content[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                                  time_list.append(effectnum(content[q][concindex],digits))
                                                      elif len(content[q][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                            if content[q][nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(content[q][concindex],digits))
                                                
                                          norm_Room_tem_dict[n] = time_list
                                    Room_tem_dict[norm[m]] = norm_Room_tem_dict

                                    norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                                    for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], len(content)):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                                if "Stability" in content[q][nameindex] and "Refrigerate" in content[q][nameindex]: # 判断是否为冷藏条件                                                                     
                                                      if len(content[q][nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                            if content[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                                  time_list.append(effectnum(content[q][concindex],digits))
                                                      elif len(content[q][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                            if content[q][nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(content[q][concindex],digits)) 
                                                
                                          norm_Refrigerate_tem_dict[n] = time_list
                                    Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

                  elif manufacturers == "Waters":
                        # 若是最新的 2.0.1 版本的xlrd包，只支持 .xls 文件，读取.xlsx文件会报错。若要正常读取，需安装旧版本的xlrd：pip3 install xlrd==1.2.0
                        data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                        file_data = data.sheets()[0]
                        nrows = file_data.nrows
                        ncols = file_data.ncols

                        norm = []  # 化合物列表
                        norm_row = []  # 化合物所在行
                        for i in range(nrows):
                              if "Compound" in str(file_data.row_values(i)[0]) and ":" in str(file_data.row_values(i)[0]):  # 如果某一行第一列含有关键词"Compound"，则该行中含有化合物名称，化合物名称在：后
                                    norm.append(file_data.row_values(i)[0].split(":")[1].strip()) # strip()的作用是去除前后空格
                                    norm_row.append(i)    

                        print(norm_row)          

                        nameindex = 0
                        concindex = 0
                        # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                        # xlsx数据抓取，file_data.row_values(行索引)[列索引]
                        for i in range(len(file_data.row_values(norm_row[0]+2))):
                              if file_data.row_values(norm_row[0]+2)[i] == "Name":
                                    nameindex = i
                              elif "实际浓度" in file_data.row_values(norm_row[0]+2)[i]:
                                    concindex = i

                        # 生成第二步的字典(Waters的数据格式与Agilent不同，每个化合物不共用同一个samplename，因此稳定性下的验证时间可能不同，需要分开定义时间列表)         
                        for m in range(len(norm)):  # 循环每个化合物
                              # 定义室温和冷藏条件下的时间列表
                              Room_tem_list = []
                              Refrigerate_tem_list = []

                              # 分别添加每一个化合物的验证时间列表
                              if m < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                    for i in range(norm_row[m], norm_row[m+1]):
                                          if "Stability" in file_data.row_values(i)[nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                                # 含有关键字段"RT",说明为室温
                                                if "RT" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中
                                                # 含有关键字段"Refrigerate",说明为冷藏
                                                elif "Refrigerate" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中

                                    norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                                    for n in Room_tem_list: # 先循环室温条件下的时间列表
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                                if "Stability" in file_data.row_values(q)[nameindex] and "RT" in file_data.row_values(q)[nameindex]: # 判断是否为室温条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                
                                          norm_Room_tem_dict[n] = time_list
                                    Room_tem_dict[norm[m]] = norm_Room_tem_dict

                                    norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                                    for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                                if "Stability" in file_data.row_values(q)[nameindex] and "Refrigerate" in file_data.row_values(q)[nameindex]: # 判断是否为冷藏条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits)) 
                                                
                                          norm_Refrigerate_tem_dict[n] = time_list
                                    Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

                              else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                                    for i in range(norm_row[m], nrows):
                                          if "Stability" in file_data.row_values(i)[nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                                # 含有关键字段"RT",说明为室温
                                                if "RT" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中
                                                # 含有关键字段"Refrigerate",说明为冷藏
                                                elif "Refrigerate" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中

                                    norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                                    for n in Room_tem_list: # 先循环室温条件下的时间列表
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], nrows):  # 循环每一行，如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                                if "Stability" in file_data.row_values(q)[nameindex] and "RT" in file_data.row_values(q)[nameindex]: # 判断是否为室温条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                
                                          norm_Room_tem_dict[n] = time_list
                                    Room_tem_dict[norm[m]] = norm_Room_tem_dict

                                    norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                                    for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], nrows):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                                if "Stability" in file_data.row_values(q)[nameindex] and "Refrigerate" in file_data.row_values(q)[nameindex]: # 判断是否为冷藏条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits)) 
                                                
                                          norm_Refrigerate_tem_dict[n] = time_list
                                    Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

                  # Thermo厂家需先在后台管理系统中设置本项目的化合物名称，以便查找上传文件中相应化合物的表格
                  elif manufacturers == "Thermo":
                        Thermo = Special.objects.get(project=project)
                        pt_special = PTspecial.objects.get(special=Thermo)
                        pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)

                        # 后台管理系统中设置的本项目化合物名称
                        PTnorm = []  
                        for i in pt_accept:
                              PTnorm.append(i.norm)

                        data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                        norm = []  # Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                        sheetindex = []  # 
                        for index in range(len(data.sheet_names())):
                              if data.sheet_names()[index] in PTnorm:
                                    norm.append(data.sheet_names()[index])
                                    sheetindex.append(index)

                        # 生成第二步的字典(Thermo的数据格式与Agilent不同，每个化合物的数据在不同的sheet工作表下，因此稳定性下的验证时间可能不同，需要分开定义时间列表)     
                        # 循环读取每个sheet工作表,即为每个化合物的表
                        for index in range(len(sheetindex)):
                              file_data = data.sheets()[sheetindex[index]]
                              nrows = file_data.nrows
                              ncols = file_data.ncols

                              # 第一行确定samplename和浓度所在列
                              nameindex = 0
                              concindex = 0
                              for i in range(len(file_data.row_values(0))):
                                    if file_data.row_values(0)[i] == "Compound":
                                          nameindex = i
                                    elif file_data.row_values(0)[i] == "Calculated Amt":
                                          concindex = i

                              # 定义室温和冷藏条件下的时间列表
                              Room_tem_list = []
                              Refrigerate_tem_list = []

                              for i in range(nrows):
                                    if "Stability" in file_data.row_values(i)[nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                          # 含有关键字段"RT",说明为室温
                                          if "RT" in file_data.row_values(i)[nameindex]:  
                                                # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                if len(file_data.row_values(i)[nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                      if file_data.row_values(i)[nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                            Room_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                elif len(file_data.row_values(i)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                      if file_data.row_values(i)[nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                            Room_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中
                                          # 含有关键字段"Refrigerate",说明为冷藏
                                          elif "Refrigerate" in file_data.row_values(i)[nameindex]:  
                                                # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                if len(file_data.row_values(i)[nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                      if file_data.row_values(i)[nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                            Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                elif len(file_data.row_values(i)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                      if file_data.row_values(i)[nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                            Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中

                              norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                              for n in Room_tem_list: # 先循环室温条件下的时间列表
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(nrows):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in file_data.row_values(q)[nameindex] and "RT" in file_data.row_values(q)[nameindex]: # 判断是否为室温条件                                                                     
                                                if len(file_data.row_values(q)[nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                      if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                            time_list.append(float(effectnum(file_data.row_values(q)[concindex],digits)))  # 取出浓度放入每个时间列表中,需转换为float格式
                                                elif len(file_data.row_values(q)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                      if file_data.row_values(q)[nameindex][11:13] == n: 
                                                            time_list.append(float(effectnum(file_data.row_values(q)[concindex],digits)))
                                          
                                    norm_Room_tem_dict[n] = time_list
                              Room_tem_dict[norm[index]] = norm_Room_tem_dict

                              norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                              for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(nrows):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in file_data.row_values(q)[nameindex] and "Refrigerate" in file_data.row_values(q)[nameindex]: # 判断是否为冷藏条件                                                                     
                                                if len(file_data.row_values(q)[nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                      if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                            time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                elif len(file_data.row_values(q)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                      if file_data.row_values(q)[nameindex][11:13] == n: 
                                                            time_list.append(effectnum(file_data.row_values(q)[concindex],digits)) 
                                          
                                    norm_Refrigerate_tem_dict[n] = time_list
                              Refrigerate_tem_dict[norm[index]] = norm_Refrigerate_tem_dict
                  
                  # AB厂家需先在后台设置化合物和相应离子对数值，以便判断需要读取的表格（定量表格）和不需要读取的表格(定性表格)
                  # 后台管理系统中进行设置时，化合物的设置顺序需与文件中的化合物排列顺序一致，即从上向下进行设置
                  elif manufacturers == "AB":

                        # 测试文件中的三个化合物（1  Clozapine定性 (327.2 / 192.2)，Clozapine定量 (327.2 / 270.2);
                        #                     2  Sertraline定性 (306.1 / 275.1)，Sertraline定量 (306.1 / 159.1)）
                        #                     3  Aripiprazole定性 (448.1 / 176.2),Aripiprazole定量 (448.1 / 285.2)

                        # 定义化合物列表，列表统一命名为norm
                        norm = normAB

                        # 获取上传的文件
                        file_data = Document(file)

                        # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                        paragraphs = [] 

                        # 若标题长度为0或为换行等，不添加进入标题内容列表
                        for p in file_data.paragraphs:
                              if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                                    paragraphs.append(p.text.strip())

                        # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                        tableindex = []
                        for i in range(len(paragraphs)):
                              for j in range(len(ZP_Method_precursor_ion)):
                                    if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                          tableindex.append(2*i+1)

                        tables = file_data.tables  # 获取文件中的表格集

                        # 循环定量表格的索引,生成第二步的字典(AB的数据格式与Agilent不同，每个化合物不共用同一个samplename，因此稳定性下的验证时间可能不同，需要分开定义时间列表) 
                        for k in range(len(tableindex)):
                              # 获取文件中的定量表格
                              tablequantify = tables[tableindex[k]] 

                              # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                              cells = tablequantify._cells
                              ROWS = len(tablequantify.rows)
                              COLUMNS = len(tablequantify.columns)
                              rowdatalist = []  # 每一行的数据列表
                              rowdatagatherlist = []  # 每一行的数据列表汇总列表
                              for i in range(ROWS*COLUMNS):
                                    text = cells[i].text.replace("\n", "")
                                    text = text.strip()  # 去除空白符
                                    if i % 12 != 0 or i == 0:  # docx文件固定为12列
                                          rowdatalist.append(text)
                                    else:
                                          rowdatagatherlist.append(rowdatalist)
                                          rowdatalist = []
                                          rowdatalist.append(text)
                              rowdatagatherlist.append(rowdatalist)

                              # 读取表格的第一行的单元格,判断实验号和浓度索引
                              nameindex = 0
                              concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列
                              for i in range(len(rowdatagatherlist[0])):
                                    if rowdatagatherlist[0][i] == "Sample Name":
                                          nameindex = i
                                    elif "Calculated Conc" in rowdatagatherlist[0][i]:
                                          concindex = i

                              # 定义室温和冷藏条件下的时间列表
                              Room_tem_list = []
                              Refrigerate_tem_list = []

                              # 分别添加每一个化合物的验证时间列表
                              for i in range(len(rowdatagatherlist)):
                                    if "Stability" in rowdatagatherlist[i][nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                          # 含有关键字段"RT",说明为室温
                                          if "RT" in rowdatagatherlist[i][nameindex]:  
                                                # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                if len(rowdatagatherlist[i][nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                      if rowdatagatherlist[i][nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                            Room_tem_list.append(rowdatagatherlist[i][nameindex][11:12])  # 取出时间位放入列表中
                                                elif len(rowdatagatherlist[i][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                      if rowdatagatherlist[i][nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                            Room_tem_list.append(rowdatagatherlist[i][nameindex][11:13])  # 取出时间位放入列表中
                                          # 含有关键字段"Refrigerate",说明为冷藏
                                          elif "Refrigerate" in rowdatagatherlist[i][nameindex]:  
                                                # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                if len(rowdatagatherlist[i][nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                      if rowdatagatherlist[i][nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                            Refrigerate_tem_list.append(rowdatagatherlist[i][nameindex][11:12])  # 取出时间位放入列表中
                                                elif len(rowdatagatherlist[i][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                      if rowdatagatherlist[i][nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                            Refrigerate_tem_list.append(rowdatagatherlist[i][nameindex][11:13])  # 取出时间位放入列表中

                              norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                              for n in Room_tem_list: # 先循环室温条件下的时间列表
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(len(rowdatagatherlist)): 
                                          if "Stability" in rowdatagatherlist[q][nameindex] and "RT" in rowdatagatherlist[q][nameindex]: # 判断是否为室温条件                                                                     
                                                if len(rowdatagatherlist[q][nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                      if rowdatagatherlist[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间  
                                                            time_list.append(effectnum(rowdatagatherlist[q][concindex],digits))  # 取出浓度放入每个时间列表中,需转换为float格式                
                                                            # time_list.append(float(effectnum(rowdatagatherlist[q][concindex],digits)))  # 取出浓度放入每个时间列表中,需转换为float格式
                                                elif len(rowdatagatherlist[q][nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                      if rowdatagatherlist[q][nameindex][11:13] == n: 
                                                            time_list.append(effectnum(rowdatagatherlist[q][concindex],digits))  # 取出浓度放入每个时间列表中,需转换为float格式
                                                            # time_list.append(float(effectnum(rowdatagatherlist[q][concindex],digits)))
                                          
                                    norm_Room_tem_dict[n] = time_list
                              Room_tem_dict[norm[k]] = norm_Room_tem_dict

                              norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                              for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in rowdatagatherlist[q][nameindex] and "Refrigerate" in rowdatagatherlist[q][nameindex]: # 判断是否为冷藏条件                                                                     
                                                if len(rowdatagatherlist[q][nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                      if rowdatagatherlist[q][nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                            time_list.append(float(effectnum(rowdatagatherlist[q][concindex],digits)))  # 取出浓度放入每个时间列表中,需转换为float格式
                                                elif len(rowdatagatherlist[q][nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                      if rowdatagatherlist[q][nameindex][11:13] == n: 
                                                            time_list.append(float(effectnum(rowdatagatherlist[i][concindex],digits))) 
                                          
                                    norm_Refrigerate_tem_dict[n] = time_list
                              Refrigerate_tem_dict[norm[k]] = norm_Refrigerate_tem_dict

            elif platform == "液相":
                  if manufacturers == "Agilent":
                        data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                        file_data = data.sheets()[0]  # 默认只读取第一个工作簿
                        nrows = file_data.nrows
                        ncols = file_data.ncols

                        # 确定化合物名称,并添加进入化合物列表
                        norm = []  # 化合物列表
                        norm_row = []  # 化合物所在行
                        for j in range(nrows):
                              # 如果某一行的第一个元素为'化合物'，则添加第三个元素进入化合物列表
                              if file_data.row_values(j)[0] == "化合物:":
                                    norm.append(file_data.row_values(j)[2])
                                    norm_row.append(j)

                        # 第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                        nameindex = 0
                        concindex = 0                     
                        for i in range(len(file_data.row_values(norm_row[0]+2))):
                              if file_data.row_values(norm_row[0]+2)[i] == "样品名称":
                                    nameindex = i
                              elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                                    concindex = i

                        # 生成第二步的字典(液相Agilent数据格式同液质Waters)         
                        for m in range(len(norm)):  # 循环每个化合物
                              # 定义室温和冷藏条件下的时间列表
                              Room_tem_list = []
                              Refrigerate_tem_list = []

                              # 分别添加每一个化合物的验证时间列表
                              if m < len(norm)-1:  # 如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                    for i in range(norm_row[m], norm_row[m+1]):
                                          if "Stability" in file_data.row_values(i)[nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                                # 含有关键字段"RT",说明为室温
                                                if "RT" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中
                                                # 含有关键字段"Refrigerate",说明为冷藏
                                                elif "Refrigerate" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中

                                    norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                                    for n in Room_tem_list: # 先循环室温条件下的时间列表
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                                if "Stability" in file_data.row_values(q)[nameindex] and "RT" in file_data.row_values(q)[nameindex]: # 判断是否为室温条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                
                                          norm_Room_tem_dict[n] = time_list
                                    Room_tem_dict[norm[m]] = norm_Room_tem_dict

                                    norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                                    for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], norm_row[m+1]):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                                if "Stability" in file_data.row_values(q)[nameindex] and "Refrigerate" in file_data.row_values(q)[nameindex]: # 判断是否为冷藏条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits)) 
                                                
                                          norm_Refrigerate_tem_dict[n] = time_list
                                    Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

                              else:  # 如果是最后一个化合物，索引为该化合物所在行到总行数
                                    for i in range(norm_row[m], nrows):
                                          if "Stability" in file_data.row_values(i)[nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                                # 含有关键字段"RT",说明为室温
                                                if "RT" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                                  Room_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中
                                                # 含有关键字段"Refrigerate",说明为冷藏
                                                elif "Refrigerate" in file_data.row_values(i)[nameindex]:  
                                                      # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                                      if len(file_data.row_values(i)[nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                            if file_data.row_values(i)[nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                                      elif len(file_data.row_values(i)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                            if file_data.row_values(i)[nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                                  Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中

                                    norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典
                                    for n in Room_tem_list: # 先循环室温条件下的时间列表
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], nrows):  # 循环每一行，如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                                if "Stability" in file_data.row_values(q)[nameindex] and "RT" in file_data.row_values(q)[nameindex]: # 判断是否为室温条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                
                                          norm_Room_tem_dict[n] = time_list
                                    Room_tem_dict[norm[m]] = norm_Room_tem_dict

                                    norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                                    for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                          time_list=[] # 每个化合物下每个验证时间的数据列表
                                          for q in range(norm_row[m], nrows):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                                if "Stability" in file_data.row_values(q)[nameindex] and "Refrigerate" in file_data.row_values(q)[nameindex]: # 判断是否为冷藏条件                                                                     
                                                      if len(file_data.row_values(q)[nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                            if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits))
                                                      elif len(file_data.row_values(q)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                            if file_data.row_values(q)[nameindex][11:13] == n: 
                                                                  time_list.append(effectnum(file_data.row_values(q)[concindex],digits)) 
                                                
                                          norm_Refrigerate_tem_dict[n] = time_list
                                    Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict
            
            elif platform == "ICP-MS":
                  # ICP-MS平台Agilent厂家需先在后台管理系统中设置本项目的化合物名称，以便查找上传文件中相应化合物的表格
                  if manufacturers == "Agilent":
                        data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                        file_data = data.sheets()[0]
                        nrows = file_data.nrows
                        ncols = file_data.ncols

                        # 从第一行确定化合物名称
                        norm = []
                        for j in range(ncols):
                              for i in PTnorm:
                                    if i in file_data.row_values(0)[j]:
                                          norm.append(i)

                        # 从第二行确定实验号（Sample Name）的索引和化合物浓度索引
                        nameindex = 0  # 实验号索引
                        concindex = []  # 浓度索引
                        for j in range(ncols):
                              if file_data.row_values(1)[j] == "样品名称":
                                    nameindex = j
                              elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                                    concindex.append(j)

                        # 定义室温和冷藏条件下的时间列表
                        Room_tem_list = []
                        Refrigerate_tem_list = []

                        # 确定每个验证条件下（室温，冷藏，冷冻）的验证时间 
                        for i in range(2, nrows): # 循环原始数据中的每一行                   
                              if "Stability" in file_data.row_values(i)[nameindex]: # 稳定性，命名前必须含有关键字段“Stability”
                                    # 含有关键字段"RT",说明为室温
                                    if "RT" in file_data.row_values(i)[nameindex]:  
                                          # samplename格式固定为“Stability-L5h-RT01”,其中“5”为时间
                                          if len(file_data.row_values(i)[nameindex]) == 18: # 依据上述格式，长度为18，说明时间仅有一位
                                                if file_data.row_values(i)[nameindex][11:12] not in Room_tem_list: # 避免重复加入列表
                                                      Room_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                          elif len(file_data.row_values(i)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位
                                                if file_data.row_values(i)[nameindex][11:13] not in Room_tem_list: # 避免重复加入列表
                                                      Room_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中
                                    # 含有关键字段"Refrigerate",说明为冷藏
                                    elif "Refrigerate" in file_data.row_values(i)[nameindex]:  
                                          # samplename格式固定为“Stability-L5h-Refrigerate01”,其中“5”为时间
                                          if len(file_data.row_values(i)[nameindex]) == 27: # 依据上述格式，长度为27，说明时间仅有一位
                                                if file_data.row_values(i)[nameindex][11:12] not in Refrigerate_tem_list: # 避免重复加入列表
                                                      Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:12])  # 取出时间位放入列表中
                                          elif len(file_data.row_values(i)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位
                                                if file_data.row_values(i)[nameindex][11:13] not in Refrigerate_tem_list: # 避免重复加入列表
                                                      Refrigerate_tem_list.append(file_data.row_values(i)[nameindex][11:13])  # 取出时间位放入列表中

                        # 生成第二步的字典    file_data.row_values(q)[nameindex]     
                        for m in range(len(norm)):  # 循环每个化合物
                              norm_Room_tem_dict = {} # 每个化合物室温条件下的数据字典  
                              for n in Room_tem_list: # 先循环室温条件下的时间列表
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(2, nrows):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in file_data.row_values(q)[nameindex] and "RT" in file_data.row_values(q)[nameindex]: # 判断是否为室温条件                                                                     
                                                if len(file_data.row_values(q)[nameindex]) == 18: # 长度为18，说明时间仅有一位,且第12位为时间
                                                      if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间                  
                                                            time_list.append(effectnum(file_data.row_values(q)[concindex[m]],digits))
                                                elif len(file_data.row_values(q)[nameindex]) == 19: # 依据上述格式，长度为19，说明时间有两位,且第12-13位为时间
                                                      if file_data.row_values(q)[nameindex][11:13] == n: 
                                                            time_list.append(effectnum(file_data.row_values(q)[concindex[m]],digits))
                                          
                                    norm_Room_tem_dict[n] = time_list
                              Room_tem_dict[norm[m]] = norm_Room_tem_dict
                              
                              norm_Refrigerate_tem_dict = {} # 每个化合物冷藏条件下的数据字典
                              for n in Refrigerate_tem_list: # 再循环冷藏条件下的时间列表，先写两个，后面需要加上冷冻条件
                                    time_list=[] # 每个化合物下每个验证时间的数据列表
                                    for q in range(2, nrows):  # 循环每一行，后续优化时可以判断每个验证指标从哪一行开始就从哪一行开始循环，减少循环数据量
                                          if "Stability" in file_data.row_values(q)[nameindex] and "Refrigerate" in file_data.row_values(q)[nameindex]: # 判断是否为冷藏条件                                                                     
                                                if len(file_data.row_values(q)[nameindex]) == 27: # 长度为27，说明时间仅有一位,且第12位为时间
                                                      if file_data.row_values(q)[nameindex][11:12] == n:  # 判断时间位是否等于循环中的时间
                                                            time_list.append(effectnum(file_data.row_values(q)[concindex[m]],digits))
                                                elif len(file_data.row_values(q)[nameindex]) == 28: # 依据上述格式，长度为28，说明时间有两位,且第12-13位为时间
                                                      if file_data.row_values(q)[nameindex][11:13] == n: 
                                                            time_list.append(effectnum(file_data.row_values(q)[concindex[m]],digits))
                                          
                                    norm_Refrigerate_tem_dict[n] = time_list
                              Refrigerate_tem_dict[norm[m]] = norm_Refrigerate_tem_dict

      print(Room_tem_dict) 
      print(Refrigerate_tem_dict) 
      print(Freeze_tem_dict)   
               
      #  第三步:计算平均值，CV，回收率

      '''
      最终需要生成几个字典,分别对应不同的储藏温度,数据格式如下：
      Room_tem_dict = {norm1:{'0h':[低浓度值1,低浓度值2,低浓度值3,中浓度值1,中浓度值2,中浓度值3,高浓度值1,高浓度值2,高浓度值3,低浓度均值,中浓度均值,高浓度均值,低浓度CV,中浓度CV,高浓度CV,
      ,低浓度回收率1,低浓度回收率2,低浓度回收率3,中浓度回收率1,中浓度回收率2,中浓度回收率3,高浓度回收率1,高浓度回收率2,高浓度回收率3], 
      '4h':[],'8h':[]},...}
      '''
      
      # 1 室温 
      # 计算均值，cv
      '''
      Room_tem_dict = {norm1:{'0h':[低浓度值1,低浓度值2,低浓度值3,中浓度值1,中浓度值2,中浓度值3,高浓度值1,高浓度值2,高浓度值3,低浓度均值,中浓度均值,高浓度均值,低浓度CV,中浓度CV,高浓度CV], 
      '4h':[],'8h':[]},...}
      '''

      # 1.1 3个浓度水平的计算
      if len(Room_conclevel_list)==3:
            for key,value in Room_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        lowmean = mean(j2[0:3]) # 低浓度均值
                        medianmean = mean(j2[3:6]) # 中浓度均值
                        highmean = mean(j2[6:9]) # 高浓度均值

                        lowcv = cv(j2[0:3]) # 低较
                        mediancv = cv(j2[3:6]) # 中浓度cv
                        highcv = cv(j2[6:9]) # 高浓度cv

                        # 添加上述六个值
                        value[i].extend([lowmean,medianmean,highmean,lowcv,mediancv,highcv])

            # 计算回收率
            for key,value in Room_tem_dict.items():
                  # 定义每个化合物0h时的低浓度，中浓度和高浓度均值
                  low0 = 0
                  median0 = 0
                  high0 = 0 
                  for i,j in value.items():  # 目前默认低中高三个水平全覆盖，因此列表中一共3*3=9个元素，后续需做判断
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加中浓度回收率
                              value[i].extend(["/","/","/"])  # 添加高浓度回收率
                              low0 = j2[9] # 设置0h低浓度均值
                              median0 = j2[10] # 设置0h中浓度均值
                              high0 = j2[11] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              medianrecycle1 = new_round(j2[3]/median0*100)
                              medianrecycle2 = new_round(j2[4]/median0*100)
                              medianrecycle3 = new_round(j2[5]/median0*100)
                              highrecycle1 = new_round(j2[6]/high0*100)
                              highrecycle2 = new_round(j2[7]/high0*100)
                              highrecycle3 = new_round(j2[8]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 添加低浓度回收率
                              value[i].extend([medianrecycle1,medianrecycle2,medianrecycle3])  # 添加中浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加高浓度回收率 

      # 1.2 两个浓度水平的计算
      else:
            for key,value in Room_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        mean1 = mean(j2[0:3]) # 较低浓度均值
                        mean2 = mean(j2[3:6]) # 较高浓度均值

                        cv1 = cv(j2[0:3]) # 较低浓度cv
                        cv2 = cv(j2[3:6]) # 较高浓度cv
                  
                        # 添加上述四个值
                        value[i].extend([mean1,mean2,cv1,cv2])

            # 计算回收率
            for key,value in Room_tem_dict.items():
                  # 定义每个化合物0h时的较低浓度和较高浓度均值
                  low0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加较低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加叫较高浓度回收率
                              low0 = j2[6] # 设置0h低浓度均值
                              high0 = j2[7] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              highrecycle1 = new_round(j2[3]/high0*100)
                              highrecycle2 = new_round(j2[4]/high0*100)
                              highrecycle3 = new_round(j2[5]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 较添加低浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加较高浓度回收率

      # 2 冷藏
      # 计算均值，cv     

      # 2.1 3个浓度水平的计算
      if len(Refrigerate_conclevel_list)==3:           
            for key,value in Refrigerate_tem_dict.items():
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        lowmean = mean(j2[0:3]) # 低浓度均值
                        medianmean = mean(j2[3:6]) # 中浓度均值
                        highmean = mean(j2[6:9]) # 高浓度均值

                        lowcv = cv(j2[0:3]) # 低浓度cv
                        mediancv = cv(j2[3:6]) # 中浓度cv
                        highcv = cv(j2[6:9]) # 高浓度cv

                        # 添加上述六个值
                        value[i].extend([lowmean,medianmean,highmean,lowcv,mediancv,highcv])
      
            # 计算回收率
            for key,value in Refrigerate_tem_dict.items():
                  # 定义每个化合物0h时的低浓度，中浓度和高浓度均值
                  low0 = 0
                  median0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加中浓度回收率
                              value[i].extend(["/","/","/"])  # 添加高浓度回收率
                              low0 = j2[9] # 设置0h低浓度均值
                              median0 = j2[10] # 设置0h中浓度均值
                              high0 = j2[11] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              medianrecycle1 = new_round(j2[3]/median0*100)
                              medianrecycle2 = new_round(j2[4]/median0*100)
                              medianrecycle3 = new_round(j2[5]/median0*100)
                              highrecycle1 = new_round(j2[6]/high0*100)
                              highrecycle2 = new_round(j2[7]/high0*100)
                              highrecycle3 = new_round(j2[8]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 添加低浓度回收率
                              value[i].extend([medianrecycle1,medianrecycle2,medianrecycle3])  # 添加中浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加高浓度回收率 

      else:
            for key,value in Refrigerate_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        mean1 = mean(j2[0:3]) # 较低浓度均值
                        mean2 = mean(j2[3:6]) # 较高浓度均值

                        cv1 = cv(j2[0:3]) # 较低浓度cv
                        cv2 = cv(j2[3:6]) # 较高浓度cv
                  
                        # 添加上述四个值
                        value[i].extend([mean1,mean2,cv1,cv2])

            # 计算回收率
            for key,value in Refrigerate_tem_dict.items():
                  # 定义每个化合物0h时的较低浓度和较高浓度均值
                  low0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加较低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加叫较高浓度回收率
                              low0 = j2[6] # 设置0h低浓度均值
                              high0 = j2[7] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              highrecycle1 = new_round(j2[3]/high0*100)
                              highrecycle2 = new_round(j2[4]/high0*100)
                              highrecycle3 = new_round(j2[5]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 较添加低浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加较高浓度回收率 

      # 3 冷冻
      # 计算均值，cv     

      # 3.1 3个浓度水平的计算
      if len(Freeze_conclevel_list)==3:           
            for key,value in Freeze_tem_dict.items():
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        lowmean = mean(j2[0:3]) # 低浓度均值
                        medianmean = mean(j2[3:6]) # 中浓度均值
                        highmean = mean(j2[6:9]) # 高浓度均值

                        lowcv = cv(j2[0:3]) # 低浓度cv
                        mediancv = cv(j2[3:6]) # 中浓度cv
                        highcv = cv(j2[6:9]) # 高浓度cv

                        # 添加上述六个值
                        value[i].extend([lowmean,medianmean,highmean,lowcv,mediancv,highcv])
      
            # 计算回收率
            for key,value in Freeze_tem_dict.items():
                  # 定义每个化合物0h时的低浓度，中浓度和高浓度均值
                  low0 = 0
                  median0 = 0
                  high0 = 0 
                  for i,j in value.items():  # 目前默认低中高三个水平全覆盖，因此列表中一共3*3=9个元素，后续需做判断
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加中浓度回收率
                              value[i].extend(["/","/","/"])  # 添加高浓度回收率
                              low0 = j2[9] # 设置0h低浓度均值
                              median0 = j2[10] # 设置0h中浓度均值
                              high0 = j2[11] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              medianrecycle1 = new_round(j2[3]/median0*100)
                              medianrecycle2 = new_round(j2[4]/median0*100)
                              medianrecycle3 = new_round(j2[5]/median0*100)
                              highrecycle1 = new_round(j2[6]/high0*100)
                              highrecycle2 = new_round(j2[7]/high0*100)
                              highrecycle3 = new_round(j2[8]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 添加低浓度回收率
                              value[i].extend([medianrecycle1,medianrecycle2,medianrecycle3])  # 添加中浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加高浓度回收率 

      else:
            for key,value in Freeze_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        mean1 = mean(j2[0:3]) # 较低浓度均值
                        mean2 = mean(j2[3:6]) # 较高浓度均值

                        cv1 = cv(j2[0:3]) # 较低浓度cv
                        cv2 = cv(j2[3:6]) # 较高浓度cv
                  
                        # 添加上述四个值
                        value[i].extend([mean1,mean2,cv1,cv2])

            # 计算回收率
            for key,value in Freeze_tem_dict.items():
                  # 定义每个化合物0h时的较低浓度和较高浓度均值
                  low0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加较低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加叫较高浓度回收率
                              low0 = j2[6] # 设置0h低浓度均值
                              high0 = j2[7] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              highrecycle1 = new_round(j2[3]/high0*100)
                              highrecycle2 = new_round(j2[4]/high0*100)
                              highrecycle3 = new_round(j2[5]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 较添加低浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加较高浓度回收率 

      print(Room_tem_dict)
      print(Refrigerate_tem_dict)
      print(Freeze_tem_dict)
      
      # 判断各验证条件下是否有验证数据，没有则不返回结果，在前端不显示

      # # 1 室温下没有数据，返回冷藏和冷冻数据
      # if Refrigerate_tem_dict[norm[0]]=={}:
      #       return {"Room_conclevel_list": Room_conclevel_list,"Room_tem_dict": Room_tem_dict, "Unit": Unit, "lowrecycle": lowrecycle, "uprecycle": uprecycle}
      # else:
      #       return {"Room_conclevel_list": Room_conclevel_list,"Room_tem_dict": Room_tem_dict, "Refrigerate_tem_dict": Refrigerate_tem_dict, "Unit": Unit, "lowrecycle": lowrecycle, "uprecycle": uprecycle}

      return {"Room_conclevel_list": Room_conclevel_list,"Room_tem_dict": Room_tem_dict, 
      "Refrigerate_conclevel_list": Refrigerate_conclevel_list,"Refrigerate_tem_dict": Refrigerate_tem_dict, "Unit": Unit,
      "Freeze_conclevel_list": Freeze_conclevel_list,"Freeze_tem_dict": Freeze_tem_dict, 
      "lowrecycle": lowrecycle, "uprecycle": uprecycle}


def data_scrap(id):

      # 第一步：后台描述性内容数据提取
      # 1 根据id找到项目
      project = ReportInfo.objects.get(id=id).project

      # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
      #特殊参数设置描述性内容
      textlist_special = []
      try:
            special_1 = Special.objects.get(project=project) 
            prepared_Sample_Stability_special = Prepared_Sample_Stability_special.objects.get(special=special_1)           
            if Prepared_Sample_Stability_special_texts.objects.filter(prepared_Sample_Stability_special=prepared_Sample_Stability_special).count()>0:
                  text_special = Prepared_Sample_Stability_special_texts.objects.filter(prepared_Sample_Stability_special=prepared_Sample_Stability_special)  
                  for i in text_special:
                        textlist_special.append(i.text)
      except:
            pass
    
      # 3 通用数据抓取
      general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
      stability_general = Stabilitygeneral.objects.get(general=general_1)
      text_general = Stabilitygeneraltexts.objects.filter(stabilitygeneral=stability_general)   

      # 描述性内容
      textlist_general = [] 
      for i in text_general:
            textlist_general.append(i.text)

      # 4 查找是否单独设置了每个化合物的有效位数
      DIGITS_TABLE = Special.objects.get(project=project) 
      pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
      pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
      Digitslist=[] # 每个化合物有效位数列表
      Digitsdict={} # 每个化合物有效位数字典

      for i in pt_accept:
            Digitslist.append(i.digits)

      if Digitslist==[] or Digitslist[0]=="": #如果全部没设置或者只是单位没设置
            pass
      else:
            for i in pt_accept:
                  Digitsdict[i.norm]=i.digits


      # 第二步：数据库抓取数据

      '''
      1  需要生成几个字典，分别对应不同的验证温度；
            Room_tem_dict = {norm1:{'0h':[低浓度值1,低浓度值2,低浓度值3,中浓度值1,中浓度值2,中浓度值3,高浓度值1,高浓度值2,高浓度值3],'4h':[],'8h':[]},...}
      2  总结性结论：实验结果如表xx-xx所示，表明处理后的样品在室温和冷藏条件下至少可以稳定存放72h；
      '''

      # 定义需要生成的字典
      Room_tem_dict = {} # 室温字典
      Refrigerate_tem_dict = {} # 冷藏字典
      Freeze_tem_dict = {} # 冷冻字典

      # 定义各验证条件下(室温，冷藏，冷冻)的浓度水平(低中高)列表
      Room_conclevel_list = []
      Refrigerate_conclevel_list = []
      Freeze_conclevel_list = []

      try:  
            # 1 基础数据抓取

            # 1.1 室温数据
            Room_tem_data = Stability.objects.filter(reportinfo_id=id,temperature="Room_tem")

            # 化合物名称
            Room_tem_norm = []  
            for i in Room_tem_data:
                  if i.norm not in Room_tem_norm:
                        Room_tem_norm.append(i.norm)

            for j in Room_tem_norm:
                  middle_dict = {}  # 每个化合物的数据
                  middle_list = []  # 每个化合物下各验证时间点的数据
                  middle_table = Stability.objects.filter(reportinfo_id=id,temperature="Room_tem",norm = j)
                  
                  for index,k in enumerate(middle_table):
                        # 首次循环判断该验证条件下各化合物的浓度水平(是否低中高全覆盖)
                        if index == 0:
                              # 3个水平全覆盖
                              if "/" not in [k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]:
                                    ["低","中","高"]
                                    Room_conclevel_list = ["低","中","高"]
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表

                              else:
                                    # 没有低浓度水平
                                    if "/" in [k.L01,k.L02,k.L03]:
                                          Room_conclevel_list = ["中","高"]
                                          middle_list.extend([k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表
                                    # 没有中浓度水平     
                                    elif "/" in [k.M01,k.M02,k.M03]:
                                          Room_conclevel_list = ["低","高"]
                                          middle_list.extend([k.L01,k.L02,k.L03,k.H01,k.H02,k.H03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表
                                    # 没有高浓度水平
                                    else:
                                          Room_conclevel_list = ["低","中"]
                                          middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表

                        # 非首次循环无需再判断浓度水平
                        else:
                              if Room_conclevel_list == ["低","中","高"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Room_conclevel_list == ["中","高"]:
                                    middle_list.extend([k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Room_conclevel_list == ["低","高"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Room_conclevel_list == ["低","中"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表

                  Room_tem_dict[j]=middle_dict

            # 1.2 冷藏数据
            Refrigerate_tem_data = Stability.objects.filter(reportinfo_id=id,temperature="Refrigerate_tem")

            # 化合物名称
            Refrigerate_tem_norm = []  
            for i in Refrigerate_tem_data:
                  if i.norm not in Refrigerate_tem_norm:
                        Refrigerate_tem_norm.append(i.norm)

            for j in Refrigerate_tem_norm:
                  middle_dict = {}  # 每个化合物的数据
                  middle_list = []  # 每个化合物下各验证时间点的数据
                  middle_table = Stability.objects.filter(reportinfo_id=id,temperature="Refrigerate_tem",norm = j)

                  for index,k in enumerate(middle_table):
                        # 首次循环判断该验证条件下各化合物的浓度水平(是否低中高全覆盖)
                        if index == 0:
                              # 3个水平全覆盖
                              if "/" not in [k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]:
                                    Refrigerate_conclevel_list = ["低","中","高"]
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表

                              else:
                                    # 没有低浓度水平
                                    if "/" in [k.L01,k.L02,k.L03]:
                                          Refrigerate_conclevel_list = ["中","高"]
                                          middle_list.extend([k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表
                                    # 没有中浓度水平     
                                    elif "/" in [k.M01,k.M02,k.M03]:
                                          Refrigerate_conclevel_list = ["低","高"]
                                          middle_list.extend([k.L01,k.L02,k.L03,k.H01,k.H02,k.H03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表
                                    # 没有高浓度水平
                                    else:
                                          Refrigerate_conclevel_list = ["低","中"]
                                          middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表

                        # 非首次循环无需再判断浓度水平
                        else:
                              if Refrigerate_conclevel_list == ["低","中","高"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Refrigerate_conclevel_list == ["中","高"]:
                                    middle_list.extend([k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Refrigerate_conclevel_list == ["低","高"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Refrigerate_conclevel_list == ["低","中"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表

                  Refrigerate_tem_dict[j]=middle_dict

            # 1.3 冷冻数据
            Freeze_tem_data = Stability.objects.filter(reportinfo_id=id,temperature="Freeze_tem")

            # 化合物名称
            Freeze_tem_norm = []  
            for i in Freeze_tem_data:
                  if i.norm not in Freeze_tem_norm:
                        Freeze_tem_norm.append(i.norm)

            for j in Freeze_tem_norm:
                  middle_dict = {}  # 每个化合物的数据
                  middle_list = []  # 每个化合物下各验证时间点的数据
                  middle_table = Stability.objects.filter(reportinfo_id=id,temperature="Freeze_tem",norm = j)

                  for index,k in enumerate(middle_table):
                        # 首次循环判断该验证条件下各化合物的浓度水平(是否低中高全覆盖)
                        if index == 0:
                              # 3个水平全覆盖
                              if "/" not in [k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]:
                                    Freeze_conclevel_list = ["低","中","高"]
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表

                              else:
                                    # 没有低浓度水平
                                    if "/" in [k.L01,k.L02,k.L03]:
                                          Freeze_conclevel_list = ["中","高"]
                                          middle_list.extend([k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表
                                    # 没有中浓度水平     
                                    elif "/" in [k.M01,k.M02,k.M03]:
                                          Freeze_conclevel_list = ["低","高"]
                                          middle_list.extend([k.L01,k.L02,k.L03,k.H01,k.H02,k.H03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表
                                    # 没有高浓度水平
                                    else:
                                          Freeze_conclevel_list = ["低","中"]
                                          middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03]) # 添加基础数据
                                          middle_dict[str(int(k.time))]=middle_list # 添加时间
                                          middle_list=[] # 重置列表

                        # 非首次循环无需再判断浓度水平
                        else:
                              if Freeze_conclevel_list == ["低","中","高"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Freeze_conclevel_list == ["中","高"]:
                                    middle_list.extend([k.M01,k.M02,k.M03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Freeze_conclevel_list == ["低","高"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.H01,k.H02,k.H03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表
                              elif Freeze_conclevel_list == ["低","中"]:
                                    middle_list.extend([k.L01,k.L02,k.L03,k.M01,k.M02,k.M03]) # 添加基础数据
                                    middle_dict[str(int(k.time))]=middle_list # 添加时间
                                    middle_list=[] # 重置列表

                  Freeze_tem_dict[j]=middle_dict
      
      except:
            pass
      
      #  2 计算平均值，CV，回收率

      # 2.1 室温 
      '''
      Room_tem_dict = {norm1:{'0h':[低浓度值1,低浓度值2,低浓度值3,中浓度值1,中浓度值2,中浓度值3,高浓度值1,高浓度值2,高浓度值3,低浓度均值,中浓度均值,高浓度均值,低浓度CV,中浓度CV,高浓度CV], 
      '4h':[],'8h':[]},...}
      '''

      # 三个浓度水平的计算
      if len(Room_conclevel_list)==3:
            for key,value in Room_tem_dict.items():
                  for i,j in value.items():
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        
                        lowmean = mean(j2[0:3]) # 低浓度均值
                        medianmean = mean(j2[3:6]) # 中浓度均值
                        highmean = mean(j2[6:9]) # 高浓度均值

                        lowcv = cv(j2[0:3]) # 低浓度cv
                        mediancv = cv(j2[3:6]) # 中浓度cv
                        highcv = cv(j2[6:9]) # 高浓度cv

                        # 添加上述六个值
                        value[i].extend([lowmean,medianmean,highmean,lowcv,mediancv,highcv])

            # 计算回收率
            for key,value in Room_tem_dict.items():
                  # 定义每个化合物0h时的低浓度，中浓度和高浓度均值
                  low0 = 0
                  median0 = 0
                  high0 = 0 
                  for i,j in value.items():
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加中浓度回收率
                              value[i].extend(["/","/","/"])  # 添加高浓度回收率
                              low0 = j2[9] # 设置0h低浓度均值
                              median0 = j2[10] # 设置0h中浓度均值
                              high0 = j2[11] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              medianrecycle1 = new_round(j2[3]/median0*100)
                              medianrecycle2 = new_round(j2[4]/median0*100)
                              medianrecycle3 = new_round(j2[5]/median0*100)
                              highrecycle1 = new_round(j2[6]/high0*100)
                              highrecycle2 = new_round(j2[7]/high0*100)
                              highrecycle3 = new_round(j2[8]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 添加低浓度回收率
                              value[i].extend([medianrecycle1,medianrecycle2,medianrecycle3])  # 添加中浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加高浓度回收率 

      # 两个浓度水平的计算
      else:
            for key,value in Room_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        mean1 = mean(j2[0:3]) # 较低浓度均值
                        mean2 = mean(j2[3:6]) # 较高浓度均值

                        cv1 = cv(j2[0:3]) # 较低浓度cv
                        cv2 = cv(j2[3:6]) # 较高浓度cv
                  
                        # 添加上述四个值
                        value[i].extend([mean1,mean2,cv1,cv2])

            # 计算回收率
            for key,value in Room_tem_dict.items():
                  # 定义每个化合物0h时的较低浓度和较高浓度均值
                  low0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加较低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加叫较高浓度回收率
                              low0 = j2[6] # 设置0h低浓度均值
                              high0 = j2[7] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              highrecycle1 = new_round(j2[3]/high0*100)
                              highrecycle2 = new_round(j2[4]/high0*100)
                              highrecycle3 = new_round(j2[5]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 较添加低浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加较高浓度回收率

      # 2.2 冷藏
      # 三个浓度水平的计算 
      if len(Refrigerate_conclevel_list)==3:            
            for key,value in Refrigerate_tem_dict.items():
                  for i,j in value.items():
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        lowmean = mean(j2[0:3]) # 低浓度均值
                        medianmean = mean(j2[3:6]) # 中浓度均值
                        highmean = mean(j2[6:9]) # 高浓度均值

                        lowcv = cv(j2[0:3]) # 低浓度cv
                        mediancv = cv(j2[3:6]) # 中浓度cv
                        highcv = cv(j2[6:9]) # 高浓度cv

                        # 添加上述六个值
                        value[i].extend([lowmean,medianmean,highmean,lowcv,mediancv,highcv])
            
            # 计算回收率
            for key,value in Refrigerate_tem_dict.items():
                  # 定义每个化合物0h时的低浓度，中浓度和高浓度均值
                  low0 = 0
                  median0 = 0
                  high0 = 0 
                  for i,j in value.items():
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加中浓度回收率
                              value[i].extend(["/","/","/"])  # 添加高浓度回收率
                              low0 = j2[9] # 设置0h低浓度均值
                              median0 = j2[10] # 设置0h中浓度均值
                              high0 = j2[11] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              medianrecycle1 = new_round(j2[3]/median0*100)
                              medianrecycle2 = new_round(j2[4]/median0*100)
                              medianrecycle3 = new_round(j2[5]/median0*100)
                              highrecycle1 = new_round(j2[6]/high0*100)
                              highrecycle2 = new_round(j2[7]/high0*100)
                              highrecycle3 = new_round(j2[8]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 添加低浓度回收率
                              value[i].extend([medianrecycle1,medianrecycle2,medianrecycle3])  # 添加中浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加高浓度回收率 

      else:
            for key,value in Refrigerate_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        mean1 = mean(j2[0:3]) # 较低浓度均值
                        mean2 = mean(j2[3:6]) # 较高浓度均值

                        cv1 = cv(j2[0:3]) # 较低浓度cv
                        cv2 = cv(j2[3:6]) # 较高浓度cv
                  
                        # 添加上述四个值
                        value[i].extend([mean1,mean2,cv1,cv2])

            # 计算回收率
            for key,value in Refrigerate_tem_dict.items():
                  # 定义每个化合物0h时的较低浓度和较高浓度均值
                  low0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加较低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加叫较高浓度回收率
                              low0 = j2[6] # 设置0h低浓度均值
                              high0 = j2[7] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              highrecycle1 = new_round(j2[3]/high0*100)
                              highrecycle2 = new_round(j2[4]/high0*100)
                              highrecycle3 = new_round(j2[5]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 较添加低浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加较高浓度回收率

      # 2.3 冷冻
      # 三个浓度水平的计算 
      if len(Freeze_conclevel_list)==3:            
            for key,value in Freeze_tem_dict.items():
                  for i,j in value.items():
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        lowmean = mean(j2[0:3]) # 低浓度均值
                        medianmean = mean(j2[3:6]) # 中浓度均值
                        highmean = mean(j2[6:9]) # 高浓度均值

                        lowcv = cv(j2[0:3]) # 低浓度cv
                        mediancv = cv(j2[3:6]) # 中浓度cv
                        highcv = cv(j2[6:9]) # 高浓度cv

                        # 添加上述六个值
                        value[i].extend([lowmean,medianmean,highmean,lowcv,mediancv,highcv])
            
            # 计算回收率
            for key,value in Freeze_tem_dict.items():
                  # 定义每个化合物0h时的低浓度，中浓度和高浓度均值
                  low0 = 0
                  median0 = 0
                  high0 = 0 
                  for i,j in value.items():
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加中浓度回收率
                              value[i].extend(["/","/","/"])  # 添加高浓度回收率
                              low0 = j2[9] # 设置0h低浓度均值
                              median0 = j2[10] # 设置0h中浓度均值
                              high0 = j2[11] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              medianrecycle1 = new_round(j2[3]/median0*100)
                              medianrecycle2 = new_round(j2[4]/median0*100)
                              medianrecycle3 = new_round(j2[5]/median0*100)
                              highrecycle1 = new_round(j2[6]/high0*100)
                              highrecycle2 = new_round(j2[7]/high0*100)
                              highrecycle3 = new_round(j2[8]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 添加低浓度回收率
                              value[i].extend([medianrecycle1,medianrecycle2,medianrecycle3])  # 添加中浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加高浓度回收率 

      else:
            for key,value in Freeze_tem_dict.items():
                  for i,j in value.items():  
                        # 列表中的字符串转数值，为了保证前端有效位数显示准确，原始列表中只能保存字符串格式的数值
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数

                        mean1 = mean(j2[0:3]) # 较低浓度均值
                        mean2 = mean(j2[3:6]) # 较高浓度均值

                        cv1 = cv(j2[0:3]) # 较低浓度cv
                        cv2 = cv(j2[3:6]) # 较高浓度cv
                  
                        # 添加上述四个值
                        value[i].extend([mean1,mean2,cv1,cv2])

            # 计算回收率
            for key,value in Freeze_tem_dict.items():
                  # 定义每个化合物0h时的较低浓度和较高浓度均值
                  low0 = 0
                  high0 = 0 
                  for i,j in value.items(): 
                        j2 = list(map(float,j)) # 列表中的字符串转换为浮点数
                        if i=="0":  # 0h无需计算回收率
                              value[i].extend(["/","/","/"])  # 添加较低浓度回收率
                              value[i].extend(["/","/","/"])  # 添加叫较高浓度回收率
                              low0 = j2[6] # 设置0h低浓度均值
                              high0 = j2[7] # 设置0h高浓度均值
                        elif i!="0":
                              lowrecycle1 = new_round(j2[0]/low0*100)
                              lowrecycle2 = new_round(j2[1]/low0*100)
                              lowrecycle3 = new_round(j2[2]/low0*100)
                              highrecycle1 = new_round(j2[3]/high0*100)
                              highrecycle2 = new_round(j2[4]/high0*100)
                              highrecycle3 = new_round(j2[5]/high0*100)

                              value[i].extend([lowrecycle1,lowrecycle2,lowrecycle3])  # 较添加低浓度回收率
                              value[i].extend([highrecycle1,highrecycle2,highrecycle3])  # 添加较高浓度回收率 


      if len(textlist_special) != 0: 
            return {"Room_conclevel_list": Room_conclevel_list,"Room_tem_dict": Room_tem_dict, 
            "Refrigerate_conclevel_list": Refrigerate_conclevel_list,"Refrigerate_tem_dict": Refrigerate_tem_dict,
            "Freeze_conclevel_list": Freeze_conclevel_list,"Freeze_tem_dict": Freeze_tem_dict,
            "textlist": textlist_special, "serial": len(textlist_special)+1}
      else:
            return {"Room_conclevel_list": Room_conclevel_list,"Room_tem_dict": Room_tem_dict, 
            "Refrigerate_conclevel_list": Refrigerate_conclevel_list,"Refrigerate_tem_dict": Refrigerate_tem_dict,
            "Freeze_conclevel_list": Freeze_conclevel_list,"Freeze_tem_dict": Freeze_tem_dict,
            "textlist": textlist_general, "serial": len(textlist_general)+1}

      