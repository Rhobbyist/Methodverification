from ast import Pass
import xlrd
from report.models import *
from docx import Document
from report.effectnum import *
import re

def PTfileread(files,Detectionplatform,project,platform,manufacturers,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB,Number_of_compounds):

    # 第一步 后台数据抓取（待测物质,可接受标准，单位）
    zqd = Special.objects.get(project=project) 
    pt_special = PTspecial.objects.get(special=zqd)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    PTnorm=[] # 待测物质列表
    PTrange1=[] # 可接受标准一适用范围，与待测物质列表一一对应
    PTstandard1 = [] # 可接受标准一
    PTrange2 = [] # 可接受标准二适用范围，与待测物质列表一一对应
    PTstandard2 = [] # 可接受标准二
    PTunit = Special.objects.get(project=project).unit #单位

    for i in pt_accept:
        PTnorm.append(i.norm)
        PTrange1.append(i.range1) 
        PTstandard1.append(i.accept1) 
        PTrange2.append(i.range2) 
        PTstandard2.append(i.accept2) 

    # 如果没在后台管理系统中设置化合物名称直接返回并提示
    if all(i is None for i in PTnorm):
        error_message="尚未在后台管理系统中设置PT的化合物名称,请设置后重新提交数据!"
        return {"error_message":error_message}

    # Python判断列表中是否为空，包括None
    # if all(i is None for i in PTrange1):
    #     error_message="尚未在后台管理系统中设置PT的可接受标准,请设置后重新提交数据!"
    #     return {"error_message":error_message}

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示
    if manufacturers=="AB":
        if len(normAB)!= Number_of_compounds:
            error_message="未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"  
            return {"error_message":error_message}

    #  第二步:开始文件读取
    '''
    注释:csv,txt,xlsx,docx 4种格式数据读取完毕后,需要生成一个字典PT_dict,数据格式如下：
    print(PT_dict):
    {'MN': [['PT1', 0.49, '±0.075nmol/L'], ['PT10', 3.32, '±15.0%'], ['PT19', 3.31, '±15.0%'], ['PT28', 3.29, '±15.0%']],
    'NMN': [['PT4', 5.96, '±20.0%'], ['PT13', 4.37, '±20.0%'], ['PT22', 4.38, '±20.0%'], ['PT31', 4.25, '±20.0%']], 
    '3-MT': [['PT7', 8.78, '±30.0%'], ['PT16', 1.38, '±30.0%'], ['PT25', 1.38, '±30.0%'], ['PT34', 1.37, '±30.0%']]}
    '''

    # 头部定义相关需要提取生成的结果
    PT_dict={}
    for i in range(len(PTnorm)):
        PT_dict[PTnorm[i]] =[]
    
    # 各仪器平台及各仪器厂家数据读取
    for file in files:
        if platform=="液质":
            if manufacturers =="Agilent":
                # 1 读取csv文件
                file.seek(0)  # https://www.jianshu.com/p/0d15ed85df2b
                file_data = file.read().decode('utf-8')
                lines = file_data.split('\r\n')
                for i in range(len(lines)): 
                    if len(lines[i])!=0:
                        lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                        # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                    else:
                        lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                        del lines[i] #最后一行如为空行，则删除该元素

                # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                sampleindex=0  #实验号索引
                concindex=[] #浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                for j in range(len(lines[1])):  #从第二行开始       
                    if lines[1][j] == "Sample Name" :
                        sampleindex=j
                    elif lines[1][j]  == "Final Conc." :
                        concindex.append(j)

                # 匹配原始数据中与PT相关(实验号前含有"PT-")的行  
                for j in range(len(concindex)):
                    for i in range(len(lines)): # 循环原始数据中的每一行
                        if "PT-" in lines[i][sampleindex]:                            
                            if float(effectnum(lines[i][concindex[j]],digits))<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                PT_dict[PTnorm[j]].append([lines[i][sampleindex],effectnum(lines[i][concindex[j]],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                            elif float(effectnum(lines[i][concindex[j]],digits))>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                PT_dict[PTnorm[j]].append([lines[i][sampleindex],effectnum(lines[i][concindex[j]],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

            elif manufacturers =="Waters":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                norm=[] #化合物列表
                norm_row=[] #化合物所在行
                for j in range(nrows):
                    for i in PTnorm:
                        if i in str(file_data.row_values(j)[0]):
                            norm.append(i)
                            norm_row.append(j)

                nameindex=0
                conindex=0
                for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    if file_data.row_values(norm_row[0]+2)[i]=="Name":
                        nameindex=i
                    elif "实际浓度" in file_data.row_values(norm_row[0]+2)[i]:
                        conindex=i

                for j in range(len(norm)):
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]): 
                            if "PT" in file_data.row_values(i)[nameindex]:
                                if float(file_data.row_values(i)[conindex])<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                                elif float(file_data.row_values(i)[conindex])>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

                    else:
                        for i in range(norm_row[j],nrows): 
                            if "PT" in file_data.row_values(i)[nameindex]:
                                if float(file_data.row_values(i)[conindex])<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                                elif float(file_data.row_values(i)[conindex])>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

            elif manufacturers =="Thermo":
                Thermo = Special.objects.get(project=project) 
                pt_special = PTspecial.objects.get(special=Thermo)
                pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
                PTnorm=[] # 待测物质列表
                for i in pt_accept:
                    PTnorm.append(i.norm)

                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                norm=[] #Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                sheetindex=[] #需要的化合物所在sheet索引列表
                for index in range(len(data.sheet_names())):
                    if data.sheet_names()[index] in PTnorm:
                        norm.append(data.sheet_names()[index])
                        sheetindex.append(index)

                # 循环读取每个sheet工作表,即为每个化合物的表
                for index in range(len(sheetindex)):
                    file_data = data.sheets()[sheetindex[index]]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    #第一行确定samplename和浓度所在列
                    nameindex=0
                    conindex=0
                    for i in range(len(file_data.row_values(0))):  
                        if file_data.row_values(0)[i]=="Compound":
                            nameindex=i
                        elif file_data.row_values(0)[i]=="Calculated Amt":
                            conindex=i

                    for i in range(nrows): 
                        if "PT" in file_data.row_values(i)[nameindex]:
                            if float(file_data.row_values(i)[conindex])<PTrange1[index]: # 小于range1,添加第一个可接受标准
                                PT_dict[PTnorm[index]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard1[index])+" "+PTunit])
                            elif float(file_data.row_values(i)[conindex])>=PTrange2[index]: # 大于range2,添加第二个可接受标准
                                PT_dict[PTnorm[index]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard2[index])+" "+"%"])

            elif manufacturers =="岛津":
                # 3 读取txt
                content= []
                for line in file:
                    content.append(line.decode("GB2312").replace("\r\n", "").split("\t"))

                nameindex=0
                conindex=0 #浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                norm=[] #化合物列表
                norm_row=[] #化合物所在行

                for i in range(len(content[2])):  #第二行确定samplename和浓度所在列
                    if content[2][i]=="数据文件名":
                        nameindex=i 
                    elif content[2][i]=="浓度":
                        conindex=i 

                for i in range(len(content)): 
                    if content[i][0]=="Name": #如果某一行第一列为"Name"，则该行第二列为化合物名称
                        norm.append(content[i][1])
                        norm_row.append(i)

                for j in range(len(norm)):
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]): 
                            if "PT-" in content[i][nameindex]:                             
                                if float(effectnum(content[i][conindex],digits))<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                    PT_dict[PTnorm[j]].append([content[i][nameindex],effectnum(content[i][conindex],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                                elif float(effectnum(content[i][conindex],digits))>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                    PT_dict[PTnorm[j]].append([content[i][nameindex],effectnum(content[i][conindex],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

                    else: #如果是最后一个化合物，索引为该化合物所在行到总行数
                        for i in range(norm_row[j],len(content)): 
                            if "PT-" in content[i][nameindex]:                             
                                if float(effectnum(content[i][conindex],digits))<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                    PT_dict[PTnorm[j]].append([content[i][nameindex],effectnum(content[i][conindex],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                                elif float(effectnum(content[i][conindex],digits))>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                    PT_dict[PTnorm[j]].append([content[i][nameindex],effectnum(content[i][conindex],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])
            
            elif manufacturers =="AB":
                print(normAB)
                print(PTnorm)
                # 获取上传的文件
                file_data = Document(file)

                # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                paragraphs = [] 

                # 若标题长度为0或为换行等，不添加进入标题内容列表
                for p in file_data.paragraphs:
                    if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                        paragraphs.append(p.text.strip())

                print(paragraphs)

                # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                tableindex = []
                for i in range(len(paragraphs)):
                    for j in range(len(ZP_Method_precursor_ion)):
                        if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                            tableindex.append(2*i+1)

                print(tableindex)

                tables = file_data.tables  # 获取文件中的表格集

                # 循环定量表格的索引
                for k in range(len(tableindex)): 
                    tablequantify = tables[tableindex[k]] #获取文件中的相关表格
                    nameindex=0
                    conindex=0

                    # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                    cells = tablequantify._cells
                    ROWS = len(tablequantify.rows)
                    COLUMNS = len(tablequantify.columns)
                    rowdatalist = []  # 每一行的数据列表
                    rowdatagatherlist = []  # 每一行的数据列表汇总列表

                    for i in range(ROWS*COLUMNS):
                        text = cells[i].text.replace("\n", "")
                        text = text.strip()  # 去除空白符
                        if i % COLUMNS != 0 or i == 0:
                            rowdatalist.append(text)
                        else:
                            rowdatagatherlist.append(rowdatalist)
                            rowdatalist = []
                            rowdatalist.append(text)
                    rowdatagatherlist.append(rowdatalist)

                    nameindex = 0
                    concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列
                    
                    # 读取表格的第一行的单元格,判断实验号和浓度索引
                    for i in range(len(rowdatagatherlist[0])):
                        if rowdatagatherlist[0][i] == "Sample Name":
                            nameindex = i
                        elif "Calculated Conc" in rowdatagatherlist[0][i]:
                            concindex = i

                    # 判断是否在后台管理系统中设置了可接受标准

                    # 第一种情况，未设置，前端调用可接受区间模板
                    if all(i is None for i in PTrange1):
                        templates = 1
                        for i in range(len(rowdatagatherlist)): 
                            if "PT" in rowdatagatherlist[i][nameindex]:
                                PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex],effectnum(rowdatagatherlist[i][concindex],digits)])

                    # 第二种情况，已设置，前端调用可接受标准模板
                    else:
                        templates = 2
                        for i in range(len(rowdatagatherlist)): 
                            if "PT" in rowdatagatherlist[i][nameindex]:                       
                                if float(rowdatagatherlist[i][concindex])<PTrange1[k]:
                                    PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex],effectnum(rowdatagatherlist[i][concindex],digits),"±"+" "+str(PTstandard1[0])+" "+PTunit])
                                elif float(rowdatagatherlist[i][concindex])>=PTrange2[k]:
                                    PT_dict[PTnorm[k]].append([rowdatagatherlist[i][nameindex],effectnum(rowdatagatherlist[i][concindex],digits),"±"+" "+str(PTstandard2[0])+" "+"%"])
                
            # 判断每个指标有几个样本
            PT_num = len(PT_dict[PTnorm[0]])

        elif platform=="液相":
            if manufacturers =="Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                norm=[] #化合物列表
                norm_row=[] #化合物所在行
                for j in range(nrows):
                    if file_data.row_values(j)[0] == "化合物:" :  #如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                        norm.append(file_data.row_values(j)[2])
                        norm_row.append(j)

                nameindex=0
                conindex=0
                for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                    if file_data.row_values(norm_row[0]+2)[i]=="样品名称":
                        nameindex=i
                    elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                        conindex=i

                for j in range(len(norm)):
                    if j<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                        for i in range(norm_row[j],norm_row[j+1]): 
                            if "PT" in file_data.row_values(i)[nameindex]:
                                if float(file_data.row_values(i)[conindex])<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                                elif float(file_data.row_values(i)[conindex])>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

                    else:
                        for i in range(norm_row[j],nrows): 
                            if "PT" in file_data.row_values(i)[nameindex]:
                                if float(file_data.row_values(i)[conindex])<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                                elif float(file_data.row_values(i)[conindex])>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                    PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

            PT_num = len(PT_dict[PTnorm[0]])
            print(PT_dict)

        elif platform=="ICP-MS":
            if manufacturers =="Agilent":
                data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                file_data = data.sheets()[0]
                nrows=file_data.nrows
                ncols=file_data.ncols

                # 从第一行确定化合物名称
                norm=[]
                for j in range(ncols):
                    for i in PTnorm:             
                        if i in file_data.row_values(0)[j]:
                            norm.append(i) 

                # 从第二行确定实验号（Sample Name）的索引和化合物浓度索引
                nameindex=0  #实验号索引
                conindex=[] #浓度索引
                for j in range(ncols):       
                    if file_data.row_values(1)[j] == "样品名称":
                        nameindex = j
                    elif file_data.row_values(1)[j] == "浓度 [ ppm ]" or file_data.row_values(1)[j] == "浓度 [ ppb ]":
                        conindex.append(j)

                # 匹配原始数据中与PT相关(实验号前含有"PT-")的行  
                for j in range(len(conindex)):
                    for i in range(2,nrows): # 循环原始数据中的每一行
                        if "PT-" in file_data.row_values(i)[nameindex]:                           
                            if float(effectnum(file_data.row_values(i)[conindex[j]],digits))<PTrange1[j]: # 小于range1,添加第一个可接受标准
                                PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex[j]],digits),"±"+" "+str(PTstandard1[j])+" "+PTunit])
                            elif float(effectnum(file_data.row_values(i)[conindex[j]],digits))>=PTrange2[j]: # 大于range2,添加第二个可接受标准
                                PT_dict[PTnorm[j]].append([file_data.row_values(i)[nameindex],effectnum(file_data.row_values(i)[conindex[j]],digits),"±"+" "+str(PTstandard2[j])+" "+"%"])

            PT_num = len(PT_dict[PTnorm[0]])

        print(PT_dict)
        return {"PT_dict":PT_dict,"PT_num":PT_num,"PTunit":PTunit,"templates":templates}

def Recyclefileread(files,project,platform,manufacturers,Unit,digits,ZP_Method_precursor_ion,ZP_Method_product_ion,normAB,Number_of_compounds):

    # 第一步:后台数据抓取（回收率上下限，最大允许CV）
    id1 = Special.objects.get(project=project).id  
    id2 = Recyclespecial.objects.get(special_id=id1).id

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    if Recyclespecialmethod.objects.filter(recyclespecial=id2): 
        lowvalue=Recyclespecialmethod.objects.get(recyclespecial=id2).lowvalue #回收率下限
        upvalue=Recyclespecialmethod.objects.get(recyclespecial=id2).upvalue #回收率上限
   
    else:
        general = General.objects.get(name="通用性项目")
        recycle_general = Recyclegeneral.objects.get(general=general)
        lowvalue=Recyclegeneralmethod.objects.get(recyclegeneral=recycle_general).lowvalue #回收率下限
        upvalue=Recyclegeneralmethod.objects.get(recyclegeneral=recycle_general).upvalue #回收率上限

    # AB厂家,未在后台管理系统里规范设置定量离子对数值,直接返回并提示
    if manufacturers=="AB":
        if len(normAB)!= Number_of_compounds:
            error_message="未在后台管理系统里规范设置定量离子对数值，请检查并规范设置后重新提交数据!"  
            return {"error_message":error_message}

    #  第二步:开始文件读取

    '''
    数据读取完毕后,需要生成一个字典Recycle_enddict,数据格式如下：
        {'待测物质1':{'sam1': [1.08, 0.44, 0.7, 1.12, 0.72, 0.76, 0.74, 1.01, 0.96, 1.23, 1.28, 0.85], 
        'sam2': [7.57, 8.33, 9.05, 13.69, 14.45, 13.36, 20.06, 19.24, 19.97, 30.72, 31.01, 30.64], 
        'sam3': [6.3, 6.48, 6.4, 9.68, 10.13, 10.52,12.24, 13.37, 14.89, 19.76, 20.77, 20.63]},
        '待测物质2':{'sam1': [1.08, 0.44, 0.7, 1.12, 0.72, 0.76, 0.74, 1.01, 0.96, 1.23, 1.28, 0.85], 
        'sam2': [7.57, 8.33, 9.05, 13.69, 14.45, 13.36, 20.06, 19.24, 19.97, 30.72, 31.01, 30.64], 
        'sam3': [6.3, 6.48, 6.4, 9.68, 10.13, 10.52, 12.24, 13.37, 14.89, 19.76, 20.77, 20.63]} }
    ''' 

    # 1 抓取加标回收率加标数据

    # 头部定义相关需要提取生成的结果
    Rec_dict={} # 加标数据字典
    Reclist=["Rec-1","Rec-2","Rec-3"] # 加标后样本实验号前缀

    for file in files:
        print(file.name)
        if file.name=="加标回收率-加标数据.csv":
            # 读取csv文件
            file.seek(0)  # 此网址查找到的答案:https://www.jianshu.com/p/0d15ed85df2b
            file_data = file.read().decode('gbk')
            lines = file_data.split('\r\n')
            for i in range(len(lines)): 
                if len(lines[i])!=0:
                    lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                    # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                else:
                    lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                    del lines[i] #最后一行如为空行，则删除该元素

            for i in range(len(Reclist)):
                middle_list=[] # 每个本底的加标数据列表
                for j in range(len(lines)):           
                    if Reclist[i] in lines[j][0]:
                        middle_list.extend([lines[j][1],lines[j][2],lines[j][3]])

                Rec_dict[Reclist[i]]=middle_list
            print(Rec_dict)

    # 2 抓取加标回收率原始数据

    # 头部定义相关需要提取生成的结果
    Recycle_enddict_show={} # 加班回收率最终字典(展示数据用,如用户上传了加标数据的文件,需要添加加标数据)
    Recycle_enddict_savedata={} # 加班回收率最终字典(保存数据用,如用户上传了加标数据的文件,不需要添加加标数据,在验证界面点击保存按钮后需要用到此字典)

    for file in files:
        print(file.name)
        if file.name!="加标回收率-加标数据.csv":
            if platform=="液质":
                if manufacturers =="Agilent":
                    # 1 读取csv文件（Agilent）
                    csv_file = file.seek(0)  # 此网址查找到的答案:https://www.jianshu.com/p/0d15ed85df2b
                    file_data = file.read().decode('utf-8')
                    lines = file_data.split('\r\n')
                    for i in range(len(lines)): 
                        if len(lines[i])!=0:
                            lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])  # 以逗号分隔字符串,但忽略双引号内的逗号
                            # lines[i]=lines[i].split(',') # 按逗号分隔后把每一行都变成一个列表
                        else:
                            lines[i]=re.split(r',\s*(?![^"]*\"\,)', lines[i])
                            del lines[i] #最后一行如为空行，则删除该元素

                    # 从第一行确定化合物名称(含有"-Q Results"),并添加进入化合物列表
                    norm=[] #化合物列表
                    for j in range(len(lines[0])):  #从第一行开始
                        if "-Q Results" in lines[0][j]:
                            if lines[0][j].split("-Q")[0][0]!='"':  # 若原始字符串中含有','，切割完后首位会多出一个'"',需去除  
                                norm.append(lines[0][j].split("-Q")[0])
                            else:
                                norm.append(lines[0][j].split("-Q")[0][1:])

                    # 从第二行确定实验号（Sample Name）,浓度（Exp. Conc.）的索引
                    nameindex=0  #实验号索引
                    conindex=[] #浓度索引列表（可能不止一个化合物，因此需要把索引放在一个列表里）
                    for j in range(len(lines[1])):  #从第二行开始       
                        if lines[1][j] == "Sample Name" :
                            nameindex=j
                        elif lines[1][j]  == "Final Conc." :
                            conindex.append(j)

                    # 确定本底数(含有"RecB")
                    RecBlist=[] #本底列表,长度/3即为本底数  
                    for i in range(len(lines)): # 循环原始数据中的每一行,并避免重复添加 
                        # 加标回收率本底实验号命名“RecB-1-1”                
                        if "RecB" in lines[i][nameindex][0:6] and lines[i][nameindex][0:6] not in RecBlist: 
                            RecBlist.append(lines[i][nameindex][0:6])

                    # 匹配原始数据中与加标回收相关的行  
                    for k in range(len(norm)):
                        norm_dict={} #每个化合物数据字典
                        for j in range(len(RecBlist)): #本底列表 
                            RecB_conc=[] #本底浓度列表
                            low=[] #本底加标后低浓度列表
                            median=[] #本底加标后中浓度列表
                            high=[] #本底加标后高浓度列表                 
                            for i in range(len(lines)): # 循环原始数据中的每一行                   
                                if RecBlist[j] in lines[i][nameindex]: 
                                    RecB_conc.append(effectnum(lines[i][conindex[k]],digits)) 
                                elif "L" in lines[i][nameindex] and Reclist[j] in lines[i][nameindex]:
                                    low.append(effectnum(lines[i][conindex[k]],digits))
                                elif "M" in lines[i][nameindex] and Reclist[j] in lines[i][nameindex]:
                                    median.append(effectnum(lines[i][conindex[k]],digits))
                                elif "H" in lines[i][nameindex] and Reclist[j] in lines[i][nameindex]:
                                    high.append(effectnum(lines[i][conindex[k]],digits))

                            norm_dict[RecBlist[j]]=[]
                    
                            norm_dict[RecBlist[j]].extend(RecB_conc)
                            norm_dict[RecBlist[j]].extend(low)
                            norm_dict[RecBlist[j]].extend(median)
                            norm_dict[RecBlist[j]].extend(high)
                            if Rec_dict!={}:
                                norm_dict[RecBlist[j]].extend(Rec_dict[Reclist[j]])
                            else:
                                norm_dict[RecBlist[j]].extend(["","","","","","","","",""])
                            
                        Recycle_enddict_show[norm[k]]=norm_dict

                    print(Recycle_enddict_show)

                elif manufacturers =="Waters":
                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    file_data = data.sheets()[0]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    norm=[] #化合物列表
                    norm_row=[] #化合物所在行
                    for j in range(nrows):
                        for i in PTnorm:
                            if i in str(file_data.row_values(j)[0]):
                                norm.append(i)
                                norm_row.append(j)

                    nameindex=0
                    conindex=0
                    for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                        if file_data.row_values(norm_row[0]+2)[i]=="Name":
                            nameindex=i
                        elif "实际浓度" in file_data.row_values(norm_row[0]+2)[i]:
                            conindex=i

                    # 确定本底数，含有"Recycle"及"background"(以第一个化合物为准确定本底数)
                    background=[] #本底列表,长度/3即为本底数
                    if len(norm)==1: #如果只有一个化合物,则循环第一个化合物所在行到最后一行      
                        for i in range(norm_row[0],nrows):                    
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]: 
                                background.append(file_data.row_values(i)[nameindex])
                    else: #如果有多个化合物,则循环第一个化合物所在行到第二个化合物所在行 
                        for i in range(norm_row[0],norm_row[1]):                    
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]: 
                                background.append(file_data.row_values(i)[nameindex])
        
                    # 匹配原始数据中与加标回收相关的行  
                    for k in range(len(norm)):
                        group_Recycle={} #每个化合物数据字典
                        for j in range(int(len(background)/3)): #本底列表,长度/3即为本底数  
                            background_conc=[] #本底浓度列表
                            low=[] #本底加标后低浓度列表
                            median=[] #本底加标后中浓度列表
                            high=[] #本底加标后高浓度列表                 
                            if k<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                for i in range(norm_row[k],norm_row[k+1]):                  
                                    if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                        background_conc.append(effectnum(file_data.row_values(i)[conindex],digits)) 
                                    elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        low.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        median.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        high.append(effectnum(file_data.row_values(i)[conindex],digits))
                            else:
                                for i in range(norm_row[k],nrows):                  
                                    if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                        background_conc.append(effectnum(file_data.row_values(i)[conindex],digits)) 
                                    elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        low.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        median.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        high.append(effectnum(file_data.row_values(i)[conindex],digits))

                            group_Recycle[Recycle_background[j]]=[]
                            for i in background_conc:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in low:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in median:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in high:
                                group_Recycle[Recycle_background[j]].append(i) 
                        Recycle_enddict[norm[k]]=group_Recycle

                    print(Recycle_enddict)
    
                elif manufacturers =="Thermo":
                    Thermo = Special.objects.get(project=project) 
                    pt_special = PTspecial.objects.get(special=Thermo)
                    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
                    PTnorm=[] # 待测物质列表
                    for i in pt_accept:
                        PTnorm.append(i.norm)

                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    norm=[] #Thermo的原始数据格式为一个化合物一个sheet,获取每个sheet的名字,与PTnorm相等的即为需要的sheet
                    sheetindex=[] #需要的化合物所在sheet索引列表
                    for index in range(len(data.sheet_names())):
                        if data.sheet_names()[index] in PTnorm:
                            norm.append(data.sheet_names()[index])
                            sheetindex.append(index)

                    # 循环读取每个sheet工作表,即为每个化合物的表
                    for index in range(len(sheetindex)):
                        file_data = data.sheets()[sheetindex[index]]
                        nrows=file_data.nrows
                        ncols=file_data.ncols

                        #第一行确定samplename和浓度所在列
                        nameindex=0
                        conindex=0
                        for i in range(len(file_data.row_values(0))):  
                            if file_data.row_values(0)[i]=="Compound":
                                nameindex=i
                            elif file_data.row_values(0)[i]=="Calculated Amt":
                                conindex=i

                        # 确定本底数，含有"Recycle"及"background"(以第一个化合物为准确定本底数)
                        background=[] #本底列表,长度/3即为本底数
                        for i in range(nrows):                    
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]: 
                                background.append(file_data.row_values(i)[nameindex])

                        # 匹配原始数据中与加标回收相关的行  
                        group_Recycle={} #每个化合物数据字典
                        for j in range(int(len(background)/3)): #本底列表,长度/3即为本底数  
                            background_conc=[] #本底浓度列表
                            low=[] #本底加标后低浓度列表
                            median=[] #本底加标后中浓度列表
                            high=[] #本底加标后高浓度列表                 
                            for i in range(nrows):                  
                                if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                    background_conc.append(effectnum(file_data.row_values(i)[conindex],digits)) 
                                elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                    low.append(effectnum(file_data.row_values(i)[conindex],digits))
                                elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                    median.append(effectnum(file_data.row_values(i)[conindex],digits))
                                elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                    high.append(effectnum(file_data.row_values(i)[conindex],digits))

                            group_Recycle[Recycle_background[j]]=[]
                            for i in background_conc:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in low:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in median:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in high:
                                group_Recycle[Recycle_background[j]].append(i) 
                        Recycle_enddict[norm[index]]=group_Recycle

                    print(Recycle_enddict)

                elif manufacturers =="岛津":
                    # 读取txt
                    content= []
                    for line in file:
                        content.append(line.decode("GB2312").replace("\r\n", "").split("\t"))

                    nameindex=0
                    conindex=0 #浓度索引，岛津的数据格式决定每个化合物的浓度所在列一定是同一列
                    norm=[] #化合物列表
                    norm_row=[] #化合物所在行

                    for i in range(len(content[2])):  #第二行确定samplename和浓度所在列
                        if content[2][i]=="数据文件名":
                            nameindex=i 
                        elif content[2][i]=="浓度":
                            conindex=i 

                    for i in range(len(content)): 
                        if content[i][0]=="Name": #如果某一行第一列为"Name"，则该行第二列为化合物名称
                            norm.append(content[i][1])
                            norm_row.append(i)
                    
                    # 确定本底数，含有"Recycle"及"background"(以第一个化合物为准确定本底数)
                    background=[] #本底列表,长度/3即为本底数
                    if len(norm)==1: #如果只有一个化合物      
                        for i in range(norm_row[0],len(content)):                    
                            if "Recycle" in content[i][nameindex] and "background" in content[i][nameindex]: 
                                background.append(content[i][nameindex])
                    else:
                        for i in range(norm_row[0],norm_row[1]):                    
                            if "Recycle" in content[i][nameindex] and "background" in content[i][nameindex]: 
                                background.append(content[i][nameindex])
        
                    # 匹配原始数据中与加标回收相关的行  
                    for k in range(len(norm)):
                        group_Recycle={} #每个化合物数据字典
                        for j in range(int(len(background)/3)): #本底列表,长度/3即为本底数  
                            background_conc=[] #本底浓度列表
                            low=[] #本底加标后低浓度列表
                            median=[] #本底加标后中浓度列表
                            high=[] #本底加标后高浓度列表                 
                            if k<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                for i in range(norm_row[k],norm_row[k+1]):                  
                                    if "background" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                        background_conc.append(effectnum(content[i][conindex],digits)) 
                                    elif "low" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:
                                        low.append(effectnum(content[i][conindex],digits))
                                    elif "median" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:
                                        median.append(effectnum(content[i][conindex],digits))
                                    elif "high" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:
                                        high.append(effectnum(content[i][conindex],digits))
                            else:
                                for i in range(norm_row[k],len(content)):                  
                                    if "background" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                        background_conc.append(effectnum(content[i][conindex],digits)) 
                                    elif "low" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:
                                        low.append(effectnum(content[i][conindex],digits))
                                    elif "median" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:
                                        median.append(effectnum(content[i][conindex],digits))
                                    elif "high" in content[i][nameindex] and Recycle_background[j] in content[i][nameindex]:
                                        high.append(effectnum(content[i][conindex],digits))

                            group_Recycle[Recycle_background[j]]=[]
                            for i in background_conc:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in low:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in median:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in high:
                                group_Recycle[Recycle_background[j]].append(i) 
                        Recycle_enddict[norm[k]]=group_Recycle

                    print("111")
                    print(Recycle_enddict)
                    
                elif manufacturers =="AB":
                    # 定义化合物列表，列表统一命名为norm
                    norm = normAB

                    # 获取上传的文件
                    file_data = Document(file)

                    # 每个表格最上方的标题内容列表，含有母离子和子离子的信息。需依此及母离子和子离子列表判断table索引
                    paragraphs = [] 

                    # 若标题长度为0或为换行等，不添加进入标题内容列表
                    for p in file_data.paragraphs:
                        if len(p.text) != 0 and p.text != "\n" and len(p.text.strip()) != 0:
                            paragraphs.append(p.text.strip())

                    # 确定table索引，母离子和子离子都与后台管理系统中设置的数值相同才证明是需要读取的定量表格
                    tableindex = []
                    for i in range(len(paragraphs)):
                        for j in range(len(ZP_Method_precursor_ion)):
                            if ZP_Method_precursor_ion[j] in paragraphs[i] and ZP_Method_product_ion[j] in paragraphs[i]:
                                tableindex.append(2*i+1)

                    tables = file_data.tables  # 获取文件中的表格集

                    # 循环定量表格的索引
                    for k in range(len(tableindex)):
                        # 获取文件中的定量表格
                        tablequantify = tables[tableindex[k]] 

                        # 先把表格里的所有数据取出来放进一个列表中，读取速度会比直接读表格快很多
                        cells = tablequantify._cells
                        ROWS = len(tablequantify.rows)
                        COLUMNS = len(tablequantify.columns)
                        rowdatalist = []  # 每一行的数据列表
                        rowdatagatherlist = []  # 每一行的数据列表汇总列表

                        for i in range(ROWS*COLUMNS):
                            text = cells[i].text.replace("\n", "")
                            text = text.strip()  # 去除空白符
                            if i % COLUMNS != 0 or i == 0:
                                rowdatalist.append(text)
                            else:
                                rowdatagatherlist.append(rowdatalist)
                                rowdatalist = []
                                rowdatalist.append(text)
                        rowdatagatherlist.append(rowdatalist)

                        nameindex = 0
                        concindex = 0  # 浓度索引，AB的数据格式决定每个化合物的浓度所在列一定是同一列

                        # 读取表格的第一行的单元格,判断实验号和浓度索引
                        for i in range(len(rowdatagatherlist[0])):
                            if rowdatagatherlist[0][i] == "Sample Name":
                                nameindex = i
                            elif "Calculated Conc" in rowdatagatherlist[0][i]:
                                concindex = i
                        
                        # 确定本底数(含有"RecB")
                        RecBlist=[] #本底列表,长度/3即为本底数  
                        for i in range(len(rowdatagatherlist)): # 循环原始数据中的每一行,并避免重复添加 
                            # 加标回收率本底实验号命名“RecB-1-1”                
                            if "RecB" in rowdatagatherlist[i][nameindex][0:6] and rowdatagatherlist[i][nameindex][0:6] not in RecBlist: 
                                RecBlist.append(rowdatagatherlist[i][nameindex][0:6])

                        # 匹配原始数据中与加标回收相关的行  
                        for k in range(len(norm)):
                            Recycle_enddict_show[norm[k]]={} # 数据展示字典
                            Recycle_enddict_savedata[norm[k]]={} # 数据保存字典

                            for j in range(len(RecBlist)): #本底列表 
                                RecB_conc=[] #本底浓度列表
                                low=[] #本底加标后低浓度列表
                                median=[] #本底加标后中浓度列表
                                high=[] #本底加标后高浓度列表                 
                                for i in range(len(rowdatagatherlist)): # 循环原始数据中的每一行                   
                                    if RecBlist[j] in rowdatagatherlist[i][nameindex]: 
                                        RecB_conc.append(effectnum(rowdatagatherlist[i][concindex],digits)) 
                                    elif "L" in rowdatagatherlist[i][nameindex] and Reclist[j] in rowdatagatherlist[i][nameindex]:
                                        low.append(effectnum(rowdatagatherlist[i][concindex],digits))
                                    elif "M" in rowdatagatherlist[i][nameindex] and Reclist[j] in rowdatagatherlist[i][nameindex]:
                                        median.append(effectnum(rowdatagatherlist[i][concindex],digits))
                                    elif "H" in rowdatagatherlist[i][nameindex] and Reclist[j] in rowdatagatherlist[i][nameindex]:
                                        high.append(effectnum(rowdatagatherlist[i][concindex],digits))

                                # 数据展示字典
                                Recycle_enddict_show[norm[k]][RecBlist[j]]=[]               
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(RecB_conc)
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(low)
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(median)
                                Recycle_enddict_show[norm[k]][RecBlist[j]].extend(high)

                                if Rec_dict!={}:
                                    Recycle_enddict_show[norm[k]][RecBlist[j]].extend(Rec_dict[Reclist[j]])
                                else:
                                    Recycle_enddict_show[norm[k]][RecBlist[j]].extend(["","","","","","","","",""])

                                # 数据保存字典
                                Recycle_enddict_savedata[norm[k]][RecBlist[j]]=[]               
                                Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(RecB_conc)
                                Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(low)
                                Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(median)
                                Recycle_enddict_savedata[norm[k]][RecBlist[j]].extend(high)                            

            elif platform=="液相":
                if manufacturers =="Agilent":
                    data = xlrd.open_workbook(filename=None, file_contents=file.read())  # 读取表格
                    file_data = data.sheets()[0]
                    nrows=file_data.nrows
                    ncols=file_data.ncols

                    norm=[] #化合物列表
                    norm_row=[] #化合物所在行
                    for j in range(nrows):
                        if file_data.row_values(j)[0] == "化合物:" :  #如果某一行的第一个元素为“化合物”，则添加第三个元素进入化合物列表
                            norm.append(file_data.row_values(j)[2])
                            norm_row.append(j)

                    nameindex=0
                    conindex=0
                    for i in range(len(file_data.row_values(norm_row[0]+2))):  #第一个化合物表格确定samplename和浓度所在列，norm_row[0]为第一个化合物所在行，+2是该化合物表格位于该化合物所在行的下两行
                        if file_data.row_values(norm_row[0]+2)[i]=="样品名称":
                            nameindex=i
                        elif "含量" in file_data.row_values(norm_row[0]+2)[i]:
                            conindex=i

                    # 确定本底数，含有"Recycle"及"background"(以第一个化合物为准确定本底数)
                    background=[] #本底列表,长度/3即为本底数
                    if len(norm)==1: #如果只有一个化合物,则循环第一个化合物所在行到最后一行      
                        for i in range(norm_row[0],nrows):                    
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]: 
                                background.append(file_data.row_values(i)[nameindex])
                                
                    else: #如果有多个化合物,则循环第一个化合物所在行到第二个化合物所在行 
                        for i in range(norm_row[0],norm_row[1]):                    
                            if "Recycle" in file_data.row_values(i)[nameindex] and "background" in file_data.row_values(i)[nameindex]: 
                                background.append(file_data.row_values(i)[nameindex])
        
                    # 匹配原始数据中与加标回收相关的行  
                    for k in range(len(norm)):
                        group_Recycle={} #每个化合物数据字典
                        for j in range(int(len(background)/3)): #本底列表,长度/3即为本底数  
                            background_conc=[] #本底浓度列表
                            low=[] #本底加标后低浓度列表
                            median=[] #本底加标后中浓度列表
                            high=[] #本底加标后高浓度列表                 
                            if k<len(norm)-1: #如果不是最后一个化合物，索引为该化合物所在行到后一个化合物所在行
                                for i in range(norm_row[k],norm_row[k+1]):                  
                                    if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                        background_conc.append(effectnum(file_data.row_values(i)[conindex],digits)) 
                                    elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        low.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        median.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        high.append(effectnum(file_data.row_values(i)[conindex],digits))
                            else:
                                for i in range(norm_row[k],nrows):                  
                                    if "background" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:  # 如果实验号命名方式匹配上，则在相应列表中添加相应数据
                                        background_conc.append(effectnum(file_data.row_values(i)[conindex],digits)) 
                                    elif "low" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        low.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "median" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        median.append(effectnum(file_data.row_values(i)[conindex],digits))
                                    elif "high" in file_data.row_values(i)[nameindex] and Recycle_background[j] in file_data.row_values(i)[nameindex]:
                                        high.append(effectnum(file_data.row_values(i)[conindex],digits))

                            group_Recycle[Recycle_background[j]]=[]
                            for i in background_conc:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in low:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in median:
                                group_Recycle[Recycle_background[j]].append(i)
                            for i in high:
                                group_Recycle[Recycle_background[j]].append(i) 
                        Recycle_enddict_show[norm[k]]=group_Recycle

                    print(Recycle_enddict)

    print(Recycle_enddict_show)
               
    return {"Recycle_enddict_show":Recycle_enddict_show,"Recycle_enddict_savedata":Recycle_enddict_savedata,"Unit":Unit,"lowvalue":lowvalue,"upvalue":upvalue}

# PT数据关联进入最终报告
def related_PT(id): 
    # 第一步：后台描述性内容数据提取
    # 1 根据id找到项目
    project = ReportInfo.objects.get(id=id).project

    # 2 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    # 特殊参数设置描述性内容
    textlist_special = []
    try:
        special_1 = Special.objects.get(project=project) 
        special_2 = PTspecial.objects.get(special=special_1)           
        if PTspecialtexts.objects.filter(pTspecial=special_2).count()>0:
            text_special = PTspecialtexts.objects.filter(pTspecial=special_2)  
            for i in text_special:
                textlist_special.append(i.text)
    except:
        pass

    # 3 通用数据抓取
    # 描述性内容
    textlist_general = [] 
    general_1 = General.objects.get(name="通用性项目") #通用参数设置描述性内容
    general_2 = PTgeneral.objects.get(general=general_1)
    text_general = PTgeneraltexts.objects.filter(pTgeneral=general_2)      
    for i in text_general:
        textlist_general.append(i.text)

    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project)
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist = []  # 每个化合物有效位数列表
    Digitsdict = {}  # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist == [] or Digitslist[0] == "":  # 如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm] = i.digits


    # 第二步：报告数据提取

    # 定义需要生成的字典
    PT_dict = {}  # 最终需要的字典

    try:
        # 1 基础数据抓取
        PT_data = PT.objects.filter(reportinfo_id=id)
        
        PT_norm=[]
        for i in PT_data:
            if i.norm not in PT_norm:
                PT_norm.append(i.norm)
        
        for i in PT_norm:
            middle_list=[]  # 每个化合物的数据列表
            middle_table = PT.objects.filter(reportinfo_id=id,norm=i)
            for j in middle_table:
                #没有为每个化合物单独设置有效位数，则调用通用性设置
                if Digitsdict=={} or list(Digitsdict.values())[0]==None: 
                    rowlist = []  # 每一行的小列表
                    rowlist.append(j.Experimentnum)
                    rowlist.append(j.value)
                    rowlist.append(j.target)
                    rowlist.append(j.received)
                    rowlist.append(j.bias)
                    middle_list.append(rowlist)
                #为每个化合物单独设置了有效位数，则调用每个化合物的设置
                else:
                    rowlist=[]
                    rowlist.append(j.Experimentnum)
                    rowlist.append(effectnum(j.value,Digitsdict[i]))                   
                    rowlist.append(j.target)
                    rowlist.append(j.received)
                    rowlist.append(j.bias)
                    middle_list.append(rowlist)
            PT_dict[i]=middle_list
        
        if len(textlist_special)!=0:
            return {"PT_dict":PT_dict,"textlist":textlist_special,"serial":len(textlist_special)+1}
        else:
            return {"PT_dict":PT_dict,"textlist":textlist_general,"serial":len(textlist_general)+1}

    except:
        pass

# 加标回收据关联进入最终报告
def related_recycle(id):  
    # 第一步：后台描述性内容数据提取

    # 根据id找到项目
    project=ReportInfo.objects.get(id=id).project

    # 优先查找特殊参数设置里是否有数据，如有就调用，没有则调用通用性参数设置里的数据
    # 特殊数据抓取
    zqd_special = Special.objects.get(project=project)   
    Recycle_special = Recyclespecial.objects.get(special=zqd_special) 
    textlist_special = [] #特殊参数设置描述性内容
    if Recyclespecialtexts.objects.filter(recyclespecial=Recycle_special).count()>0: 
        text_special = Recyclespecialtexts.objects.filter(recyclespecial=Recycle_special)     
        for i in text_special:
            textlist_special.append(i.text)
    
    # 通用数据抓取
    zqd_general = General.objects.get(name="通用性项目") 
    Recycle_general = Recyclegeneral.objects.get(general=zqd_general)
    text_general = Recyclegeneraltexts.objects.filter(recyclegeneral=Recycle_general)   
    textlist_general = []
    for i in text_general:
        textlist_general.append(i.text)
    
    # 查找是否单独设置了每个化合物的有效位数
    DIGITS_TABLE = Special.objects.get(project=project) 
    pt_special = PTspecial.objects.get(special=DIGITS_TABLE)
    pt_accept = PTspecialaccept.objects.filter(pTspecial=pt_special)
    Digitslist=[] # 每个化合物有效位数列表
    Digitsdict={} # 每个化合物有效位数字典

    for i in pt_accept:
        Digitslist.append(i.digits)

    if Digitslist==[] or Digitslist[0]=="": #如果全部没设置或者只是单位没设置
        pass
    else:
        for i in pt_accept:
            Digitsdict[i.norm]=i.digits

    
    # 第二步：报告数据提取   
    datasam = RECYCLE.objects.filter(reportinfo_id=id)

    if datasam:
        Recycle_endreport_dict={}  #最终需要的字典
        Recycle_endreport_norm=[]  #去重后的指标列表
        sam=[] #去重后的本底列表

        for item in datasam:
            if item.norm not in Recycle_endreport_norm:
                Recycle_endreport_norm.append(item.norm)

        for item in datasam:
            if item.Experimentnum not in sam:
                sam.append(item.Experimentnum)
        
        Recycle_endconclusion=""
        for i in Recycle_endreport_norm:
            Recycle={}  #每个化合物的字典
            for j in sam:
                data = RECYCLE.objects.filter(reportinfo_id=id,norm=i,Experimentnum=j)
                sum=0
                recycle_sam=[]
                recycle_list=[] #每个化合物的回收率列表，方便提取最大最小值  "maxrecycle":str(max(recycle_list))+"%","minrecycle":str(min(recycle_list))+"%"
                for item in data: 
                    #没有为每个化合物单独设置有效位数，则调用通用性设置
                    if Digitsdict=={} or list(Digitsdict.values())[0]==None:                         
                        recycle_sam.append(item.sam_conc)
                        sum+=float(item.sam_conc)
                        recycle_sam.append(item.theory_conc)
                        recycle_sam.append(item.level)
                        recycle_sam.append(item.end_conc1)
                        recycle_sam.append(item.end_conc2)
                        recycle_sam.append(item.end_conc3)
                        recycle_sam.append(item.end_recycle1)
                        recycle_sam.append(item.end_recycle2)
                        recycle_sam.append(item.end_recycle3)
                        recycle_list.append(float(item.end_recycle1))
                        recycle_list.append(float(item.end_recycle2))
                        recycle_list.append(float(item.end_recycle3))
                    #为每个化合物单独设置了有效位数，则调用每个化合物的设置
                    else:
                        recycle_sam.append(effectnum(item.sam_conc,Digitsdict[i]))
                        sum+=float(item.sam_conc)
                        recycle_sam.append(item.theory_conc)
                        recycle_sam.append(item.level)
                        recycle_sam.append(effectnum(item.end_conc1,Digitsdict[i]))
                        recycle_sam.append(effectnum(item.end_conc2,Digitsdict[i]))
                        recycle_sam.append(effectnum(item.end_conc3,Digitsdict[i]))
                        recycle_sam.append(item.end_recycle1)
                        recycle_sam.append(item.end_recycle2)
                        recycle_sam.append(item.end_recycle3)
                        recycle_list.append(float(item.end_recycle1))
                        recycle_list.append(float(item.end_recycle2))
                        recycle_list.append(float(item.end_recycle3))
                recycle_sam.append(new_round(sum/3,1))
                Recycle[j]=recycle_sam
            Recycle_endreport_dict[i]=Recycle
            Recycle_endconclusion=Recycle_endconclusion+i+"在"+str(min(recycle_list))+"%"+"~"+str(max(recycle_list))+"%"+"范围内, "

        if len(textlist_special)!=0:
            return {"Recycle_endreport_dict":Recycle_endreport_dict,"textlist":textlist_special,"serial":len(textlist_special)+1,"Recycle_endconclusion":Recycle_endconclusion}
        else:
            return {"Recycle_endreport_dict":Recycle_endreport_dict,"textlist":textlist_general,"serial":len(textlist_general)+1,"Recycle_endconclusion":Recycle_endconclusion}

        
